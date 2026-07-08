from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/issamouladsmane/Desktop/full_stuck_ticket_gestion")
DOCS_DIR = ROOT / "docs"
OUTPUTS_DIR = ROOT / "outputs"
DOCS_DIR.mkdir(exist_ok=True)
OUTPUTS_DIR.mkdir(exist_ok=True)

MD_PATH = DOCS_DIR / "LIVRABLE_REPRISE_ANALYSE_DEPLOIEMENT.md"
DOCX_PATH = OUTPUTS_DIR / "livrable-reprise-analyse-deploiement-agce-crm.docx"


TITLE = "Livrable de reprise, analyse et déploiement"
SUBTITLE = "AGCE CRM - Système de gestion des tickets, contacts, e-mails et réunions"
DATE = "08 juillet 2026"


def source_inventory() -> list[str]:
    excluded_dirs = {".git", "node_modules", "dist", "outputs", ".agents", ".codex"}
    excluded_names = {".DS_Store", ".env"}
    files: list[str] = []
    for path in ROOT.rglob("*"):
        if not path.is_file():
            continue
        rel = path.relative_to(ROOT)
        parts = set(rel.parts)
        if parts & excluded_dirs:
            continue
        if rel.name in excluded_names:
            continue
        if rel.suffix.lower() in {".png", ".jpg", ".jpeg", ".gif", ".woff", ".woff2", ".pdf", ".docx"}:
            continue
        files.append(str(rel))
    return sorted(files)


ROUTE_ROWS = [
    ("POST", "/api/auth/login", "Connexion utilisateur, génération du JWT", "Public + login limiter"),
    ("POST", "/api/auth/logout", "Invalidation locale du token", "Token requis"),
    ("GET", "/api/employees/me", "Profil de l'utilisateur connecté", "Tous rôles authentifiés"),
    ("GET", "/api/employees/GetAllEmps", "Liste des employés", "ADMIN"),
    ("POST", "/api/employees/InsertEmp", "Création employé", "ADMIN"),
    ("PUT", "/api/employees/EditEmp/:id", "Modification employé ou profil personnel", "ADMIN ou soi-même"),
    ("PUT", "/api/employees/ChangePassword/:id", "Changement mot de passe employé", "ADMIN"),
    ("DELETE", "/api/employees/DeleteEmp/:id", "Suppression employé", "ADMIN"),
    ("GET", "/api/organizations", "Liste organisations", "SD, Manager"),
    ("POST/PUT/DELETE", "/api/organizations", "CRUD organisations", "SD"),
    ("GET", "/api/contacts", "Liste contacts", "SD, Manager"),
    ("POST/PUT/DELETE", "/api/contacts", "CRUD contacts", "SD"),
    ("GET", "/api/tickets", "Liste tickets filtrée par rôle", "Tous rôles authentifiés"),
    ("POST", "/api/tickets", "Création ticket et room associée", "SD"),
    ("GET", "/api/tickets/:ticketId", "Détail ticket", "Rôle autorisé"),
    ("PUT", "/api/tickets/:ticketId/assign", "Affectation vers IT ou PKI", "SD"),
    ("PUT", "/api/tickets/:ticketId/status", "Mise à jour statut ticket", "SD"),
    ("PUT", "/api/tickets/:ticketId/resolve", "Validation résolution finale", "SD"),
    ("GET/POST", "/api/tickets/:ticketId/comments", "Commentaires et propositions de résolution", "Rôle autorisé"),
    ("GET/PATCH", "/api/rooms", "Rooms de discussion et lecture", "Rôle autorisé"),
    ("GET/POST/PUT/DELETE", "/api/meetings", "Réunions, invitations et décisions", "Lecture selon rôle, écriture SD"),
    ("GET/POST/PUT/DELETE", "/api/meeting-rooms", "Salles de réunion", "Lecture SD/Manager, écriture SD"),
    ("GET/POST/PATCH", "/api/client-emails", "Inbox client et synchronisation Gmail", "SD"),
    ("GET", "/api/dashboard/sd", "Dashboard Service Delivery", "SD"),
    ("GET", "/api/dashboard/manager", "Dashboard Manager", "Manager"),
    ("GET", "/api/dashboard/manager/analytics", "Analytics manager filtrables", "Manager"),
]


ENV_ROWS = [
    ("PORT", "Port HTTP du backend", "2300 en local"),
    ("JWT_SECRET", "Secret de signature JWT", "Secret fort, jamais commit"),
    ("JWT_EXPIRES_IN", "Durée de vie du token", "1d par défaut"),
    ("FRONTEND_URL", "Origine autorisée CORS", "http://localhost:5173 en local"),
    ("JSON_BODY_LIMIT", "Limite taille JSON Express", "1mb"),
    ("RATE_LIMIT_WINDOW_MS", "Fenêtre de rate limit globale", "900000 ms"),
    ("RATE_LIMIT_MAX_REQUESTS", "Maximum requêtes globales", "300"),
    ("LOGIN_RATE_LIMIT_WINDOW_MS", "Fenêtre de blocage login", "900000 ms"),
    ("LOGIN_RATE_LIMIT_MAX_REQUESTS", "Tentatives login échouées autorisées", "10"),
    ("DB_HOST / DB_NAME / DB_USER / DB_PWD", "Connexion MySQL", "Secrets backend uniquement"),
    ("GMAIL_IMAP_USER", "Boîte Gmail SD surveillée", "Adresse inbox SD"),
    ("GMAIL_APP_PASSWORD", "Mot de passe d'application Gmail", "Secret backend uniquement"),
    ("GMAIL_SYNC_INTERVAL_MS", "Fréquence sync Gmail", "60000 ms"),
    ("ADMIN_INITIAL_*", "Bootstrap admin local", "À utiliser une seule fois"),
    ("ENABLE_ADMIN_BOOTSTRAP", "Activation bootstrap admin au démarrage", "false après création"),
    ("VITE_API_URL", "URL API frontend", "http://localhost:2300/api"),
    ("VITE_SOCKET_URL", "URL Socket.IO frontend", "http://localhost:2300"),
]


DB_ROWS = [
    ("services", "Référentiel des rôles/services: IT, SD, MANAGER, ADMIN, PKI."),
    ("employees", "Comptes employés, identifiants, mot de passe haché, service et statut actif/inactif."),
    ("organizations", "Organisations clientes avec secteur, e-mail, téléphone, adresse et statut."),
    ("contacts", "Contacts rattachables aux organisations, type, coordonnées et statut."),
    ("tickets", "Demandes/tickets, description, statut, résolution, créateur et client/organisation."),
    ("rooms", "Salon de discussion associé à un ticket, avec services autorisés."),
    ("messages", "Historique des messages temps réel par room."),
    ("comments", "Commentaires et propositions de résolution liées aux tickets."),
    ("ticket_assignment_history", "Traçabilité des affectations SD vers IT/PKI."),
    ("meeting_rooms", "Référentiel des salles de réunion."),
    ("meetings", "Réunions, organisateur, invité, ticket associé, décision et motif de rejet."),
    ("client_emails", "E-mails clients importés/reçus pour le Service Delivery."),
    ("client_email_attachments", "Pièces jointes filtrées et conservées pour les e-mails clients."),
    ("client_email_reads", "Suivi lu/non-lu par employé pour l'inbox."),
    ("room_message_reads", "Suivi lecture des rooms par employé."),
]


FUNCTIONAL_MODULES = [
    ("Authentification et session", "Connexion par identifiant/mot de passe, JWT signé, persistance localStorage, redirection selon rôle et invalidation au logout."),
    ("Administration", "Création, activation/désactivation, modification, changement de mot de passe et suppression des employés par ADMIN."),
    ("Service Delivery", "Gestion des organisations/contacts, création de tickets, affectation aux équipes, pilotage de la résolution, inbox e-mail et réunions."),
    ("Manager", "Vue de supervision: dashboards, analytics, contacts/tickets/messages/meetings en lecture contrôlée."),
    ("PKI et IT", "Traitement des tickets assignés, accès aux rooms autorisées, messages et propositions de résolution."),
    ("Tickets", "Cycle de vie: Pending, In Progress, Warning, Critical, Resolved. Un ticket résolu est verrouillé."),
    ("Messagerie temps réel", "Socket.IO, rooms par ticket, contrôle d'accès côté serveur et historique MySQL."),
    ("E-mails clients", "Import IMAP Gmail, création d'e-mails entrants, filtrage des pièces jointes, marquage lu."),
    ("Meetings", "Planification SD, réponse invité Accepté/Rejeté, visibilité selon rôle."),
    ("Dashboards", "Indicateurs SD et Manager, statistiques, files d'attente, activité récente et analytics."),
]


TECH_STACK = [
    ("Frontend", "React 19, Vite 8, React Router DOM, Recharts, Socket.IO Client, lucide-react, react-hot-toast, Tailwind/shadcn utilities."),
    ("Backend", "Node.js, Express 5, MySQL2, JWT, bcrypt, dotenv, helmet, cors, express-rate-limit, Socket.IO, ImapFlow, mailparser."),
    ("Base de données", "MySQL relationnel avec schéma applicatif, clés étrangères, historiques et tables de lecture."),
    ("Sécurité", "JWT, hachage bcrypt, contrôle rôles backend, CORS strict, Helmet, rate limiting, filtrage fichiers attachés, secrets `.env`."),
]


TEST_CHECKS = [
    ("Frontend lint", "cd frontend_ticket_gestion && npm run lint -- --quiet"),
    ("Frontend build", "cd frontend_ticket_gestion && npm run build"),
    ("Backend syntax", "node --check backend_ticket_gestion/src/app.js et contrôle des fichiers JS."),
    ("Health backend", "GET http://localhost:2300/test -> {\"message\":\"test ok\"}"),
    ("API base", "GET http://localhost:2300/api -> {\"message\":\"API is working\"}"),
    ("Flux API", "Login, profil, employés, organisations, contacts, tickets, meetings, dashboards, e-mails."),
]


def md_table(headers: list[str], rows: list[tuple[str, ...]]) -> str:
    lines = ["| " + " | ".join(headers) + " |", "| " + " | ".join(["---"] * len(headers)) + " |"]
    for row in rows:
        lines.append("| " + " | ".join(str(cell).replace("\n", "<br>") for cell in row) + " |")
    return "\n".join(lines)


def build_markdown() -> str:
    files = source_inventory()
    api_rows = [(m, p, d, r) for m, p, d, r in ROUTE_ROWS]
    modules_rows = [(m, d) for m, d in FUNCTIONAL_MODULES]
    db_rows = [(t, d) for t, d in DB_ROWS]
    env_rows = [(k, d, e) for k, d, e in ENV_ROWS]
    stack_rows = [(a, b) for a, b in TECH_STACK]
    test_rows = [(a, b) for a, b in TEST_CHECKS]

    source_list = "\n".join(f"- `{item}`" for item in files)

    return f"""# {TITLE}

**Projet :** {SUBTITLE}  
**Date :** {DATE}  
**Objet :** livrable de reprise, d'analyse, de configuration et de déploiement sur plateforme de test.

## 1. Résumé exécutif

AGCE CRM est une application web de gestion de tickets et de relation client construite autour d'un backend Express/MySQL et d'un frontend React/Vite. Elle couvre la réception des demandes, la gestion des contacts et organisations, la création et l'affectation des tickets, les échanges internes en temps réel, les réunions, les dashboards et la synchronisation d'e-mails clients.

Le présent livrable donne à une nouvelle équipe toutes les informations nécessaires pour reprendre le code, comprendre le fonctionnement métier, configurer les environnements, déployer la solution et réaliser les premiers contrôles sur une plateforme de test.

## 2. Périmètre du livrable

- Code source complet de l'application présent dans le dépôt local.
- Documentation fonctionnelle par rôle et par module.
- Documentation technique: architecture, modules, API, base de données, sécurité et temps réel.
- Prérequis d'installation et de configuration.
- Procédure de déploiement local/test.
- Checklist de reprise, contrôle qualité et points d'attention.

## 3. Code source complet livré

Le code source complet est livré dans le dépôt `full_stuck_ticket_gestion`. Les dossiers importants sont:

- `backend_ticket_gestion/`: API Express, logique métier, accès MySQL, Socket.IO, synchronisation Gmail et configuration serveur.
- `frontend_ticket_gestion/`: application React/Vite, routes, dashboards, composants métier et appels API.
- `package.json`: scripts racine pour lancer backend et frontend ensemble.
- `backend_ticket_gestion/.env.example` et `frontend_ticket_gestion/.env.example`: modèles de configuration sans secrets.
- `backend_ticket_gestion/issam.session.sql`: script SQL historique de création de base et tables principales.

> Les vrais fichiers `.env`, les secrets, `node_modules`, `dist` et les outputs générés ne doivent pas être considérés comme du code source à publier.

### Inventaire du code source maintenable

{source_list}

## 4. Architecture générale

```mermaid
flowchart LR
  U[Utilisateur] --> F[Frontend React/Vite]
  F -->|HTTP JSON + JWT| API[Backend Express]
  F <-->|Socket.IO| RT[Serveur temps réel]
  API --> S[Services métier]
  RT --> S
  S --> R[Repositories SQL]
  R --> DB[(MySQL)]
  Gmail[Boîte Gmail SD] -->|IMAP| Mail[Service Gmail Sync]
  Mail --> DB
```

Le frontend consomme l'API via `VITE_API_URL` et les échanges temps réel via `VITE_SOCKET_URL`. Le backend centralise les contrôles de sécurité, les rôles, la validation métier, la persistance MySQL et les sockets.

## 5. Technologies

{md_table(["Couche", "Technologies et rôle"], stack_rows)}

## 6. Documentation fonctionnelle

{md_table(["Module", "Description fonctionnelle"], modules_rows)}

### Rôles et permissions

- `ADMIN`: administration des employés et accès à l'interface admin.
- `SD`: Service Delivery, rôle principal de création/pilotage des tickets, contacts, organisations, e-mails et meetings.
- `Manager`: supervision et lecture contrôlée via dashboards, tickets, messages, contacts et meetings.
- `PKI`: traitement des tickets et conversations assignés au service PKI.
- `IT`: traitement des tickets et conversations assignés au service IT.

Les permissions sont appliquées côté frontend pour l'expérience utilisateur et côté backend pour empêcher les accès directs non autorisés.

## 7. Documentation technique backend

Le backend démarre depuis `backend_ticket_gestion/src/index.js`, charge les variables `.env`, initialise le schéma complémentaire, crée éventuellement l'admin initial si `ENABLE_ADMIN_BOOTSTRAP=true`, démarre Socket.IO puis écoute sur `PORT`.

`backend_ticket_gestion/src/app.js` configure:

- `helmet()` pour les en-têtes de sécurité;
- CORS strict selon `FRONTEND_URL`;
- parsing JSON limité par `JSON_BODY_LIMIT`;
- rate limit global;
- routes REST sous `/api`.

L'organisation métier suit le pattern `routes -> controllers -> services -> repositories`. Les services contiennent la validation métier, les repositories isolent les requêtes SQL.

## 8. Documentation API

{md_table(["Méthode", "Endpoint", "Usage", "Accès"], api_rows)}

## 9. Base de données

La base cible est MySQL. Le nom local observé est `ticket_gestion`, mais il doit rester configurable via `DB_NAME`. Les tables principales sont:

{md_table(["Table", "Rôle"], db_rows)}

### Initialisation

1. Créer la base MySQL.
2. Importer le script SQL initial si nécessaire.
3. Vérifier les services de base: `IT`, `SD`, `MANAGER`, `ADMIN`, `PKI`.
4. Démarrer le backend: `src/database/initSchema.js` ajoute/ajuste plusieurs tables complémentaires.

## 10. Configuration des environnements

{md_table(["Variable", "Utilité", "Remarque"], env_rows)}

### Règles de sécurité des `.env`

- Les vrais secrets vont uniquement dans `.env`.
- Les fichiers `.env.example` contiennent seulement la structure et des valeurs factices.
- Les variables frontend commençant par `VITE_` sont publiques dans le navigateur.
- Ne jamais mettre `JWT_SECRET`, mot de passe DB, mot de passe Gmail ou mot de passe admin dans le frontend.

## 11. Prérequis de déploiement

- Node.js récent compatible React 19/Vite 8.
- npm.
- MySQL Server accessible depuis le backend.
- Accès réseau entre frontend et backend.
- Compte Gmail avec IMAP activé et mot de passe d'application si la synchronisation e-mail est utilisée.
- Variables `.env` créées sur la plateforme de test.
- Port backend disponible ou configuré autrement.

## 12. Procédure de déploiement plateforme de test

### Backend

```bash
cd backend_ticket_gestion
npm install
cp .env.example .env
# remplir les secrets réels
npm run dev
```

Contrôles:

```bash
curl http://localhost:2300/test
curl http://localhost:2300/api
```

### Frontend

```bash
cd frontend_ticket_gestion
npm install
cp .env.example .env
npm run dev
```

En local, ouvrir `http://localhost:5173/`.

### Lancement racine

```bash
npm install
npm run dev
```

Cette commande lance backend et frontend en parallèle avec `concurrently`.

## 13. Sécurité et durcissement

- Mots de passe employés hachés avec `bcrypt`.
- JWT signé avec `JWT_SECRET` et durée `JWT_EXPIRES_IN`.
- Rate limit global et login limité à 10 tentatives échouées par fenêtre de 15 minutes.
- CORS limité aux origines définies dans `FRONTEND_URL`.
- En-têtes HTTP renforcés via `helmet`.
- Comptes inactifs bloqués côté backend et socket.
- Accès aux tickets/rooms contrôlé par rôle et services autorisés.
- Pièces jointes e-mail filtrées par type MIME/extension.
- Bootstrap admin désactivable après création.

## 14. Tests et contrôles recommandés

{md_table(["Contrôle", "Commande ou attendu"], test_rows)}

## 15. Reprise par une nouvelle équipe

1. Lire ce document et les fichiers `.env.example`.
2. Installer Node.js, npm et MySQL.
3. Créer une base propre de test.
4. Importer/initialiser le schéma.
5. Configurer backend `.env` avec secrets réels.
6. Configurer frontend `.env` avec URLs publiques.
7. Créer/synchroniser l'admin initial.
8. Lancer `npm run dev`.
9. Tester login admin, création employé SD, création organisation/contact/ticket.
10. Vérifier dashboards, meetings, messages et inbox e-mail.

## 16. Points d'attention

- Un ancien processus Node peut occuper `2300` ou `5173`; libérer le port avant un nouveau lancement.
- Le frontend peut basculer sur `5174` si `5173` est déjà utilisé; dans ce cas, mettre à jour `FRONTEND_URL` côté backend ou libérer le port.
- `ENABLE_ADMIN_BOOTSTRAP` doit rester à `false` après création de l'admin initial.
- Les vrais secrets ne doivent jamais être ajoutés à Git.
- La logique actuelle de création des rooms démarre avec `SD`; les affectations IT/PKI sont ajoutées par action SD.
- Tester l'import Gmail uniquement avec une boîte dédiée Service Delivery.

## 17. Livrables utiles associés

- Code source complet du dépôt.
- Fichiers `.env.example` documentés.
- Script SQL initial `backend_ticket_gestion/issam.session.sql`.
- Documentation Markdown: `docs/LIVRABLE_REPRISE_ANALYSE_DEPLOIEMENT.md`.
- Documentation Word: `outputs/livrable-reprise-analyse-deploiement-agce-crm.docx`.
- Présentation déjà générée: `outputs/agce-crm-documentation-slides.pptx` si disponible.
"""


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for margin in ("top", "start", "bottom", "end"):
        node = tbl_cell_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), "120")
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_table(doc: Document, headers: list[str], rows: list[tuple[str, ...]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    set_cell_margins(table)
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
        set_cell_shading(hdr[i], "E8EEF5")
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    set_repeat_table_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.size = Pt(8.6)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    doc.add_paragraph()
    return table


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_code_block(doc: Document, code: str) -> None:
    for line in code.strip().splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(31, 77, 120)
    doc.add_paragraph()


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.text = "AGCE CRM - Livrable de reprise et déploiement"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(85, 85, 85)

    footer = section.footer.paragraphs[0]
    footer.text = "Document technique et fonctionnel - Plateforme de test"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(85, 85, 85)


def build_docx() -> None:
    doc = Document()
    style_document(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    r = subtitle.add_run(SUBTITLE)
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(70, 70, 70)

    add_table(doc, ["Information", "Valeur"], [
        ("Projet", "AGCE CRM"),
        ("Date", DATE),
        ("Objectif", "Reprise, analyse, configuration et déploiement sur plateforme de test"),
        ("Périmètre", "Frontend React/Vite, backend Express/MySQL, Socket.IO, Gmail IMAP, documentation environnement"),
    ], [1.6, 4.7])

    doc.add_heading("1. Résumé exécutif", level=1)
    doc.add_paragraph("AGCE CRM est une application web de gestion de tickets et de relation client construite autour d'un backend Express/MySQL et d'un frontend React/Vite. Elle couvre la réception des demandes, la gestion des contacts et organisations, la création et l'affectation des tickets, les échanges internes en temps réel, les réunions, les dashboards et la synchronisation d'e-mails clients.")
    doc.add_paragraph("Ce livrable donne à une équipe de reprise les informations nécessaires pour comprendre le périmètre, installer les prérequis, configurer les environnements, déployer sur une plateforme de test et valider les principaux flux applicatifs.")

    doc.add_heading("2. Périmètre du livrable", level=1)
    add_bullets(doc, [
        "Code source complet livré dans le dépôt local.",
        "Documentation technique et fonctionnelle par module.",
        "Prérequis de déploiement, configuration backend/frontend et base de données.",
        "Endpoints API, rôles, sécurité, tests et points d'attention.",
        "Checklist de reprise et de validation plateforme de test.",
    ])

    doc.add_heading("3. Code source complet de l'application", level=1)
    doc.add_paragraph("Le code source complet est disponible dans le dépôt full_stuck_ticket_gestion. Les secrets locaux, fichiers .env, node_modules, dist et outputs ne doivent pas être publiés. Le code maintenable est inventorié en annexe et dans la version Markdown.")
    add_table(doc, ["Dossier / fichier", "Contenu livré"], [
        ("backend_ticket_gestion/", "API Express, modules métier, accès MySQL, Socket.IO, Gmail Sync et sécurité."),
        ("frontend_ticket_gestion/", "Application React/Vite, routes, dashboards, composants et appels API."),
        ("backend_ticket_gestion/issam.session.sql", "Script SQL historique de création des tables principales."),
        ("backend_ticket_gestion/.env.example", "Structure backend sans secrets."),
        ("frontend_ticket_gestion/.env.example", "Structure frontend publique Vite."),
        ("package.json", "Scripts racine pour lancement simultané backend/frontend."),
    ], [2.1, 4.2])

    doc.add_heading("4. Architecture générale", level=1)
    add_bullets(doc, [
        "Utilisateur -> Frontend React/Vite.",
        "Frontend -> API Express via HTTP JSON et JWT.",
        "Frontend <-> Socket.IO pour la messagerie temps réel.",
        "Backend -> services métier -> repositories SQL -> MySQL.",
        "Boîte Gmail SD -> ImapFlow/mailparser -> tables client_emails et pièces jointes.",
    ])

    doc.add_heading("5. Technologies", level=1)
    add_table(doc, ["Couche", "Technologies et rôle"], TECH_STACK, [1.45, 4.85])

    doc.add_heading("6. Documentation fonctionnelle", level=1)
    add_table(doc, ["Module", "Description"], FUNCTIONAL_MODULES, [1.65, 4.65])

    doc.add_heading("7. Rôles et permissions", level=1)
    add_table(doc, ["Rôle", "Responsabilité", "Permissions principales"], [
        ("ADMIN", "Administration comptes", "CRUD employés, statut, mot de passe, accès admin."),
        ("SD", "Service Delivery", "Contacts, organisations, tickets, affectations, meetings, inbox e-mail, dashboards SD."),
        ("Manager", "Supervision", "Dashboards, analytics, lecture contrôlée, contacts/tickets/messages/meetings."),
        ("PKI", "Traitement PKI", "Tickets/rooms assignés, messages, propositions de résolution."),
        ("IT", "Traitement IT", "Tickets/rooms assignés, messages, propositions de résolution."),
    ], [1.0, 1.65, 3.65])

    doc.add_heading("8. Documentation technique backend", level=1)
    doc.add_paragraph("Le backend démarre depuis src/index.js, charge les variables d'environnement, initialise le schéma complémentaire, démarre Socket.IO puis écoute sur PORT. L'application Express applique Helmet, CORS strict, limite JSON, rate limiting global et routes REST sous /api.")
    doc.add_paragraph("La structure suit le pattern routes -> controllers -> services -> repositories. Les services portent la logique métier et les repositories isolent les requêtes SQL.")

    doc.add_heading("9. Documentation API", level=1)
    add_table(doc, ["Méthode", "Endpoint", "Usage", "Accès"], ROUTE_ROWS, [0.85, 1.7, 2.2, 1.55])

    doc.add_heading("10. Base de données", level=1)
    doc.add_paragraph("La base cible est MySQL. Le nom observé en local est ticket_gestion, mais la valeur doit rester configurable avec DB_NAME. Le backend ajoute/ajuste plusieurs tables au démarrage via src/database/initSchema.js.")
    add_table(doc, ["Table", "Rôle"], DB_ROWS, [1.75, 4.55])

    doc.add_heading("11. Configuration des environnements", level=1)
    add_table(doc, ["Variable", "Utilité", "Remarque"], ENV_ROWS, [1.75, 2.55, 2.0])
    doc.add_paragraph("Règle importante: les vrais secrets restent uniquement dans .env. Les fichiers .env.example documentent la structure avec des valeurs factices. Les variables frontend VITE_* sont publiques dans le navigateur.")

    doc.add_heading("12. Prérequis de déploiement", level=1)
    add_bullets(doc, [
        "Node.js récent, npm et accès à Internet ou registre npm interne pour installer les dépendances.",
        "MySQL Server opérationnel et accessible depuis le backend.",
        "Compte Gmail avec IMAP et mot de passe d'application si la synchronisation e-mail est activée.",
        "Ports disponibles: 2300 pour le backend, 5173 pour le frontend local.",
        "Variables d'environnement créées côté backend et frontend.",
    ])

    doc.add_heading("13. Procédure de déploiement test", level=1)
    doc.add_heading("Backend", level=2)
    add_code_block(doc, """
cd backend_ticket_gestion
npm install
cp .env.example .env
# remplir les valeurs reelles
npm run dev
""")
    doc.add_heading("Frontend", level=2)
    add_code_block(doc, """
cd frontend_ticket_gestion
npm install
cp .env.example .env
npm run dev
""")
    doc.add_heading("Lancement racine", level=2)
    add_code_block(doc, """
npm install
npm run dev
""")

    doc.add_heading("14. Sécurité", level=1)
    add_bullets(doc, [
        "Mots de passe hachés avec bcrypt.",
        "JWT signé avec JWT_SECRET et durée JWT_EXPIRES_IN.",
        "Login limité à 10 tentatives échouées par fenêtre de 15 minutes.",
        "CORS strict selon FRONTEND_URL.",
        "Helmet pour les en-têtes de sécurité.",
        "Comptes inactifs bloqués côté API et Socket.IO.",
        "Contrôle d'accès backend par rôle et accès aux rooms/tickets.",
        "Pièces jointes e-mail filtrées par MIME/extension.",
    ])

    doc.add_heading("15. Tests et contrôles recommandés", level=1)
    add_table(doc, ["Contrôle", "Commande ou résultat attendu"], TEST_CHECKS, [2.0, 4.3])

    doc.add_heading("16. Checklist de reprise", level=1)
    add_numbered(doc, [
        "Installer Node.js, npm et MySQL.",
        "Créer la base de test et importer/initialiser le schéma.",
        "Créer backend .env avec secrets réels.",
        "Créer frontend .env avec VITE_API_URL et VITE_SOCKET_URL.",
        "Créer ou synchroniser l'admin initial.",
        "Lancer npm run dev depuis la racine.",
        "Tester login admin puis création d'un employé SD.",
        "Tester organisation, contact, ticket, affectation, message, meeting et dashboard.",
        "Vérifier que les secrets ne sont pas committés.",
        "Préparer le build frontend pour la plateforme de test.",
    ])

    doc.add_heading("17. Points d'attention", level=1)
    add_bullets(doc, [
        "Si le port 2300 ou 5173 est occupé, arrêter l'ancien processus avant de relancer.",
        "Si Vite démarre sur 5174, libérer 5173 ou adapter FRONTEND_URL.",
        "ENABLE_ADMIN_BOOTSTRAP doit rester false après création de l'admin.",
        "Ne jamais livrer de vrais secrets dans .env.example ou dans le frontend.",
        "Tester Gmail avec une boîte dédiée Service Delivery.",
    ])

    doc.add_heading("18. Annexe - Inventaire du code source", level=1)
    doc.add_paragraph("Inventaire synthétique des fichiers maintenables, hors node_modules, build, outputs, secrets et fichiers binaires.")
    files = source_inventory()
    for item in files[:160]:
        doc.add_paragraph(item, style="List Bullet")
    if len(files) > 160:
        doc.add_paragraph(f"Inventaire complet dans {MD_PATH.relative_to(ROOT)} ({len(files)} fichiers listés).")

    doc.save(DOCX_PATH)


def main() -> None:
    MD_PATH.write_text(build_markdown(), encoding="utf-8")
    build_docx()
    print(f"Created {MD_PATH}")
    print(f"Created {DOCX_PATH}")


if __name__ == "__main__":
    main()
