from copy import deepcopy
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from lxml import etree


ROOT = Path("/Users/issamouladsmane/Desktop/full_stuck_ticket_gestion")
SOURCE = Path("/Users/issamouladsmane/Downloads/Memoire_AGCE_CRM_Final_Professionnel-2 (1).docx")
ASSETS = ROOT / "livrables/revision_prof/preuves_existantes"
OUTPUT = ROOT / "livrables/Memoire_AGCE_CRM_Corrige_Remarques_Enseignante.docx"
INTERMEDIATE = ROOT / "livrables/revision_prof/memoire_revision_intermediaire.docx"


def paragraph_starting(doc, prefix):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise ValueError(f"Paragraphe introuvable : {prefix}")


def body_heading_starting(doc, prefix):
    for paragraph in doc.paragraphs:
        if paragraph.style.name in {"Titre", "Titre 2"} and paragraph.text.strip().startswith(prefix):
            return paragraph
    raise ValueError(f"Titre introuvable : {prefix}")


def set_text(paragraph, text):
    paragraph.clear()
    paragraph.add_run(text)


def replace_in_paragraph(doc, prefix, old, new):
    paragraph = paragraph_starting(doc, prefix)
    set_text(paragraph, paragraph.text.replace(old, new))


def append_citation(doc, prefix, citation):
    paragraph = paragraph_starting(doc, prefix)
    if citation not in paragraph.text:
        set_text(paragraph, f"{paragraph.text.rstrip()} {citation}")


def remove_element(element):
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def remove_subsection(doc, heading_text):
    paragraphs = doc.paragraphs
    start = next(i for i, p in enumerate(paragraphs) if p.text.strip() == heading_text)
    remove = [paragraphs[start]._element]
    for paragraph in paragraphs[start + 1 :]:
        if paragraph.style.name in {"Titre", "Titre 2"}:
            break
        remove.append(paragraph._element)
    for element in remove:
        remove_element(element)


def move_arabic_summary_before_english(doc):
    paragraphs = doc.paragraphs
    english = next(i for i, p in enumerate(paragraphs) if p.text.strip() == "Abstract")
    arabic = next(i for i, p in enumerate(paragraphs) if p.text.strip() == "الملخص")
    arabic_block = [p._element for p in paragraphs[arabic : arabic + 4]]
    anchor = paragraphs[english]._element
    parent = anchor.getparent()
    position = parent.index(anchor)
    for element in arabic_block:
        parent.remove(element)
        parent.insert(position, element)
        position += 1

    toc_arabic = next(
        p for p in doc.paragraphs
        if p.style.name.startswith("TOC") and "الملخص" in p.text
    )
    toc_english = next(
        p for p in doc.paragraphs
        if p.style.name.startswith("TOC") and "Abstract" in p.text
    )
    parent = toc_english._element.getparent()
    parent.remove(toc_arabic._element)
    parent.insert(parent.index(toc_english._element), toc_arabic._element)
    set_text(toc_arabic, "الملخص\t4")
    set_text(toc_english, "Abstract\t5")


def rename_structure(doc):
    replacements = {
        "Chapitre 1 : Contexte général et revue de la littérature": "Section 1 : Contexte général et revue de la littérature",
        "Chapitre 2 : Analyse des besoins et conception": "Section 2 : Analyse des besoins et conception",
        "Chapitre 3 : Réalisation, implémentation et validation": "Section 3 : Réalisation, présentation et validation",
        "1.2 Présentation de l’organisme d’accueil : l’AGCE": "1.1 Présentation de l’organisme d’accueil : l’AGCE",
        "1.3 Cadre réglementaire et contraintes de souveraineté": "1.2 Cadre réglementaire et contraintes de souveraineté",
        "1.4 Problématique métier": "1.3 Problématique métier",
        "1.5 Étude de l’existant": "1.4 Étude de l’existant",
        "1.6 Solution proposée": "1.5 Solution proposée",
        "2.2 Méthodologie de développement": "2.1 Méthodologie de développement",
        "2.3 Identification des acteurs": "2.2 Identification des acteurs",
        "2.4 Besoins fonctionnels": "2.3 Besoins fonctionnels",
        "2.5 Besoins non fonctionnels": "2.4 Besoins non fonctionnels",
        "2.6 Matrice d’habilitation RBAC": "2.5 Matrice d’habilitation RBAC",
        "2.7 Diagramme de cas d’utilisation": "2.6 Diagramme de cas d’utilisation",
        "2.8 Description textuelle de cas d’utilisation": "2.7 Description textuelle de cas d’utilisation",
        "2.9 Architecture générale": "2.8 Architecture générale",
        "2.10 Modèle conceptuel des classes": "2.9 Modèle conceptuel des classes",
        "2.11 Modèle relationnel de la base de données": "2.10 Modèle relationnel de la base de données",
        "2.12 Conception de la sécurité": "2.11 Conception de la sécurité",
        "3.2 Environnement de développement": "3.1 Environnement de développement",
        "Conclusion générale et perspectives": "Conclusion générale : contributions et perspectives",
    }
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        for old, new in replacements.items():
            if text == old:
                set_text(paragraph, new)
                break
            if text.startswith(old + " "):
                set_text(paragraph, paragraph.text.replace(old, new, 1))
                break
    replace_in_paragraph(
        doc,
        "Le présent mémoire est organisé en trois",
        "trois chapitres",
        "trois sections",
    )
    replace_in_paragraph(
        doc,
        "Le présent mémoire est organisé en trois",
        "Le premier chapitre",
        "La première section",
    )
    replace_in_paragraph(
        doc,
        "Le présent mémoire est organisé en trois",
        "Le deuxième chapitre",
        "La deuxième section",
    )
    replace_in_paragraph(
        doc,
        "Le présent mémoire est organisé en trois",
        "Le troisième chapitre",
        "La troisième section",
    )


def revise_lists(doc):
    figure_lines = [
        "Figure 2.1 - Diagramme de cas d’utilisation du système AGCE CRM",
        "Figure 2.2 - Architecture générale du système AGCE CRM",
        "Figure 2.3 - Modèle conceptuel simplifié des classes",
        "Figure 3.1 - Séquence de création et d’affectation d’un ticket",
        "Figure 3.2 - Séquence de messagerie temps réel",
        "Figure 3.3 - Écran sécurisé d’authentification",
        "Figure 3.4 - Tableau de bord et accès rapides du Service Delivery",
        "Figure 3.5 - Création d’un ticket par le Service Delivery",
        "Figure 3.6 - Interface de gestion des tickets",
        "Figure 3.7 - Messagerie associée à un ticket",
        "Figure 3.8 - Planification et suivi des réunions",
        "Figure 3.9 - Profil utilisateur et gestion du mot de passe",
    ]
    start = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "Liste des figures")
    end = next(i for i, p in enumerate(doc.paragraphs[start + 1 :], start + 1) if p.style.name == "Titre")
    existing = doc.paragraphs[start + 1 : end]
    for paragraph, text in zip(existing, figure_lines):
        set_text(paragraph, text)
    anchor = doc.paragraphs[end]
    for text in figure_lines[len(existing) :]:
        anchor.insert_paragraph_before(text, style="Corps")

    table_lines = [
        "Tableau 1.1 - Comparaison critique des solutions existantes",
        "Tableau 2.1 - Acteurs du système",
        "Tableau 2.2 - Besoins fonctionnels du système",
        "Tableau 2.3 - Besoins non fonctionnels",
        "Tableau 2.4 - Matrice d’habilitation RBAC",
        "Tableau 2.5 - Description synthétique de cas d’utilisation",
        "Tableau 2.6 - Synthèse du modèle relationnel",
        "Tableau 3.1 - Environnement logiciel du projet",
        "Tableau 3.2 - Jeux de tests fonctionnels et sécurité",
        "Tableau 3.3 - Synthèse des résultats par objectif",
    ]
    start = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "Liste des tableaux")
    end = next(i for i, p in enumerate(doc.paragraphs[start + 1 :], start + 1) if p.style.name == "Titre")
    existing = doc.paragraphs[start + 1 : end]
    for paragraph, text in zip(existing, table_lines):
        set_text(paragraph, text)
    anchor = doc.paragraphs[end]
    for text in table_lines[len(existing) :]:
        anchor.insert_paragraph_before(text, style="Corps")


def update_static_toc_labels(doc):
    substitutions = {
        "Chapitre 1 :": "Section 1 :",
        "Chapitre 2 :": "Section 2 :",
        "Chapitre 3 : Réalisation, implémentation et validation": "Section 3 : Réalisation, présentation et validation",
        "Conclusion générale et perspectives": "Conclusion générale : contributions et perspectives",
        "1.2 Présentation": "1.1 Présentation",
        "1.3 Cadre": "1.2 Cadre",
        "1.4 Problématique": "1.3 Problématique",
        "1.5 Étude": "1.4 Étude",
        "1.6 Solution": "1.5 Solution",
        "2.2 Méthodologie": "2.1 Méthodologie",
        "2.3 Identification": "2.2 Identification",
        "2.4 Besoins fonctionnels": "2.3 Besoins fonctionnels",
        "2.5 Besoins non fonctionnels": "2.4 Besoins non fonctionnels",
        "2.6 Matrice": "2.5 Matrice",
        "2.7 Diagramme": "2.6 Diagramme",
        "2.8 Description": "2.7 Description",
        "2.9 Architecture": "2.8 Architecture",
        "2.10 Modèle conceptuel": "2.9 Modèle conceptuel",
        "2.11 Modèle relationnel": "2.10 Modèle relationnel",
        "2.12 Conception": "2.11 Conception",
        "3.2 Environnement": "3.1 Environnement",
    }
    remove_prefixes = (
        "1.1 Introduction",
        "1.7 Synthèse",
        "2.1 Introduction",
        "2.13 Synthèse",
        "3.1 Introduction",
        "3.15 Synthèse",
    )
    for paragraph in list(doc.paragraphs):
        if not paragraph.style.name.startswith("TOC"):
            continue
        text = paragraph.text.strip()
        if any(text.startswith(prefix) for prefix in remove_prefixes):
            remove_element(paragraph._element)
            continue
        for old, new in substitutions.items():
            if text.startswith(old):
                set_text(paragraph, paragraph.text.replace(old, new, 1))
                break
    frontend_entry = next(
        p for p in doc.paragraphs
        if p.style.name.startswith("TOC") and p.text.strip().startswith("3.3 Organisation du frontend")
    )
    frontend_entry.insert_paragraph_before("3.2 Présentation de l’application réalisée\t27", style="TOC 2")
    set_text(frontend_entry, frontend_entry.text.replace("3.3 Organisation du frontend", "3.3 Organisation du frontend"))


def revise_tools_table(doc):
    table = doc.tables[8]
    compact_rows = [
        ("Interface", "React + Vite{{FN1}}{{FN2}}", "Composants et assemblage frontend"),
        ("Serveur API", "Node.js + Express.js{{FN3}}{{FN4}}", "Routes REST et contrôles d’accès"),
        ("Données", "MySQL{{FN5}}", "Stockage relationnel"),
        ("Temps réel", "Socket.IO{{FN6}}", "Messagerie par ticket"),
        ("Sécurité", "JWT + bcrypt", "Session et hachage des mots de passe"),
        ("Courriels", "ImapFlow + mailparser{{FN7}}{{FN8}}", "Import des demandes reçues"),
        ("Interface", "Tailwind CSS + Recharts{{FN9}}{{FN10}}", "Mise en forme et indicateurs"),
        ("Validation", "Scénarios métier et navigateur", "Contrôles fonctionnels"),
    ]
    for row, values in zip(table.rows[1:], compact_rows):
        for cell, value in zip(row.cells, values):
            cell.text = value
    caption = paragraph_starting(doc, "Tableau 3.1 -")
    next_heading = body_heading_starting(doc, "3.3 Organisation du frontend")
    note = next_heading.insert_paragraph_before(
        "Les technologies sont retenues pour leur adéquation avec l’architecture définie en conception : interface découplée, API contrôlée, persistance locale et échanges temps réel.",
        style="Corps",
    )
    note.paragraph_format.space_before = Pt(4)
    note.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.space_after = Pt(4)


def add_application_presentation(doc):
    anchor = body_heading_starting(doc, "3.3 Organisation du frontend")
    anchor.insert_paragraph_before("3.2 Présentation de l’application réalisée", style="Titre 2")
    anchor.insert_paragraph_before(
        "L’application AGCE CRM matérialise la conception présentée dans la section précédente. Après authentification, chaque utilisateur accède uniquement aux vues autorisées par son rôle : administration des employés, pilotage Service Delivery, traitement IT/PKI ou supervision Manager. Les écrans suivants illustrent le parcours principal, de l’accès sécurisé au suivi d’un ticket.",
        style="Corps",
    )

    figures = [
        ("image4.jpeg", "Figure 3.3 - Écran sécurisé d’authentification"),
        ("image6.jpeg", "Figure 3.4 - Tableau de bord et accès rapides du Service Delivery"),
        ("image7.jpeg", "Figure 3.5 - Création d’un ticket par le Service Delivery"),
        ("image9.jpeg", "Figure 3.6 - Interface de gestion et de filtrage des tickets"),
        ("image10.jpeg", "Figure 3.7 - Messagerie temps réel associée à un ticket"),
        ("image11.jpeg", "Figure 3.8 - Planification et suivi des réunions"),
        ("image_profile.png", "Figure 3.9 - Profil utilisateur et gestion du mot de passe"),
    ]
    for image, caption in figures:
        image_paragraph = anchor.insert_paragraph_before()
        image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_paragraph.paragraph_format.space_before = Pt(8)
        image_paragraph.paragraph_format.space_after = Pt(3)
        picture = image_paragraph.add_run().add_picture(str(ASSETS / image), width=Inches(5.9))
        picture._inline.docPr.set("descr", caption)
        picture._inline.docPr.set("title", caption)
        caption_paragraph = anchor.insert_paragraph_before(caption, style="Corps")
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_paragraph.paragraph_format.space_after = Pt(8)
        for run in caption_paragraph.runs:
            run.italic = True

    anchor.insert_paragraph_before(
        "Ces interfaces traduisent directement les choix de conception : le tableau de bord synthétise l’état des tickets et propose des accès rapides aux opérations courantes ; le formulaire matérialise la réception des demandes ; la liste assure la traçabilité du workflow ; la room de discussion limite les échanges aux personnes habilitées ; le calendrier structure les réunions de suivi. L’espace Profile présente l’identité, le rôle, l’état du compte et l’accès à la modification du mot de passe.",
        style="Corps",
    )


def mark_table_headers(doc):
    for table in doc.tables:
        row_props = table.rows[0]._tr.get_or_add_trPr()
        if row_props.find(qn("w:tblHeader")) is None:
            etree.SubElement(row_props, qn("w:tblHeader")).set(qn("w:val"), "true")


def add_citations(doc):
    additions = [
        ("L’Autorité Gouvernementale de Certification Électronique est", "[4]"),
        ("La certification électronique et la signature électronique", "[1]–[3]"),
        ("Plusieurs solutions de ticketing existent", "[20]–[23]"),
        ("Le projet suit une démarche inspirée", "[9]"),
        ("La modélisation utilise UML", "[8]"),
        ("La matrice d’habilitation formalise", "[5]"),
        ("La sécurité du système est conçue", "[6], [7], [18], [19]"),
        ("La réalisation s’appuie sur une architecture web moderne", "[10]–[13], [24]"),
        ("Socket.IO est utilisé pour la communication", "[12]"),
        ("Le projet intègre un service d’import d’e-mails", "[25], [26]"),
    ]
    for prefix, references in additions:
        try:
            append_citation(doc, prefix, references)
        except ValueError:
            # The removed intermediary introduction contained one technology sentence.
            pass


def restructure_conclusion(doc):
    title = paragraph_starting(doc, "Conclusion générale : contributions et perspectives")
    first_content = paragraph_starting(doc, "Ce mémoire a présenté la conception")
    first_content.insert_paragraph_before("Contributions", style="Titre 2")
    perspective_anchor = paragraph_starting(doc, "Même si le système répond aux objectifs principaux du projet")
    perspective_anchor.insert_paragraph_before("Perspectives", style="Titre 2")
    set_text(
        perspective_anchor,
        "Le prolongement prioritaire consiste à évaluer l’application dans un environnement réel de l’AGCE afin de mesurer les temps de traitement, la disponibilité et l’adoption par les utilisateurs. Cette phase permettra également de formaliser les procédures de déploiement, de sauvegarde et de journalisation.",
    )
    perspective_next = paragraph_starting(doc, "Les perspectives d’amélioration concernent également")
    set_text(
        perspective_next,
        "Sur le plan fonctionnel, les perspectives portent sur la gestion des SLA, la génération de rapports, l’export des indicateurs, un historique d’audit plus détaillé et une politique de renouvellement initial des mots de passe. Ces évolutions prolongent la conception actuelle sans remettre en cause la séparation des rôles.",
    )
    title.paragraph_format.space_after = Pt(10)


def add_footnotes_and_reference_section(input_path, output_path):
    word_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    rel_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
    content_ns = "http://schemas.openxmlformats.org/package/2006/content-types"
    ns = {"w": word_ns}
    footnotes = {
        1: "React, site officiel : https://react.dev/ (consulté le 25 mai 2026).",
        2: "Vite, site officiel : https://vite.dev/ (consulté le 25 mai 2026).",
        3: "Node.js, site officiel : https://nodejs.org/ (consulté le 25 mai 2026).",
        4: "Express.js, site officiel : https://expressjs.com/ (consulté le 25 mai 2026).",
        5: "MySQL, documentation officielle : https://dev.mysql.com/doc/ (consulté le 25 mai 2026).",
        6: "Socket.IO, site officiel : https://socket.io/ (consulté le 25 mai 2026).",
        7: "ImapFlow, documentation officielle : https://imapflow.com/ (consulté le 25 mai 2026).",
        8: "Mailparser, documentation officielle : https://nodemailer.com/extras/mailparser/ (consulté le 25 mai 2026).",
        9: "Tailwind CSS, site officiel : https://tailwindcss.com/ (consulté le 25 mai 2026).",
        10: "Recharts, site officiel : https://recharts.org/ (consulté le 25 mai 2026).",
    }

    with ZipFile(input_path, "r") as original:
        files = {name: original.read(name) for name in original.namelist()}

    document = etree.fromstring(files["word/document.xml"])
    for text_node in document.xpath(".//w:t", namespaces=ns):
        if "{{FN" not in (text_node.text or ""):
            continue
        text = text_node.text
        ids = []
        for number in footnotes:
            marker = f"{{{{FN{number}}}}}"
            if marker in text:
                text = text.replace(marker, "")
                ids.append(number)
        text_node.text = text
        run = text_node.getparent()
        parent = run.getparent()
        index = parent.index(run) + 1
        for number in ids:
            ref_run = etree.Element(qn("w:r"))
            ref_run_props = etree.SubElement(ref_run, qn("w:rPr"))
            etree.SubElement(ref_run_props, qn("w:rStyle")).set(qn("w:val"), "FootnoteReference")
            etree.SubElement(ref_run, qn("w:footnoteReference")).set(qn("w:id"), str(number))
            parent.insert(index, ref_run)
            index += 1

    body = document.find(qn("w:body"))
    final_sect = body.find(qn("w:sectPr"))
    biblio_heading = next(
        p for p in body.findall(qn("w:p"))
        if "".join(p.xpath(".//w:t/text()", namespaces=ns)).strip() == "Bibliographie"
    )
    numbered_section = deepcopy(final_sect)
    break_p = etree.Element(qn("w:p"))
    break_ppr = etree.SubElement(break_p, qn("w:pPr"))
    break_ppr.append(numbered_section)
    body.insert(body.index(biblio_heading), break_p)
    for footer_reference in final_sect.findall(qn("w:footerReference")):
        final_sect.remove(footer_reference)

    footnotes_root = etree.Element(qn("w:footnotes"), nsmap={"w": word_ns})
    for number, footnote_type in [(-1, "separator"), (0, "continuationSeparator")]:
        footnote = etree.SubElement(footnotes_root, qn("w:footnote"))
        footnote.set(qn("w:type"), footnote_type)
        footnote.set(qn("w:id"), str(number))
        para = etree.SubElement(footnote, qn("w:p"))
        run = etree.SubElement(para, qn("w:r"))
        etree.SubElement(run, qn(f"w:{footnote_type}"))
    for number, text in footnotes.items():
        footnote = etree.SubElement(footnotes_root, qn("w:footnote"))
        footnote.set(qn("w:id"), str(number))
        para = etree.SubElement(footnote, qn("w:p"))
        para_props = etree.SubElement(para, qn("w:pPr"))
        etree.SubElement(para_props, qn("w:pStyle")).set(qn("w:val"), "FootnoteText")
        ref_run = etree.SubElement(para, qn("w:r"))
        ref_props = etree.SubElement(ref_run, qn("w:rPr"))
        etree.SubElement(ref_props, qn("w:rStyle")).set(qn("w:val"), "FootnoteReference")
        etree.SubElement(ref_run, qn("w:footnoteRef"))
        space_run = etree.SubElement(para, qn("w:r"))
        space_text = etree.SubElement(space_run, qn("w:t"))
        space_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        space_text.text = " "
        text_run = etree.SubElement(para, qn("w:r"))
        etree.SubElement(text_run, qn("w:t")).text = text

    relationships = etree.fromstring(files["word/_rels/document.xml.rels"])
    relationship = etree.SubElement(relationships, f"{{{rel_ns}}}Relationship")
    relationship.set("Id", "rIdFootnotes")
    relationship.set(
        "Type",
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes",
    )
    relationship.set("Target", "footnotes.xml")

    content_types = etree.fromstring(files["[Content_Types].xml"])
    override = etree.SubElement(content_types, f"{{{content_ns}}}Override")
    override.set("PartName", "/word/footnotes.xml")
    override.set(
        "ContentType",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml",
    )
    settings = etree.fromstring(files["word/settings.xml"])
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = etree.SubElement(settings, qn("w:updateFields"))
    update_fields.set(qn("w:val"), "true")

    files["word/document.xml"] = etree.tostring(document, xml_declaration=True, encoding="UTF-8", standalone=True)
    files["word/footnotes.xml"] = etree.tostring(footnotes_root, xml_declaration=True, encoding="UTF-8", standalone=True)
    files["word/_rels/document.xml.rels"] = etree.tostring(relationships, xml_declaration=True, encoding="UTF-8", standalone=True)
    files["[Content_Types].xml"] = etree.tostring(content_types, xml_declaration=True, encoding="UTF-8", standalone=True)
    files["word/settings.xml"] = etree.tostring(settings, xml_declaration=True, encoding="UTF-8", standalone=True)
    with ZipFile(output_path, "w", compression=ZIP_DEFLATED) as final:
        for name, data in files.items():
            final.writestr(name, data)


def main():
    doc = Document(SOURCE)
    move_arabic_summary_before_english(doc)
    for subsection in [
        "1.1 Introduction",
        "1.7 Synthèse du chapitre",
        "2.1 Introduction",
        "2.13 Synthèse du chapitre",
        "3.1 Introduction",
        "3.15 Synthèse du chapitre",
    ]:
        remove_subsection(doc, subsection)
    rename_structure(doc)
    update_static_toc_labels(doc)
    revise_lists(doc)
    revise_tools_table(doc)
    add_application_presentation(doc)
    mark_table_headers(doc)
    add_citations(doc)
    restructure_conclusion(doc)
    doc.save(INTERMEDIATE)
    add_footnotes_and_reference_section(INTERMEDIATE, OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
