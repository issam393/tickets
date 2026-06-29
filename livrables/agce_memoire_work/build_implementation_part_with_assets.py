from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/issamouladsmane/Desktop/full_stuck_ticket_gestion")
OUT_DIR = ROOT / "livrables"
WORK = OUT_DIR / "agce_memoire_work"
ASSETS = WORK / "user_assets"
DOCX_OUT = OUT_DIR / "Partie_Implementation_AGCE_CRM_Avec_Arborescence_Et_Tests.docx"
MD_OUT = OUT_DIR / "Partie_Implementation_AGCE_CRM_Avec_Arborescence_Et_Tests.md"

BACKEND_TREE = ASSETS / "backend_tree_1.png"
TICKETS_TABLE = ASSETS / "tickets_table_capture.png"

BLUE = RGBColor(31, 78, 121)
DARK = RGBColor(22, 47, 75)
MUTED = RGBColor(95, 105, 115)
HEADER_FILL = "E8EEF5"
TABLE_FILL = "F2F4F7"
NOTE_FILL = "FFF7E8"

md: list[str] = []


def set_run(run, size=11, bold=None, italic=None, color=None, name="Calibri"):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), name)
    rpr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.find(qn("w:tcMar"))
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for name, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def style_table(table, widths=None, header=True):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    if widths:
        grid = table._tbl.tblGrid
        if grid is None:
            grid = OxmlElement("w:tblGrid")
            table._tbl.insert(0, grid)
        for child in list(grid):
            grid.remove(child)
        for width in widths:
            col = OxmlElement("w:gridCol")
            col.set(qn("w:w"), str(width))
            grid.append(col)
        for row in table.rows:
            for idx, width in enumerate(widths):
                if idx < len(row.cells):
                    set_cell_width(row.cells[idx], width)
    for idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell_margins(cell)
            for par in cell.paragraphs:
                par.paragraph_format.space_before = Pt(0)
                par.paragraph_format.space_after = Pt(2)
                for run in par.runs:
                    set_run(run, size=9.3)
        if header and idx == 0:
            for cell in row.cells:
                shade(cell, TABLE_FILL)
                for par in cell.paragraphs:
                    for run in par.runs:
                        set_run(run, size=9.3, bold=True, color=DARK)


def configure(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK, 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)


def h(doc, text, level=1):
    doc.add_heading(text, level=level)
    md.append(f"\n{'#' * level} {text}\n")


def p(doc, text):
    par = doc.add_paragraph()
    r = par.add_run(text)
    set_run(r)
    md.append(text + "\n")


def bullets(doc, items):
    for item in items:
        par = doc.add_paragraph(style="List Bullet")
        r = par.add_run(item)
        set_run(r, size=10.8)
        md.append(f"- {item}\n")


def caption(doc, text):
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = par.add_run(text)
    set_run(r, size=9.4, italic=True, color=MUTED)
    md.append(f"_{text}_\n")


def table(doc, headers, rows, widths=None, title=None):
    if title:
        caption(doc, title)
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, head in enumerate(headers):
        t.rows[0].cells[i].text = str(head)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    style_table(t, widths=widths)
    md.append("\n| " + " | ".join(headers) + " |\n")
    md.append("|" + "|".join(["---"] * len(headers)) + "|\n")
    for row in rows:
        md.append("| " + " | ".join(str(v).replace("\n", "<br>") for v in row) + " |\n")
    md.append("\n")
    return t


def note(doc, title, text):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.cell(0, 0)
    shade(cell, NOTE_FILL)
    cell_margins(cell, top=120, bottom=120, start=160, end=160)
    par = cell.paragraphs[0]
    r = par.add_run(title + " : ")
    set_run(r, bold=True, color=DARK)
    r2 = par.add_run(text)
    set_run(r2, size=10.5)
    md.append(f"\n**{title} :** {text}\n")


def picture(doc, path: Path, width, caption_text):
    if path.exists():
        par = doc.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = par.add_run()
        run.add_picture(str(path), width=Inches(width))
        caption(doc, caption_text)
        md.append(f"\n![{caption_text}]({path})\n\n")
    else:
        note(doc, "Image manquante", f"Le fichier {path} est introuvable.")


def title_page(doc):
    for text in [
        "Université Saad Dahleb - Blida 1",
        "Faculté des Sciences",
        "Département d'Informatique",
        "Option : Systèmes d'Information et Ingénierie Logicielle, ISIL",
    ]:
        par = doc.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = par.add_run(text)
        set_run(r, bold="Université" in text, size=12)
    doc.add_paragraph()
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = par.add_run("Partie du mémoire")
    set_run(r, size=15, bold=True, color=BLUE)
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = par.add_run("Réalisation, implémentation et tests")
    set_run(r, size=18, bold=True, color=DARK)
    doc.add_paragraph()
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.cell(0, 0)
    shade(cell, HEADER_FILL)
    cell_margins(cell, top=180, bottom=180, start=180, end=180)
    par = cell.paragraphs[0]
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = par.add_run("AGCE CRM - Système de gestion des incidents, des tickets et de la relation client")
    set_run(r, size=14, bold=True, color=DARK)
    doc.add_paragraph()
    table(doc, ["Élément", "Information"], [
        ("Projet", "Conception et réalisation d'une application de gestion des incidents et des tickets des services de l'AGCE"),
        ("Étudiants", "Mr. Laceb Karim\nMr. Ouladsmane Issam"),
        ("Encadrants", "Mrs. Arkam Meriem\nMr. Lekhchine Sami"),
        ("Année universitaire", "2025/2026"),
    ], widths=[2500, 6860])
    doc.add_page_break()


def implementation(doc):
    h(doc, "Chapitre 4 : Réalisation, implémentation et tests", 1)
    h(doc, "4.1 Introduction du chapitre", 2)
    p(doc, "Ce chapitre présente la mise en œuvre technique de l'application AGCE CRM. Il décrit l'environnement de réalisation, l'organisation du code, l'implémentation du frontend et du backend, les mécanismes de sécurité, la messagerie temps réel, l'import des e-mails clients, les interfaces principales et les tests effectués. L'objectif est de montrer comment la conception proposée a été traduite en une solution web fonctionnelle, sécurisée et adaptée au contexte de l'AGCE.")
    p(doc, "La réalisation ne se limite pas à l'affichage de formulaires. Elle met en place un cycle complet de traitement des demandes, depuis la réception de l'information jusqu'à la résolution du ticket, avec contrôle des rôles, traçabilité et conservation locale des données.")

    h(doc, "4.2 Environnement matériel", 2)
    table(doc, ["Élément", "Description"], [
        ("Ordinateur portable", "Machine utilisée pour le développement local du frontend, du backend et de la base de données."),
        ("Processeur", "Processeur compatible avec l'exécution de Node.js, MySQL et d'un navigateur web moderne."),
        ("Mémoire", "Mémoire suffisante pour exécuter simultanément l'API, le serveur de développement React et MySQL."),
        ("Système d'exploitation", "Environnement macOS ou Windows selon la machine de développement utilisée."),
    ], widths=[2400, 6960])

    h(doc, "4.3 Environnement logiciel", 2)
    table(doc, ["Outil / technologie", "Utilisation dans le projet"], [
        ("Visual Studio Code", "Édition du code source et navigation entre les modules."),
        ("Node.js et npm", "Exécution du backend, installation des dépendances et lancement des scripts."),
        ("React 19 et Vite", "Développement de l'interface utilisateur et serveur frontend de développement."),
        ("Express.js", "Définition de l'API REST, des routes et des middlewares."),
        ("MySQL Server", "Persistance relationnelle des données."),
        ("MySQL Workbench", "Création, consultation et vérification des tables."),
        ("Postman", "Test manuel des endpoints API et des réponses HTTP."),
        ("Git / GitHub", "Versionnement du code et suivi des modifications."),
        ("Socket.IO", "Échanges en temps réel dans les rooms liées aux tickets."),
        ("ImapFlow et mailparser", "Récupération et analyse des e-mails clients reçus par le Service Delivery."),
    ], widths=[2600, 6760])

    h(doc, "4.4 Installation et configuration", 2)
    p(doc, "L'installation du projet se fait en deux parties : le frontend et le backend. Après l'installation de Node.js, npm et MySQL Server, une base de données locale nommée agce_crm est préparée. Le backend initialise les tables complémentaires nécessaires au fonctionnement de l'application.")
    table(doc, ["Étape", "Commande ou action", "Objectif"], [
        ("Création de la base", "CREATE DATABASE agce_crm;", "Préparer l'espace de stockage MySQL."),
        ("Installation frontend", "npm install dans frontend_ticket_gestion", "Installer React, Vite et les dépendances UI."),
        ("Lancement frontend", "npm run dev", "Démarrer l'interface en développement."),
        ("Installation backend", "npm install dans backend_ticket_gestion", "Installer Express, MySQL2, JWT, bcrypt, Socket.IO et les dépendances e-mail."),
        ("Lancement backend", "npm run dev", "Démarrer l'API avec nodemon."),
        ("Configuration", "Fichier .env", "Définir le port, le secret JWT, MySQL et Gmail IMAP."),
    ], widths=[1800, 3600, 3960])
    p(doc, "Le fichier d'environnement contient notamment PORT, JWT_SECRET, DB_HOST, DB_NAME, DB_USER, DB_PWD, GMAIL_IMAP_USER, GMAIL_APP_PASSWORD et GMAIL_SYNC_INTERVAL_MS. Ces informations ne doivent pas être intégrées directement dans le code source, car elles peuvent contenir des paramètres sensibles.")

    h(doc, "4.5 Organisation générale du code", 2)
    p(doc, "Le projet est organisé selon une architecture client-serveur. Le frontend React gère l'affichage, les formulaires, la navigation et les retours utilisateurs. Le backend Express.js applique les règles métier et contrôle l'accès aux données. MySQL assure la persistance, tandis que Socket.IO prend en charge les échanges temps réel.")
    table(doc, ["Couche", "Rôle"], [
        ("Frontend React + Vite", "Interface utilisateur, routes, composants, formulaires, dashboards et affichage conditionnel."),
        ("API Express.js", "Endpoints REST, middlewares, validation, sécurité et orchestration des modules."),
        ("Services métier", "Règles fonctionnelles : création de ticket, affectation, résolution, meetings et e-mails."),
        ("Repositories SQL", "Accès aux tables MySQL et exécution des requêtes."),
        ("Socket.IO", "Authentification socket, rooms de tickets et diffusion des messages."),
        ("MySQL", "Stockage des employés, organisations, contacts, tickets, messages, meetings et e-mails."),
    ], widths=[2600, 6760])

    h(doc, "4.6 Arborescence du backend", 2)
    p(doc, "L'arborescence du backend montre une organisation modulaire. Chaque domaine fonctionnel possède ses propres routes, controllers, services et repositories lorsque cela est nécessaire. Cette structure facilite la lecture du code, la maintenance et l'évolution du système.")
    picture(doc, BACKEND_TREE, 6.25, "Figure 4.1 : Arborescence du backend backend_ticket_gestion.")
    p(doc, "Les dossiers principaux sont config, database, middleware, modules, socket et utils. Le dossier modules regroupe les fonctionnalités métier : auth, employees, organizations, contacts, tickets, comments, rooms, meetings, dashboard, activity et clientEmails. Cette séparation évite de concentrer toute la logique dans un seul fichier et permet d'identifier rapidement la responsabilité de chaque partie du code.")

    h(doc, "4.7 Implémentation du frontend", 2)
    p(doc, "Le frontend est construit avec React et Vite. Il repose sur des composants réutilisables, des pages par rôle et des routes protégées. Le fichier App.jsx définit les routes principales et le toaster global. Le fichier main.jsx initialise l'application et applique le thème. Le fichier index.css regroupe les variables visuelles globales, notamment pour le mode sombre et le mode clair.")
    table(doc, ["Élément frontend", "Rôle"], [
        ("RoleBasedRoute", "Protège les pages selon le token, le rôle et l'état actif du compte."),
        ("Sidebar", "Affiche les entrées de navigation adaptées à chaque rôle."),
        ("Register.jsx", "Gère l'authentification de l'utilisateur."),
        ("Tickets", "Liste, recherche, détail, affectation et suivi des tickets."),
        ("Messages", "Affiche les rooms et la messagerie temps réel."),
        ("Meetings", "Permet de planifier et suivre les réunions."),
        ("Dashboards", "Affiche les indicateurs d'activité selon le rôle."),
        ("AccessDenied", "Informe l'utilisateur lorsqu'une route lui est interdite."),
    ], widths=[2600, 6760])

    h(doc, "4.8 Concepts React utilisés", 2)
    table(doc, ["Concept", "Utilisation"], [
        ("Composants et props", "Découpage de l'interface en blocs réutilisables et transmission de données entre composants."),
        ("useState", "Gestion des formulaires, chargements, erreurs, onglets, tickets sélectionnés, modals et thème."),
        ("useEffect", "Chargement des données, vérification du compte actif, redirections, application du thème et connexion aux rooms."),
        ("useRef", "Conservation de la connexion socket, de la room active, des menus et de la zone de messages."),
        ("useMemo", "Optimisation des filtres, calculs de KPI, sélection de room et regroupement de tickets."),
        ("useCallback", "Stabilisation des fonctions joinRoom, sendMessage, loadRooms et chargements partagés."),
        ("Affichage conditionnel", "Affichage des boutons selon le rôle, gestion des états loading, error, empty state, statut et room non lue."),
    ], widths=[2500, 6860])

    h(doc, "4.9 Implémentation du backend", 2)
    p(doc, "Le backend est développé avec Node.js et Express.js. Le fichier index.js démarre le serveur HTTP, initialise Socket.IO, prépare le schéma complémentaire et lance la synchronisation Gmail lorsque la configuration est disponible. Le fichier app.js centralise le montage des routes Express.")
    table(doc, ["Couche backend", "Responsabilité"], [
        ("Routes", "Définissent les URL, les méthodes HTTP et les middlewares à appliquer."),
        ("Controllers", "Reçoivent les requêtes, appellent les services et construisent les réponses."),
        ("Services", "Appliquent les règles métier et les validations fonctionnelles."),
        ("Repositories", "Exécutent les requêtes SQL vers la base MySQL."),
        ("Middlewares", "Contrôlent l'authentification, les rôles, l'accès aux tickets et les tickets résolus."),
        ("Utils", "Regroupent les réponses API, la blacklist et les vérifications d'accès aux rooms."),
    ], widths=[2500, 6860])

    h(doc, "4.10 Authentification et sécurité", 2)
    p(doc, "L'authentification repose sur un nom d'utilisateur et un mot de passe. Le backend recherche l'employé, vérifie le mot de passe avec bcrypt, contrôle que le compte est actif puis génère un token JWT. Le token est stocké côté client et envoyé dans l'en-tête Authorization des requêtes protégées.")
    p(doc, "La sécurité est vérifiée à deux niveaux. Le frontend adapte les routes, la sidebar et les actions visibles selon le rôle. Le backend vérifie systématiquement le token, l'état actif du compte, le rôle et l'accès à la ressource. Ainsi, une requête API manuelle ne doit pas permettre de contourner les règles visibles dans l'interface.")
    table(doc, ["Mécanisme", "Rôle dans la sécurité"], [
        ("bcrypt", "Hachage et vérification des mots de passe."),
        ("JWT", "Authentification stateless des requêtes protégées."),
        ("auth.js", "Vérification du token, de la blacklist et du compte actif."),
        ("requireServiceDelivery", "Autorise uniquement le Service Delivery."),
        ("requireManager", "Autorise uniquement le Manager."),
        ("requireServiceDeliveryOrManager", "Autorise les lectures partagées SD/Manager."),
        ("requireTicketAccess", "Vérifie que l'utilisateur peut accéder au ticket demandé."),
        ("blockResolvedTicket", "Bloque les modifications d'un ticket clôturé."),
        ("Code 403", "Refuse les actions interdites côté backend."),
    ], widths=[3000, 6360])

    h(doc, "4.11 Gestion des tickets", 2)
    p(doc, "Le module tickets constitue le cœur fonctionnel du système. Un ticket représente une demande ou un incident lié à une organisation, à une application, à un type de problème et à un niveau de criticité. Lorsqu'il est créé, il reçoit le statut Pending. Les statuts utilisés sont Pending, In Progress, Warning, Critical et Resolved.")
    table(doc, ["Statut", "Signification"], [
        ("Pending", "Ticket créé, en attente de traitement ou d'affectation."),
        ("In Progress", "Traitement en cours."),
        ("Warning", "Situation nécessitant une surveillance particulière."),
        ("Critical", "Incident urgent ou fortement prioritaire."),
        ("Resolved", "Ticket clôturé ; les modifications métier sont bloquées."),
    ], widths=[2200, 7160])
    p(doc, "Le Service Delivery peut créer un ticket, le consulter, l'affecter à IT ou PKI et suivre son état. Les services IT et PKI consultent uniquement les tickets de leur périmètre. Le Manager dispose d'une supervision en lecture, ce qui limite les risques de modification non souhaitée des données métier.")

    h(doc, "4.12 Historique des affectations et commentaires", 2)
    p(doc, "Chaque affectation ou réaffectation d'un ticket vers IT ou PKI est conservée dans un historique. Cet historique indique le ticket concerné, le service affecté, la personne ayant réalisé l'action et la date de l'opération. Il améliore la traçabilité et permet de comprendre le chemin suivi par une demande.")
    p(doc, "Les commentaires complètent cette traçabilité. Ils permettent de documenter les actions réalisées, les remarques techniques et les décisions prises autour du ticket. Lorsqu'un ticket est résolu, l'ajout de nouveaux commentaires est bloqué afin de préserver la cohérence de la clôture.")

    h(doc, "4.13 Messagerie temps réel", 2)
    p(doc, "La messagerie interne repose sur Socket.IO. Chaque ticket peut être associé à une room permettant aux acteurs autorisés d'échanger en temps réel. Le frontend utilise le hook useChatRoom pour rejoindre une room, charger l'historique, envoyer un message et recevoir les nouveaux messages.")
    p(doc, "Les messages sont stockés en base et les états de lecture sont suivis individuellement. Ce fonctionnement évite de disperser les discussions dans des canaux externes et permet de conserver les échanges dans le contexte du ticket concerné.")

    h(doc, "4.14 Import des e-mails clients", 2)
    p(doc, "Le Service Delivery dispose d'une inbox e-mail intégrée. La synchronisation utilise Gmail IMAP avec ImapFlow. Le contenu des messages est analysé avec mailparser afin d'extraire l'expéditeur, le sujet, le corps du message et les pièces jointes. Les e-mails acceptés sont enregistrés dans MySQL avec un état lu/non lu propre à chaque employé autorisé.")
    p(doc, "Cette fonctionnalité rapproche le système d'une logique Email-to-Ticket. Elle contribue à réduire la perte d'information et facilite la transformation d'une demande reçue par e-mail en donnée exploitable dans le processus interne.")

    h(doc, "4.15 Base de données", 2)
    table(doc, ["Table", "Rôle"], [
        ("employees", "Stocke les employés, leurs identifiants, leur service et leur statut."),
        ("organizations", "Stocke les organisations clientes."),
        ("contacts", "Stocke les contacts associés aux organisations."),
        ("tickets", "Stocke les demandes et incidents suivis par le système."),
        ("rooms", "Stocke les conversations liées aux tickets."),
        ("messages", "Stocke les messages envoyés dans les rooms."),
        ("room_message_reads", "Stocke les indicateurs lu/non lu par employé."),
        ("comments", "Stocke les commentaires attachés aux tickets."),
        ("ticket_assignment_history", "Stocke l'historique des affectations IT/PKI."),
        ("meetings", "Stocke les réunions de suivi."),
        ("activity_logs", "Stocke les actions significatives du système."),
        ("client_emails", "Stocke les e-mails clients importés."),
        ("client_email_attachments", "Stocke les pièces jointes."),
        ("client_email_reads", "Stocke l'état de lecture des e-mails par employé."),
    ], widths=[2800, 6560])

    h(doc, "4.16 Interfaces et preuves de fonctionnement", 2)
    p(doc, "Les captures suivantes constituent des preuves visuelles de la réalisation. Elles peuvent être conservées dans ce chapitre ou déplacées en annexe selon les consignes de mise en page du mémoire final.")
    picture(doc, TICKETS_TABLE, 6.45, "Figure 4.2 : Exemple de tableau de suivi de tickets et de résolutions.")
    p(doc, "La capture du tableau de tickets illustre un jeu d'exemples utilisé pour représenter des incidents réels ou réalistes : numéro du ticket, niveau, utilisateur, entité, application, type de problème, description, statut, résolution et service affecté. Elle montre l'intérêt de structurer les informations afin de retrouver rapidement le contexte et la résolution d'un incident.")

    h(doc, "4.17 Tests fonctionnels", 2)
    table(doc, ["Test", "Scénario", "Résultat attendu", "Résultat obtenu", "Statut"], [
        ("TF01", "Connexion avec compte actif", "Accès au dashboard du rôle", "Accès obtenu", "Validé"),
        ("TF02", "Connexion avec mauvais mot de passe", "Refus de connexion", "Erreur affichée", "Validé"),
        ("TF03", "Accès à une route interdite", "Page Access Denied ou redirection", "Accès bloqué", "Validé"),
        ("TF04", "Désactivation d'un employé", "Compte refusé", "Session bloquée", "Validé"),
        ("TF05", "Création d'organisation", "Organisation enregistrée", "Liste actualisée", "Validé"),
        ("TF06", "Création de contact", "Contact rattaché", "Contact visible", "Validé"),
        ("TF07", "Création de ticket", "Ticket Pending créé", "Ticket et room créés", "Validé"),
        ("TF08", "Affectation à IT/PKI", "Historique enregistré", "Service affecté", "Validé"),
        ("TF09", "Modification de statut", "Statut mis à jour", "Statut actualisé", "Validé"),
        ("TF10", "Ajout de commentaire", "Commentaire conservé", "Historique enrichi", "Validé"),
        ("TF11", "Envoi de message", "Diffusion temps réel", "Message reçu", "Validé"),
        ("TF12", "Création de meeting", "Meeting créé", "Meeting visible", "Validé"),
        ("TF13", "Consultation dashboard", "Indicateurs affichés", "KPI visibles", "Validé"),
    ], widths=[850, 2650, 2500, 2300, 1060])

    h(doc, "4.18 Tests de sécurité", 2)
    table(doc, ["Test", "Scénario", "Résultat attendu", "Résultat obtenu"], [
        ("TS01", "Accès sans token", "Refus 401", "Requête refusée"),
        ("TS02", "Token invalide", "Refus 401", "Requête refusée"),
        ("TS03", "Token expiré", "Retour login ou refus", "Accès refusé"),
        ("TS04", "Compte désactivé", "Refus 403", "Accès bloqué"),
        ("TS05", "Rôle non autorisé", "Refus 403", "Action interdite"),
        ("TS06", "Modification d'un ticket résolu", "Blocage", "Action refusée"),
        ("TS07", "Accès PKI à un ticket IT", "Refus", "Accès interdit"),
        ("TS08", "Accès IT à un ticket PKI", "Refus", "Accès interdit"),
        ("TS09", "Vérification du code 403", "Message clair", "Erreur retournée"),
    ], widths=[900, 3300, 2700, 2460])

    h(doc, "4.19 Résultats obtenus", 2)
    p(doc, "La solution réalisée permet de centraliser les organisations, les contacts, les tickets, les commentaires, les messages, les meetings et les e-mails clients. Elle améliore l'organisation du support en structurant le cycle de vie des demandes. Elle renforce la traçabilité grâce aux historiques d'affectation, aux commentaires et aux messages conservés. Elle offre également au Manager une vue de supervision sans lui donner des droits de modification métier.")
    p(doc, "Sur le plan technique, le projet met en œuvre une séparation claire entre frontend, backend, services métier, repositories SQL et base MySQL. Le contrôle d'accès est appliqué dans l'interface et dans le backend, ce qui renforce la fiabilité du système.")

    h(doc, "4.20 Difficultés rencontrées", 2)
    bullets(doc, [
        "structuration du code en modules cohérents ;",
        "synchronisation des échanges entre frontend et backend ;",
        "définition précise des permissions selon les rôles ;",
        "gestion de Socket.IO et des rooms de discussion ;",
        "sécurisation des routes sensibles ;",
        "intégration de Gmail IMAP et traitement des e-mails ;",
        "cohérence du modèle relationnel MySQL ;",
        "tests des accès selon les rôles et l'état des tickets.",
    ])

    h(doc, "4.21 Limites du projet", 2)
    p(doc, "Certaines limites doivent être signalées. Les tests ont été principalement réalisés de manière manuelle et doivent être complétés par des tests automatisés. La solution n'a pas encore été évaluée sur un grand volume de tickets en production. La blacklist JWT est conservée en mémoire et devrait être remplacée par un mécanisme persistant. Le monitoring, les sauvegardes automatiques, l'authentification multifactorielle et les rapports avancés restent également à renforcer.")

    h(doc, "4.22 Conclusion du chapitre", 2)
    p(doc, "Ce chapitre a présenté la réalisation technique de l'application AGCE CRM. L'implémentation obtenue concrétise la conception proposée à travers une interface React, une API Express.js sécurisée, une base MySQL, une messagerie temps réel, une inbox e-mail et des tableaux de bord. Les tests fonctionnels et de sécurité confirment que la solution répond aux objectifs principaux du projet, tout en laissant des perspectives d'amélioration pour une version plus industrialisée.")


def appendix(doc):
    doc.add_page_break()
    h(doc, "Annexe : preuves visuelles complémentaires", 1)
    p(doc, "Cette annexe peut être utilisée si les captures ne sont pas conservées dans le corps principal du chapitre. Elle regroupe les éléments visuels transmis afin d'illustrer la réalisation.")
    picture(doc, BACKEND_TREE, 6.25, "Annexe A : Arborescence détaillée du backend.")
    picture(doc, TICKETS_TABLE, 6.45, "Annexe B : Exemple de tableau de suivi des tickets.")


def main():
    doc = Document()
    configure(doc)
    title_page(doc)
    note(doc, "Note sur le modèle fourni", "Le fichier is09.pages a été utilisé comme référence de mise en page : titres bleus, encadré de titre, tableaux simples et style académique sobre. Les annotations Pages exploitables n'ont pas été trouvées dans le paquet ; les remarques visibles disponibles ont donc été prises en compte à partir des éléments fournis et de la structure du modèle.")
    implementation(doc)
    appendix(doc)
    doc.save(DOCX_OUT)
    MD_OUT.write_text("\n".join(md), encoding="utf-8")
    print(DOCX_OUT)
    print(MD_OUT)


if __name__ == "__main__":
    main()
