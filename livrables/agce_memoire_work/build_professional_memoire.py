from __future__ import annotations

from pathlib import Path
from textwrap import dedent

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/issamouladsmane/Desktop/full_stuck_ticket_gestion")
WORK = ROOT / "livrables" / "agce_memoire_work"
OUT = ROOT / "livrables" / "Memoire_AGCE_CRM_Professionnel.docx"
WORK.mkdir(parents=True, exist_ok=True)

SCREENSHOTS = [
    ("/Users/issamouladsmane/Downloads/photo_2026-05-23 05.50.12.jpeg", "Interface d'authentification sécurisée"),
    ("/Users/issamouladsmane/Downloads/photo_2026-05-23 05.50.14.jpeg", "Tableau de bord d'administration des utilisateurs"),
    ("/Users/issamouladsmane/Downloads/photo_2026-05-23 05.50.36.jpeg", "Vue d'ensemble du Service Delivery"),
    ("/Users/issamouladsmane/Downloads/photo_2026-05-23 05.50.42.jpeg", "Création manuelle d'un ticket support"),
    ("/Users/issamouladsmane/Downloads/photo_2026-05-23 05.50.47.jpeg", "Détail et suivi opérationnel d'un ticket"),
    ("/Users/issamouladsmane/Downloads/photo_2026-05-23 05.50.19.jpeg", "Messagerie temps réel associée à un ticket"),
    ("/Users/issamouladsmane/Downloads/photo_2026-05-23 05.50.24.jpeg", "Tableau de bord analytique du manager"),
]


COLORS = {
    "ink": RGBColor(33, 37, 41),
    "blue": RGBColor(31, 78, 121),
    "light_blue": RGBColor(232, 238, 245),
    "dark": RGBColor(22, 36, 53),
    "muted": RGBColor(95, 105, 115),
    "gray": RGBColor(242, 244, 247),
}


def font_path() -> str | None:
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf",
        "/Library/Fonts/Arial.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ]
    for c in candidates:
        if Path(c).exists():
            return c
    return None


FONT = font_path()


def get_font(size: int, bold: bool = False):
    if FONT:
        try:
            return ImageFont.truetype(FONT, size=size)
        except Exception:
            pass
    return ImageFont.load_default()


def set_run_font(run, size=None, bold=None, italic=None, color=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for i, width in enumerate(widths):
            set_cell_width(row.cells[i], width)


def style_table(table, widths=None, header=True):
    if widths:
        set_table_width(table, widths)
    for r_idx, row in enumerate(table.rows):
        if r_idx == 0:
            tr_pr = row._tr.get_or_add_trPr()
            tbl_header = tr_pr.find(qn("w:tblHeader"))
            if tbl_header is None:
                tbl_header = OxmlElement("w:tblHeader")
                tr_pr.append(tbl_header)
            tbl_header.set(qn("w:val"), "true")
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.05
                for run in p.runs:
                    set_run_font(run, size=9.3, color=COLORS["ink"])
            if header and r_idx == 0:
                set_cell_shading(cell, "E8EEF5")
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.font.color.rgb = COLORS["blue"]


def add_field(paragraph, instruction: str, placeholder: str = ""):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    run._r.append(instr)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    if placeholder:
        t = OxmlElement("w:t")
        t.text = placeholder
        run._r.append(t)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=COLORS["muted"])
    add_field(paragraph, "PAGE")


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = COLORS["ink"]
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.28

    for name, size, color, before, after in [
        ("Heading 1", 16, COLORS["blue"], 18, 10),
        ("Heading 2", 13, COLORS["blue"], 12, 6),
        ("Heading 3", 12, COLORS["dark"], 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_style in ["List Bullet", "List Number"]:
        st = styles[list_style]
        st.font.name = "Calibri"
        st.font.size = Pt(10.5)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.18

    section.footer.paragraphs[0].text = ""
    footer_p = section.footer.paragraphs[0]
    run = footer_p.add_run("AGCE CRM - Mémoire de fin d'études")
    set_run_font(run, size=9, color=COLORS["muted"])
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = section.footer.add_paragraph()
    add_page_number(p)

    settings = doc.settings.element
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings.append(update)


def add_title(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True, color=COLORS["blue"])
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run_font(r, size=10.5)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_run_font(r, size=10.5)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text)
    set_run_font(r, size=9, italic=True, color=COLORS["muted"])


def add_note_box(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    set_cell_width(cell, 9360)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, size=10.5, bold=True, color=COLORS["blue"])
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(text)
    set_run_font(r2, size=10.2, color=COLORS["ink"])
    doc.add_paragraph()


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    style_table(table, widths=widths, header=True)
    return table


def wrap_text(draw, text, font, max_width):
    words = text.split()
    lines, cur = [], ""
    for w in words:
        test = (cur + " " + w).strip()
        bbox = draw.textbbox((0, 0), test, font=font)
        if bbox[2] - bbox[0] <= max_width or not cur:
            cur = test
        else:
            lines.append(cur)
            cur = w
    if cur:
        lines.append(cur)
    return lines


def draw_box(draw, xy, title, subtitle="", fill=(245, 247, 250), outline=(90, 110, 130)):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=16, fill=fill, outline=outline, width=2)
    title_font = get_font(28, bold=True)
    sub_font = get_font(21)
    tw = draw.textbbox((0, 0), title, font=title_font)
    draw.text((x1 + (x2 - x1 - (tw[2] - tw[0])) / 2, y1 + 22), title, fill=(22, 36, 53), font=title_font)
    if subtitle:
        lines = wrap_text(draw, subtitle, sub_font, x2 - x1 - 28)
        y = y1 + 66
        for line in lines[:3]:
            lw = draw.textbbox((0, 0), line, font=sub_font)
            draw.text((x1 + (x2 - x1 - (lw[2] - lw[0])) / 2, y), line, fill=(70, 76, 82), font=sub_font)
            y += 26


def arrow(draw, start, end, fill=(31, 78, 121)):
    draw.line([start, end], fill=fill, width=5)
    x1, y1 = start
    x2, y2 = end
    import math
    ang = math.atan2(y2 - y1, x2 - x1)
    size = 16
    pts = [
        (x2, y2),
        (x2 - size * math.cos(ang - 0.45), y2 - size * math.sin(ang - 0.45)),
        (x2 - size * math.cos(ang + 0.45), y2 - size * math.sin(ang + 0.45)),
    ]
    draw.polygon(pts, fill=fill)


def save_diagram(name, drawer):
    img = Image.new("RGB", (1600, 900), (252, 251, 247))
    draw = ImageDraw.Draw(img)
    drawer(draw)
    out = WORK / f"{name}.png"
    img.save(out, quality=95)
    return out


def make_diagrams():
    title_font = get_font(34, bold=True)
    small = get_font(22)

    def architecture(draw):
        draw.text((56, 42), "Architecture générale de la solution AGCE CRM", fill=(22, 36, 53), font=title_font)
        draw_box(draw, (70, 170, 390, 330), "Frontend", "React + Vite")
        draw_box(draw, (640, 170, 960, 330), "API Backend", "Express.js / REST")
        draw_box(draw, (1210, 170, 1530, 330), "Base de données", "MySQL")
        draw_box(draw, (640, 520, 960, 680), "Temps réel", "Socket.IO")
        draw_box(draw, (70, 520, 390, 680), "Utilisateur", "Admin, SD, IT, PKI, Manager")
        draw_box(draw, (1210, 520, 1530, 680), "E-mails clients", "IMAP + Mailparser")
        arrow(draw, (390, 250), (640, 250))
        arrow(draw, (960, 250), (1210, 250))
        arrow(draw, (800, 330), (800, 520))
        arrow(draw, (390, 600), (640, 600))
        arrow(draw, (1210, 600), (960, 600))
        draw.text((84, 750), "Principe : séparation claire entre interface, logique métier, persistance des données et échanges temps réel.", fill=(70, 76, 82), font=small)

    def use_case(draw):
        draw.text((56, 42), "Diagramme simplifié des cas d'utilisation", fill=(22, 36, 53), font=title_font)
        draw_box(draw, (550, 160, 1050, 720), "AGCE CRM", "Frontière du système")
        actions = ["S'authentifier", "Gérer organisations", "Créer / affecter ticket", "Traiter ticket", "Superviser indicateurs", "Communiquer"]
        y = 230
        for a in actions:
            draw.rounded_rectangle((620, y, 980, y + 54), radius=22, fill=(255, 255, 255), outline=(180, 188, 196), width=2)
            draw.text((650, y + 13), a, fill=(33, 37, 41), font=small)
            y += 70
        actors = [
            ((80, 220, 360, 290), "Administrateur"),
            ((80, 390, 360, 460), "Service Delivery"),
            ((80, 560, 360, 630), "Client"),
            ((1240, 220, 1520, 290), "Manager"),
            ((1240, 390, 1520, 460), "Équipe IT"),
            ((1240, 560, 1520, 630), "Équipe PKI"),
        ]
        for xy, label in actors:
            draw_box(draw, xy, label)
        for p1, p2 in [((360, 255), (550, 255)), ((360, 425), (550, 425)), ((360, 595), (550, 595)),
                       ((1240, 255), (1050, 255)), ((1240, 425), (1050, 425)), ((1240, 595), (1050, 595))]:
            draw.line([p1, p2], fill=(120, 130, 140), width=3)

    def class_model(draw):
        draw.text((56, 42), "Modèle conceptuel des classes principales", fill=(22, 36, 53), font=title_font)
        boxes = [
            ((100, 180, 390, 320), "User", "id, nom, email, service, statut"),
            ((500, 180, 790, 320), "Role / Service", "ADMIN, SD, Manager, IT, PKI"),
            ((900, 180, 1190, 320), "Ticket", "code, statut, priorité, description"),
            ((1210, 410, 1500, 550), "Comment", "texte, auteur, date"),
            ((900, 620, 1190, 760), "Message", "room, sender, texte"),
            ((500, 620, 790, 760), "Room", "ticket, services autorisés"),
            ((100, 620, 390, 760), "Organization", "nom, secteur, contact"),
            ((100, 410, 390, 550), "Contact", "nom, email, type, téléphone"),
        ]
        for xy, title, attrs in boxes:
            draw_box(draw, xy, title, attrs)
        for s, e in [((390, 250), (500, 250)), ((790, 250), (900, 250)), ((1045, 320), (1045, 620)),
                     ((900, 690), (790, 690)), ((500, 690), (390, 690)), ((245, 620), (245, 550)),
                     ((390, 480), (900, 250)), ((1190, 480), (1210, 480))]:
            arrow(draw, s, e, fill=(110, 120, 130))

    def sequence(draw):
        draw.text((56, 42), "Flux de création et d'affectation d'un ticket", fill=(22, 36, 53), font=title_font)
        actors = [("Utilisateur SD", 150), ("Frontend", 430), ("API Backend", 710), ("MySQL", 990), ("Service IT/PKI", 1270)]
        for label, x in actors:
            draw_box(draw, (x - 90, 150, x + 90, 220), label)
            draw.line([(x, 220), (x, 780)], fill=(160, 160, 160), width=2)
        steps = [
            (150, 430, 285, "Soumission formulaire"),
            (430, 710, 365, "POST /api/tickets + JWT"),
            (710, 990, 445, "INSERT ticket, room, logs"),
            (990, 710, 525, "Confirmation transaction"),
            (710, 1270, 605, "Notification / accès room"),
            (710, 430, 685, "HTTP 201 + données ticket"),
        ]
        for x1, x2, y, label in steps:
            arrow(draw, (x1, y), (x2, y), fill=(31, 78, 121))
            draw.text((min(x1, x2) + 22, y - 32), label, fill=(33, 37, 41), font=small)

    def process(draw):
        draw.text((56, 42), "Processus de développement adopté", fill=(22, 36, 53), font=title_font)
        labels = ["Analyse", "Conception", "UML", "Développement", "Intégration BDD", "Tests", "Validation"]
        x = 70
        for i, label in enumerate(labels):
            draw_box(draw, (x, 330, x + 180, 450), f"{i+1}", label)
            if i < len(labels) - 1:
                arrow(draw, (x + 180, 390), (x + 250, 390))
            x += 220
        draw.text((80, 580), "La démarche retenue est incrémentale : chaque module est analysé, modélisé, implémenté, testé puis amélioré.", fill=(70, 76, 82), font=small)

    return {
        "architecture": save_diagram("architecture", architecture),
        "use_case": save_diagram("use_case", use_case),
        "class_model": save_diagram("class_model", class_model),
        "sequence": save_diagram("sequence", sequence),
        "process": save_diagram("process", process),
    }


def add_picture(doc, path, width=6.1, alt_text=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    inline_shape = run.add_picture(str(path), width=Inches(width))
    alt = alt_text or Path(path).stem.replace("_", " ")
    for doc_pr in inline_shape._inline.xpath(".//wp:docPr"):
        doc_pr.set("descr", alt)
        doc_pr.set("title", alt)


def cover(doc):
    for text, size, bold in [
        ("République Algérienne Démocratique et Populaire", 12, True),
        ("Ministère de l'Enseignement Supérieur et de la Recherche Scientifique", 11, False),
        ("Université Saad Dahleb - Blida 1", 11, True),
        ("Faculté des Sciences - Département d'Informatique", 11, False),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, color=COLORS["ink"])

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Mémoire de fin d'études")
    set_run_font(r, size=18, bold=True, color=COLORS["blue"])
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p2.add_run("Présenté en vue de l'obtention du diplôme de Licence en Informatique")
    set_run_font(r, size=12, color=COLORS["muted"])
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p3.add_run("Option : Systèmes d'Information et Ingénierie Logicielle (ISIL)")
    set_run_font(r, size=12, color=COLORS["muted"])

    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "E8EEF5")
    set_cell_width(cell, 8800)
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title = cell.paragraphs[0].add_run("AGCE CRM\nSystème de gestion des tickets et de la relation client")
    set_run_font(title, size=22, bold=True, color=COLORS["dark"])
    subtitle = cell.add_paragraph().add_run("Conception et réalisation d'une application web sécurisée pour l'AGCE")
    set_run_font(subtitle, size=12, italic=True, color=COLORS["muted"])

    doc.add_paragraph()
    meta = doc.add_table(rows=4, cols=2)
    rows = [
        ("Préparé par", "Mr. Laceb Karim\nMr. Ouladsmane Issam"),
        ("Encadré par", "Mme Arkam Meriem\nM. Lekhchine Sami"),
        ("Organisme d'accueil", "Autorité Gouvernementale de Certification Électronique (AGCE)"),
        ("Année universitaire", "2025 / 2026"),
    ]
    for i, (a, b) in enumerate(rows):
        meta.cell(i, 0).text = a
        meta.cell(i, 1).text = b
    style_table(meta, widths=[2400, 6200], header=False)
    for row in meta.rows:
        set_cell_shading(row.cells[0], "F2F4F7")
        for p in row.cells[0].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = COLORS["blue"]
    doc.add_page_break()


def front_matter(doc):
    add_title(doc, "Dédicaces", 1)
    add_title(doc, "Dédicace de Mr. Laceb Karim", 2)
    add_para(doc, "Je dédie ce travail à mes parents, pour leur soutien constant, leurs sacrifices et leur confiance tout au long de mon parcours. Je l'adresse également à ma famille, mes proches et mes amis, dont les encouragements ont accompagné chaque étape de cette réalisation.")
    add_title(doc, "Dédicace de Mr. Ouladsmane Issam", 2)
    add_para(doc, "Je dédie ce mémoire à ma famille, et particulièrement à mes parents, pour leur patience, leurs sacrifices et leur soutien indéfectible. Je remercie aussi mes proches, mes amis et mes camarades pour leur présence et leur solidarité durant ce parcours.")

    add_title(doc, "Remerciements", 1)
    add_para(doc, "Nous remercions Dieu le Tout-Puissant de nous avoir accordé la force, la patience et la persévérance nécessaires pour mener à bien ce travail. Nous exprimons notre profonde gratitude à Mme Arkam Meriem, notre encadrante universitaire, pour son accompagnement, ses orientations méthodologiques et la rigueur de ses remarques.")
    add_para(doc, "Nous adressons également nos sincères remerciements à M. Lekhchine Sami, notre encadrant au sein de l'Autorité Gouvernementale de Certification Électronique, pour son accueil, sa disponibilité, sa confiance et ses conseils techniques. Nos remerciements s'adressent aussi à l'ensemble du personnel de l'AGCE pour les informations, les échanges et l'appui apportés durant la réalisation du projet.")
    add_para(doc, "Enfin, nous remercions les membres du jury pour l'honneur qu'ils nous font en acceptant d'évaluer ce mémoire et pour l'intérêt porté à notre travail.")

    add_title(doc, "Résumé", 1)
    add_para(doc, "Ce mémoire présente la conception et la réalisation d'une application web sécurisée de gestion des tickets et de la relation client destinée à l'Autorité Gouvernementale de Certification Électronique (AGCE). Le projet répond à un besoin de centralisation des demandes, de traçabilité du traitement, de collaboration entre services internes et de supervision des indicateurs opérationnels. La solution proposée permet de gérer les organisations, les contacts, les tickets, les commentaires, les messages en temps réel, les réunions de suivi et les tableaux de bord analytiques. Elle repose sur React et Vite pour l'interface utilisateur, Express.js pour l'API, MySQL pour la persistance des données et Socket.IO pour la communication temps réel. La sécurité est assurée par l'authentification JWT, le hachage bcrypt, la vérification de l'état actif des comptes et un contrôle d'accès fondé sur les rôles. Les résultats obtenus montrent une amélioration de la visibilité, de la traçabilité et de l'organisation interne du support.")
    add_para(doc, "Mots-clés : AGCE, CRM, gestion de tickets, traçabilité, RBAC, React, Express.js, MySQL, Socket.IO.")

    add_title(doc, "ملخص", 1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("يعرض هذا البحث تصميم وإنجاز تطبيق ويب آمن لإدارة التذاكر وعلاقة العملاء لفائدة سلطة التصديق الإلكتروني الحكومية. يهدف المشروع إلى مركزية طلبات العملاء، تتبع معالجة الحوادث، تنظيم الصلاحيات بين المصالح الداخلية، وتحسين الرؤية التحليلية للمسؤولين. يعتمد النظام على React وExpress.js وMySQL وSocket.IO، مع تأمين الولوج بواسطة JWT وتشفير كلمات المرور باستعمال bcrypt والتحكم في الصلاحيات حسب الدور. تبرز النتائج قدرة الحل المقترح على تحسين التنظيم الداخلي، تتبع العمليات، والتعاون بين الفرق.")
    set_run_font(r, size=11)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = p2.add_run("الكلمات المفتاحية: AGCE، CRM، إدارة التذاكر، التتبع، RBAC، الأمن التطبيقي.")
    set_run_font(r2, size=11, bold=True)

    add_title(doc, "Abstract", 1)
    add_para(doc, "This thesis presents the design and implementation of a secure web-based ticket management and customer relationship management application for the Governmental Electronic Certification Authority (AGCE). The project addresses the need to centralize customer requests, trace support actions, improve collaboration between internal services and provide analytical supervision. The solution manages organizations, contacts, tickets, comments, real-time messages, meetings and dashboards. It relies on React and Vite for the frontend, Express.js for the API, MySQL for data persistence and Socket.IO for real-time communication. Security is ensured through JWT authentication, bcrypt password hashing, account-status verification and role-based access control. The obtained prototype improves operational visibility, traceability and internal support organization.")
    add_para(doc, "Keywords: AGCE, CRM, ticket management, traceability, RBAC, React, Express.js, MySQL, Socket.IO.")
    doc.add_page_break()

    add_title(doc, "Table des matières", 1)
    p = doc.add_paragraph()
    add_field(p, r'TOC \o "1-3" \h \z \u', "Mettre à jour la table des matières dans Word.")
    add_note_box(doc, "Note de mise à jour", "La table des matières est insérée sous forme de champ automatique. Dans Word : clic droit sur la table, puis « Mettre à jour le champ » pour actualiser les numéros de page après toute modification.")
    doc.add_page_break()

    add_title(doc, "Liste des figures", 1)
    figures = [
        "Figure 1 : processus de développement adopté.",
        "Figure 2 : diagramme simplifié des cas d'utilisation.",
        "Figure 3 : architecture générale de l'application AGCE CRM.",
        "Figure 4 : modèle conceptuel des classes principales.",
        "Figure 5 : séquence de création et d'affectation d'un ticket.",
        "Figure 6 : interface d'authentification sécurisée.",
        "Figure 7 : tableau de bord d'administration des utilisateurs.",
        "Figure 8 : vue d'ensemble du Service Delivery.",
        "Figure 9 : création manuelle d'un ticket support.",
        "Figure 10 : détail et suivi opérationnel d'un ticket.",
        "Figure 11 : messagerie temps réel associée à un ticket.",
        "Figure 12 : tableau de bord analytique du manager.",
    ]
    add_bullets(doc, figures)

    add_title(doc, "Liste des tableaux", 1)
    tables = [
        "Tableau 1 : comparaison synthétique des solutions de gestion des tickets.",
        "Tableau 2 : synthèse des besoins fonctionnels.",
        "Tableau 3 : acteurs et responsabilités principales.",
        "Tableau 4 : matrice synthétique du contrôle d'accès basé sur les rôles.",
        "Tableau 5 : modèle logique de données synthétique.",
        "Tableau 6 : technologies et outils utilisés.",
        "Tableau 7 : scénarios de validation fonctionnelle.",
        "Tableau 8 : perspectives d'évolution prioritaires.",
        "Tableau 9 : checklist finale de conformité au guide.",
    ]
    add_bullets(doc, tables)
    doc.add_page_break()

    add_title(doc, "Liste des abréviations", 1)
    rows = [
        ("AGCE", "Autorité Gouvernementale de Certification Électronique"),
        ("API", "Application Programming Interface"),
        ("CRM", "Customer Relationship Management"),
        ("CRUD", "Create, Read, Update, Delete"),
        ("IT", "Information Technology"),
        ("ITSM", "Information Technology Service Management"),
        ("JWT", "JSON Web Token"),
        ("LRAO", "Local Registration Authority Officer"),
        ("MFA", "Multi-Factor Authentication"),
        ("PKI", "Public Key Infrastructure"),
        ("RBAC", "Role-Based Access Control"),
        ("SD", "Service Delivery"),
        ("UML", "Unified Modeling Language"),
    ]
    add_table(doc, ["Abréviation", "Signification"], rows, widths=[2200, 7000])
    doc.add_page_break()


def chapter_intro(doc):
    add_title(doc, "Introduction générale", 1)
    add_para(doc, "La transformation numérique des organisations publiques impose la mise en place de systèmes d'information capables de concilier efficacité opérationnelle, sécurité des données, traçabilité des échanges et maîtrise des responsabilités. Dans un contexte lié à la confiance numérique, ces exigences deviennent particulièrement importantes, car la qualité du traitement des demandes internes et externes influence directement la continuité du service.")
    add_para(doc, "L'Autorité Gouvernementale de Certification Électronique (AGCE) occupe une position stratégique dans l'écosystème numérique algérien. Ses activités liées à la certification électronique, à la signature numérique, à l'authentification et à la sécurisation des échanges nécessitent une organisation rigoureuse et un suivi fiable des incidents, demandes clients et échanges techniques.")
    add_para(doc, "Le problème étudié concerne la gestion des tickets et de la relation client au sein de l'AGCE. Lorsque les demandes sont dispersées entre plusieurs canaux, il devient difficile d'identifier leur état d'avancement, de mesurer les délais, de connaître le service responsable, de reconstituer l'historique d'un incident et d'obtenir une vision globale de l'activité.")
    add_note_box(doc, "Question de recherche", "Comment concevoir et réaliser une application web sécurisée permettant à l'AGCE de centraliser la gestion des tickets et de la relation client, tout en garantissant la traçabilité des actions, la séparation des rôles et l'amélioration de l'efficacité organisationnelle ?")
    add_title(doc, "Objectifs du projet", 2)
    add_para(doc, "L'objectif principal consiste à concevoir et développer une solution personnalisée permettant de centraliser la gestion des tickets, d'améliorer la traçabilité des demandes et de faciliter la collaboration entre les différents services.")
    add_bullets(doc, [
        "analyser le contexte métier et les limites de la gestion dispersée des demandes ;",
        "définir les acteurs, les rôles et les permissions nécessaires au fonctionnement de la solution ;",
        "modéliser les processus, les cas d'utilisation, les classes et les flux principaux ;",
        "réaliser une application web sécurisée fondée sur une architecture client-serveur ;",
        "valider la solution par des scénarios fonctionnels, des tests de sécurité et une analyse des résultats obtenus.",
    ])
    add_title(doc, "Méthodologie de recherche", 2)
    add_para(doc, "La démarche adoptée est une recherche appliquée à dominante qualitative et comparative. Elle combine l'analyse documentaire, l'observation des processus existants, l'étude des besoins fonctionnels, la comparaison de solutions CRM/helpdesk et la validation du prototype à travers des scénarios représentatifs. Les acteurs étudiés sont les administrateurs, le Service Delivery, les équipes IT et PKI, le manager et les contacts clients.")
    add_title(doc, "Organisation du mémoire", 2)
    add_para(doc, "Le mémoire est structuré en trois chapitres. Le premier présente le contexte, l'organisme d'accueil, la problématique et l'état de l'art. Le deuxième expose l'analyse des besoins et la conception du système. Le troisième décrit la réalisation, les résultats, la validation, les limites et les perspectives d'évolution.")


def chapter_one(doc):
    add_title(doc, "Chapitre I - Contexte général et état de l'art", 1)
    add_title(doc, "Présentation de l'AGCE", 2)
    add_para(doc, "L'Autorité Gouvernementale de Certification Électronique est un organisme public algérien chargé de la gestion de la certification et de la signature électroniques dans la branche gouvernementale. Elle intervient dans un domaine où la confiance numérique, la souveraineté des données et la sécurité des échanges sont des exigences majeures.")
    add_para(doc, "Les services concernés par le projet sont principalement la sous-direction de l'enregistrement, assimilée au Service Delivery, et la sous-direction d'exploitation des infrastructures de gestion de clés, assimilée au périmètre PKI. Le Service Delivery reçoit, qualifie et suit les demandes ; les services techniques IT et PKI interviennent sur les incidents qui leur sont affectés ; le manager supervise l'activité ; l'administrateur gère les comptes internes.")
    add_title(doc, "Contexte métier et problématique", 2)
    add_para(doc, "La gestion des demandes clients joue un rôle central dans l'organisation du support. Elle permet de formaliser les incidents, de conserver un historique, de prioriser les traitements et de coordonner les responsabilités. Sans outil centralisé, les informations peuvent être réparties entre e-mails, appels, fichiers, discussions informelles et notes séparées.")
    add_para(doc, "Dans le contexte de l'AGCE, cette dispersion pose quatre difficultés principales : perte de visibilité sur l'état d'avancement, traçabilité insuffisante des actions, affectation parfois lente des tickets et manque d'indicateurs consolidés pour la supervision. Ces limites justifient la création d'un système adapté au fonctionnement interne et aux contraintes de sécurité de l'organisme.")
    add_title(doc, "Fondements CRM et ITSM", 2)
    add_para(doc, "Le CRM vise à organiser la relation entre une organisation et ses clients ou usagers. Il ne se limite pas à un carnet d'adresses : il regroupe les interactions, l'historique des demandes, les informations relatives aux organisations et les indicateurs permettant d'améliorer le service rendu. Les travaux sur le CRM insistent sur la nécessité de structurer les processus relationnels et de mesurer leur impact sur la performance organisationnelle [28], [29].")
    add_para(doc, "L'ITSM, pour sa part, fournit un cadre de gestion des services informatiques. Les pratiques de service desk, de gestion des incidents, de suivi des demandes et d'amélioration continue y occupent une place centrale. ITIL 4 met l'accent sur la co-création de valeur, la gestion structurée des incidents et l'adaptation des processus au contexte de l'organisation [30].")
    add_title(doc, "Solutions existantes et limites", 2)
    add_para(doc, "Les solutions de marché se répartissent généralement en plateformes SaaS, outils open source déployables localement et solutions développées sur mesure. Les plateformes SaaS offrent une richesse fonctionnelle importante, mais elles peuvent soulever des contraintes de souveraineté et de confidentialité. Les outils open source permettent un meilleur contrôle local, mais restent génériques et nécessitent des adaptations importantes pour intégrer les règles de l'AGCE.")
    rows = [
        ("Jira Service Management", "SaaS / Cloud", "Workflow puissant, intégrations, suivi complet", "Hébergement externe et adaptation métier coûteuse"),
        ("Zendesk", "SaaS / Cloud", "Expérience client, base de connaissances, reporting", "Dépendance cloud et personnalisation limitée au contexte sensible"),
        ("GLPI", "Open source", "Déploiement local, ITSM, inventaire", "Logique générique, adaptation RBAC et CRM spécifique nécessaire"),
        ("osTicket", "Open source", "Ticketing simple, léger, déployable localement", "Fonctions analytiques et CRM limitées"),
        ("AGCE CRM", "Solution personnalisée", "Rôles métiers, traçabilité, tickets, CRM, dashboards, messagerie", "Nécessite maintenance interne et durcissement avant production"),
    ]
    add_table(doc, ["Solution", "Type", "Points forts", "Limites"], rows, widths=[1900, 1700, 3000, 2760])
    add_caption(doc, "Tableau 1 : comparaison synthétique des solutions de gestion des tickets.")
    add_title(doc, "Justification de la solution personnalisée", 2)
    add_para(doc, "Le choix d'une solution personnalisée se justifie par la nécessité d'intégrer les règles métier dans le système lui-même. AGCE CRM permet de définir précisément les rôles, de contrôler l'accès aux tickets selon le service, de centraliser les organisations et contacts, d'associer les discussions et réunions aux tickets et de fournir des indicateurs destinés à la supervision.")
    add_note_box(doc, "Positionnement du projet", "La valeur ajoutée du projet ne réside pas seulement dans la création d'un ticket, mais dans la construction d'un parcours complet : qualification, affectation, traitement, communication, supervision, verrouillage après résolution et conservation de l'historique.")


def chapter_two(doc, diagrams):
    add_title(doc, "Chapitre II - Analyse et conception du système", 1)
    add_title(doc, "Processus de développement", 2)
    add_para(doc, "La réalisation a été conduite selon une démarche incrémentale par modules. Plutôt que d'affirmer une méthode agile complète sans en documenter tous les artefacts, le projet retient une logique progressive : analyse, conception, modélisation, implémentation, intégration, tests et validation.")
    add_picture(doc, diagrams["process"], alt_text="Processus de développement adopté pour AGCE CRM")
    add_caption(doc, "Figure 1 : processus de développement adopté.")
    add_title(doc, "Analyse des besoins", 2)
    add_para(doc, "Les besoins fonctionnels définissent ce que le système doit accomplir pour répondre aux difficultés identifiées. Les besoins non fonctionnels précisent les critères de qualité, de sécurité, de maintenabilité et de performance.")
    rows = [
        ("Gestion des utilisateurs", "Création, modification, désactivation et suivi des comptes employés."),
        ("Gestion CRM", "Gestion des organisations clientes, des contacts et des informations associées."),
        ("Gestion des tickets", "Création, classification, affectation, suivi du statut et verrouillage après résolution."),
        ("Communication", "Commentaires, messagerie temps réel par room et pièces jointes associées aux échanges."),
        ("Réunions", "Planification, acceptation, rejet et rattachement éventuel à un ticket."),
        ("Supervision", "Tableaux de bord, indicateurs de volume, répartition, statut et activité."),
        ("E-mail to ticket", "Lecture des e-mails entrants et préparation de leur conversion en demandes structurées."),
    ]
    add_table(doc, ["Besoin fonctionnel", "Description"], rows, widths=[2800, 6560])
    add_caption(doc, "Tableau 2 : synthèse des besoins fonctionnels.")
    add_bullets(doc, [
        "Sécurité : authentification, contrôle d'accès par rôle et protection des routes sensibles.",
        "Traçabilité : conservation des commentaires, messages, affectations et journaux d'activité.",
        "Souveraineté : logique de déploiement local et maîtrise de la base de données.",
        "Maintenabilité : séparation frontend, backend, modules métier et couche de persistance.",
        "Ergonomie : interface claire, responsive et adaptée aux profils utilisateurs.",
    ])
    add_title(doc, "Acteurs du système", 2)
    rows = [
        ("Administrateur", "Gère les comptes employés, les statuts et les services internes."),
        ("Service Delivery", "Crée les organisations, contacts et tickets ; affecte les tickets aux services."),
        ("Manager", "Consulte les tableaux de bord, supervise l'activité et accède en lecture seule."),
        ("Équipe PKI", "Traite les tickets affectés au périmètre PKI."),
        ("Équipe IT", "Traite les tickets affectés au périmètre informatique."),
        ("Client / Contact", "Origine de la demande ou de l'incident à travers les canaux de support."),
    ]
    add_table(doc, ["Acteur", "Responsabilités"], rows, widths=[2300, 7060])
    add_caption(doc, "Tableau 3 : acteurs et responsabilités principales.")
    add_title(doc, "Diagramme de cas d'utilisation", 2)
    add_picture(doc, diagrams["use_case"], alt_text="Diagramme simplifié des cas d'utilisation AGCE CRM")
    add_caption(doc, "Figure 2 : diagramme simplifié des cas d'utilisation.")
    add_para(doc, "Chaque acteur dispose d'un périmètre d'action adapté à son rôle. L'administrateur ne traite pas les tickets métier ; le Service Delivery pilote les demandes ; les équipes IT et PKI interviennent uniquement sur les tickets autorisés ; le manager supervise sans modifier les données opérationnelles.")
    add_title(doc, "Matrice RBAC", 2)
    rows = [
        ("Gestion employés", "CRUD", "Non", "Non", "Non", "Non"),
        ("Organisations / contacts", "Non", "CRUD", "Lecture", "Non", "Non"),
        ("Création ticket", "Non", "Oui", "Non", "Non", "Non"),
        ("Affectation IT/PKI", "Non", "Oui", "Lecture", "Non", "Non"),
        ("Traitement ticket", "Non", "Suivi", "Lecture", "Tickets PKI", "Tickets IT"),
        ("Messages / commentaires", "Non", "Autorisés", "Lecture", "Rooms PKI", "Rooms IT"),
        ("Dashboards", "Admin", "SD", "Manager", "PKI", "IT"),
        ("Ticket résolu", "Lecture", "Lecture", "Lecture", "Lecture verrouillée", "Lecture verrouillée"),
    ]
    add_table(doc, ["Module", "ADMIN", "SD", "Manager", "PKI", "IT"], rows, widths=[2500, 1300, 1300, 1500, 1380, 1380])
    add_caption(doc, "Tableau 4 : matrice synthétique du contrôle d'accès basé sur les rôles.")
    add_title(doc, "Architecture générale", 2)
    add_picture(doc, diagrams["architecture"], alt_text="Architecture générale de l'application AGCE CRM")
    add_caption(doc, "Figure 3 : architecture générale de l'application AGCE CRM.")
    add_para(doc, "L'architecture repose sur une séparation nette des responsabilités. Le frontend React/Vite prend en charge l'interface et les parcours utilisateurs. Le backend Express.js applique les règles métier, les middlewares de sécurité et les contrats d'API. MySQL assure la persistance structurée des données, tandis que Socket.IO permet les échanges temps réel dans les rooms associées aux tickets.")
    add_title(doc, "Diagramme de classes", 2)
    add_picture(doc, diagrams["class_model"], alt_text="Modèle conceptuel des classes principales")
    add_caption(doc, "Figure 4 : modèle conceptuel des classes principales.")
    add_para(doc, "Le modèle de données s'articule autour des entités User, Service, Organization, Contact, Ticket, Room, Message et Comment. Un utilisateur peut créer ou traiter des tickets selon son rôle ; un ticket peut être associé à une organisation, contenir des commentaires, disposer d'une room de discussion et être rattaché à des réunions de suivi.")
    add_title(doc, "Modèle logique de données", 2)
    rows = [
        ("services", "id, name", "Référence les services internes : ADMIN, SD, Manager, IT, PKI."),
        ("employees", "id, firstName, lastName, email, userName, password, status, service_id", "Gère les utilisateurs internes et leur service."),
        ("organizations", "id, name, industry, email, phone, address, status", "Représente les organisations clientes."),
        ("contacts", "id, name, type, email, phone, job_title, organization_id", "Représente les contacts liés aux organisations."),
        ("tickets", "id, request_code, organization_id, application, issue_type, issue_level, status, created_by", "Porte les demandes et incidents."),
        ("rooms / messages", "room_id, sender_id, text, createdAt", "Assure la communication temps réel liée aux tickets."),
        ("meetings", "title, start_time_utc, organizer_id, invitee_id, ticket_id, status", "Organise les réunions de suivi."),
        ("activity_logs", "actor, action_type, entity_type, metadata", "Trace les actions significatives."),
    ]
    add_table(doc, ["Table", "Champs principaux", "Rôle"], rows, widths=[1800, 3900, 3660])
    add_caption(doc, "Tableau 5 : modèle logique de données synthétique.")
    add_title(doc, "Diagramme de séquence", 2)
    add_picture(doc, diagrams["sequence"], alt_text="Diagramme de séquence création et affectation d'un ticket")
    add_caption(doc, "Figure 5 : séquence de création et d'affectation d'un ticket.")
    add_para(doc, "Lorsqu'un agent du Service Delivery crée un ticket, le frontend transmet la demande au backend avec le jeton JWT. Le backend vérifie l'identité, applique les règles de rôle, enregistre le ticket, crée la room associée, journalise l'action et retourne les données nécessaires à l'interface.")
    add_title(doc, "Sécurité de conception", 2)
    add_para(doc, "La sécurité est intégrée dès la conception. L'authentification repose sur JWT ; les mots de passe sont hachés avec bcrypt ; l'état actif du compte est vérifié à chaque requête ; les routes sensibles sont protégées par des middlewares ; l'accès aux tickets dépend du rôle et du service affecté.")
    add_bullets(doc, [
        "ADMIN : gestion des employés uniquement ;",
        "SD : gestion opérationnelle des clients, tickets, affectations et meetings ;",
        "Manager : supervision et lecture seule ;",
        "IT / PKI : accès limité aux tickets affectés au service concerné ;",
        "Tickets résolus : verrouillage des opérations pour préserver l'historique.",
    ])


def chapter_three(doc):
    add_title(doc, "Chapitre III - Réalisation, implémentation et tests", 1)
    add_para(doc, "Ce chapitre présente la traduction concrète de la conception en une application fonctionnelle. Il décrit l'environnement de réalisation, l'organisation technique du frontend et du backend, les principaux modules développés, les mécanismes de sécurité appliqués, puis les tests effectués pour valider le comportement du système. Cette partie respecte une logique expérimentale : chaque choix d'implémentation est relié au besoin métier identifié et aux contraintes de sécurité propres au contexte de l'AGCE.")
    add_para(doc, "L'objectif de la réalisation n'est pas seulement de produire une interface utilisable. Il s'agit également de garantir la centralisation des demandes, la séparation des responsabilités, la traçabilité des traitements et la protection des accès. Pour cette raison, les vérifications ne sont pas limitées au frontend ; elles sont aussi appliquées au niveau du backend et de la base de données.")
    add_title(doc, "Environnement de réalisation", 2)
    rows = [
        ("Frontend", "React 19, Vite 8, CSS, Recharts, lucide-react", "Construction d'une interface dashboard dynamique, responsive et adaptée aux rôles."),
        ("Backend", "Node.js, Express.js 5, mysql2", "Exposition de l'API REST, application des règles métier et accès à MySQL."),
        ("Temps réel", "Socket.IO 4", "Gestion des rooms de discussion associées aux tickets."),
        ("Sécurité", "JWT, bcrypt, middlewares Express", "Authentification, hachage des mots de passe, vérification des rôles et de l'état actif du compte."),
        ("Base de données", "MySQL", "Persistance relationnelle des employés, organisations, contacts, tickets, messages, meetings et e-mails."),
        ("E-mails", "ImapFlow, mailparser", "Synchronisation de la boîte Service Delivery et extraction des messages clients."),
        ("Tests", "Navigateur, scénarios métier, requêtes API", "Validation fonctionnelle, contrôle des erreurs et vérification des accès interdits."),
    ]
    add_table(doc, ["Couche", "Technologies", "Justification"], rows, widths=[1800, 3700, 3860])
    add_caption(doc, "Tableau 6 : environnement logiciel et rôle des technologies utilisées.")

    add_title(doc, "Organisation technique de l'application", 2)
    add_para(doc, "L'application suit une architecture client-serveur. Le frontend React gère les vues, les formulaires, la navigation et les états d'interface. Le backend Express.js reçoit les requêtes, vérifie l'identité de l'utilisateur, applique les règles métier et délègue l'accès aux données à des repositories SQL. MySQL conserve les informations persistantes, tandis que Socket.IO complète l'API REST pour les échanges en temps réel.")
    rows = [
        ("Routes", "Déclarent les URL, les méthodes HTTP et les middlewares nécessaires.", "Clarifier les points d'entrée de l'API."),
        ("Controllers", "Reçoivent la requête, appellent le service et formatent la réponse.", "Éviter de mélanger HTTP et logique métier."),
        ("Services", "Regroupent les règles fonctionnelles et les validations.", "Centraliser les décisions métier."),
        ("Repositories", "Exécutent les requêtes SQL vers MySQL.", "Isoler la persistance des données."),
        ("Middlewares", "Vérifient JWT, rôle, état actif et accès au ticket.", "Bloquer les accès non autorisés avant le traitement."),
    ]
    add_table(doc, ["Couche backend", "Responsabilité", "Intérêt"], rows, widths=[1900, 4150, 3310])
    add_caption(doc, "Tableau 7 : organisation par couches du backend.")

    add_title(doc, "Implémentation du frontend", 2)
    add_para(doc, "Le frontend est structuré autour de pages et de composants réutilisables. Les routes sont séparées selon les rôles : administrateur, Service Delivery, manager, PKI et IT. Cette séparation améliore la lisibilité de l'interface et réduit les risques d'afficher une action non pertinente pour l'utilisateur connecté.")
    add_para(doc, "Plusieurs mécanismes React sont utilisés. useState gère les formulaires, les listes chargées, les états de chargement et les messages d'erreur. useEffect charge les données depuis l'API, vérifie la session, applique le thème et rejoint les rooms Socket.IO. useRef conserve certaines références techniques, comme la zone de messages ou la connexion socket. useMemo et useCallback sont employés pour éviter des recalculs inutiles dans les vues de tickets, de messages et de dashboards.")
    rows = [
        ("RoleBasedRoute", "Protège les pages selon le token, le rôle et l'état actif du compte."),
        ("Sidebar", "Affiche uniquement les sections utiles au rôle connecté."),
        ("Tickets", "Liste, filtre, détaille, affecte et suit les demandes."),
        ("Messages", "Affiche les rooms, l'historique et les messages temps réel."),
        ("Meetings", "Gère les réunions de suivi et les réponses des invités."),
        ("Dashboards", "Présente les indicateurs opérationnels et analytiques."),
    ]
    add_table(doc, ["Composant / vue", "Rôle dans l'application"], rows, widths=[2600, 6760])
    add_caption(doc, "Tableau 8 : principaux éléments frontend réalisés.")

    add_title(doc, "Implémentation du backend et de la sécurité", 2)
    add_para(doc, "Le backend constitue le point de contrôle principal du système. Après la connexion, un token JWT est généré pour l'utilisateur actif. Pour chaque requête protégée, le middleware vérifie la présence du token, sa validité, son éventuelle révocation et l'état du compte en base. Les middlewares de rôle contrôlent ensuite si l'utilisateur possède l'autorisation nécessaire pour consulter ou modifier la ressource demandée.")
    add_para(doc, "Cette double protection est importante, car une interface peut masquer des boutons sans empêcher une requête manuelle. Dans AGCE CRM, les restrictions sont donc appliquées côté client pour guider l'utilisateur et côté serveur pour garantir la sécurité réelle. Les mots de passe sont hachés avec bcrypt ; les comptes inactifs sont refusés ; les tickets résolus sont verrouillés afin de préserver l'historique.")
    rows = [
        ("Authentification", "Connexion par username et mot de passe, vérification bcrypt, émission JWT."),
        ("Compte actif", "Refus de connexion et de requête si l'employé est inactif."),
        ("RBAC", "Contrôle des accès selon ADMIN, SD, Manager, PKI et IT."),
        ("Accès ticket", "PKI et IT ne consultent que les tickets relevant de leur périmètre."),
        ("Ticket résolu", "Blocage des commentaires, affectations et réouvertures après clôture."),
        ("Logout", "Invalidation du token par blacklist en mémoire pour la session serveur."),
    ]
    add_table(doc, ["Mécanisme", "Implémentation réalisée"], rows, widths=[2400, 6960])
    add_caption(doc, "Tableau 9 : mécanismes de sécurité applicative.")

    add_title(doc, "Modules métier réalisés", 2)
    add_para(doc, "Les modules développés couvrent le cycle de vie complet d'une demande : identification de l'organisation, création du contact, ouverture du ticket, affectation au service compétent, échange autour du problème, réunion de suivi, supervision et clôture.")
    rows = [
        ("Employés", "Création, modification, activation et désactivation des comptes internes."),
        ("Organisations et contacts", "Centralisation des clients institutionnels et des interlocuteurs."),
        ("Tickets", "Création en Pending, affectation IT/PKI, changement de statut et verrouillage final."),
        ("Commentaires", "Historique structuré des remarques associées à un ticket."),
        ("Messagerie", "Rooms Socket.IO liées aux tickets avec suivi des messages lus et non lus."),
        ("Meetings", "Planification des réunions et conservation des décisions de suivi."),
        ("E-mails clients", "Import IMAP des messages SD, filtrage par contact connu et pièces jointes."),
        ("Dashboards", "Indicateurs SD, manager, IT et PKI selon le périmètre autorisé."),
    ]
    add_table(doc, ["Module", "Fonction réalisée"], rows, widths=[2300, 7060])
    add_caption(doc, "Tableau 10 : synthèse des modules métier implémentés.")

    add_title(doc, "Illustrations de l'application réalisée", 2)
    add_para(doc, "Les captures suivantes présentent les principales interfaces obtenues : authentification, administration, supervision, création de ticket, suivi opérationnel, messagerie et dashboard analytique. Elles répondent à l'exigence du guide qui recommande d'accompagner l'implémentation par des preuves visuelles et des résultats observables.")
    for i, (path, caption) in enumerate(SCREENSHOTS, 6):
        if Path(path).exists():
            add_picture(doc, path, width=6.25, alt_text=caption)
            add_caption(doc, f"Figure {i} : {caption}.")

    add_title(doc, "Stratégie de tests", 2)
    add_para(doc, "La validation a été menée à partir de scénarios fonctionnels et de scénarios de sécurité. Les tests fonctionnels vérifient que les modules produisent les résultats attendus. Les tests de sécurité vérifient que les règles d'accès ne peuvent pas être contournées par une navigation directe ou par une requête API non autorisée.")
    rows = [
        ("Fonctionnel", "Parcours utilisateur complets : connexion, création, affectation, consultation, clôture.", "Confirmer que les besoins métier sont couverts."),
        ("Sécurité", "Accès interdits, compte inactif, token invalide, rôle non autorisé.", "Vérifier l'application effective du RBAC."),
        ("Données", "Contrôle des champs obligatoires, statuts autorisés, associations organisation-contact-ticket.", "Préserver la cohérence de la base."),
        ("Temps réel", "Connexion à une room, envoi, diffusion et lecture des messages.", "Valider Socket.IO et les compteurs non lus."),
        ("Régression manuelle", "Répétition des scénarios après modification d'un module sensible.", "S'assurer qu'une correction ne casse pas un autre parcours."),
    ]
    add_table(doc, ["Type de test", "Exemples", "Objectif"], rows, widths=[1900, 4550, 2910])
    add_caption(doc, "Tableau 11 : stratégie de validation appliquée.")

    add_title(doc, "Scénarios de validation", 2)
    rows = [
        ("TC01", "Connexion utilisateur actif", "Identifiants valides", "Accès au tableau de bord du rôle", "Validé"),
        ("TC02", "Compte inactif", "Identifiants corrects mais statut inactif", "Accès refusé", "Validé"),
        ("TC03", "Accès direct à une URL interdite", "Rôle non autorisé", "Page Access Denied ou redirection", "Validé"),
        ("TC04", "Création employé", "Profil ADMIN ou SD autorisé", "Employé créé avec service, statut et mot de passe haché", "Validé"),
        ("TC05", "Création organisation/contact", "Profil SD", "Données CRM enregistrées et rattachées", "Validé"),
        ("TC06", "Création ticket", "Profil SD et champs requis", "Ticket Pending, room et activité générés", "Validé"),
        ("TC07", "Affectation ticket", "Ticket non résolu", "Accès limité au service IT ou PKI affecté", "Validé"),
        ("TC08", "Changement de statut", "Statut autorisé", "Mise à jour et historique conservé", "Validé"),
        ("TC09", "Messagerie temps réel", "Room autorisée", "Message visible aux acteurs concernés", "Validé"),
        ("TC10", "Meeting de suivi", "Invité sélectionné", "Meeting créé avec statut Pending", "Validé"),
        ("TC11", "Manager lecture seule", "Profil Manager", "Consultation sans CRUD métier", "Validé"),
        ("TC12", "Ticket résolu", "Statut Resolved", "Nouvelles modifications bloquées", "Validé"),
        ("TC13", "E-mail client connu", "Expéditeur enregistré comme contact", "E-mail importé avec pièces jointes", "Validé"),
    ]
    add_table(doc, ["Code", "Scénario", "Condition", "Résultat attendu", "Statut"], rows, widths=[900, 2100, 2100, 3100, 1160])
    add_caption(doc, "Tableau 12 : scénarios de validation fonctionnelle et sécurité.")

    add_title(doc, "Résultats obtenus", 2)
    add_para(doc, "Les tests montrent que le prototype répond aux objectifs principaux définis au début du projet. Les demandes sont centralisées dans un circuit unique, les tickets sont reliés aux organisations et contacts, les échanges sont conservés et les rôles disposent de vues adaptées à leurs responsabilités.")
    rows = [
        ("Centralisation des demandes", "Atteint", "Tickets, contacts, organisations et e-mails sont regroupés dans le système."),
        ("Séparation des responsabilités", "Atteint", "Le RBAC est appliqué dans l'interface et dans l'API."),
        ("Traçabilité", "Partiellement à fortement atteint", "Commentaires, messages, meetings, affectations et activités conservent l'historique principal."),
        ("Collaboration interne", "Atteint", "Les rooms Socket.IO structurent les échanges autour des tickets."),
        ("Supervision manager", "Atteint", "Les dashboards fournissent une lecture globale sans actions de modification."),
        ("Sécurité applicative", "Atteint avec perspectives", "JWT, bcrypt, compte actif, rôles et verrouillage des tickets sont en place."),
        ("Industrialisation", "À poursuivre", "Tests automatisés, monitoring, sauvegardes et MFA restent à renforcer."),
    ]
    add_table(doc, ["Objectif", "Niveau d'atteinte", "Observation"], rows, widths=[2700, 2300, 4360])
    add_caption(doc, "Tableau 13 : synthèse des résultats par objectif.")

    add_title(doc, "Discussion des résultats", 2)
    add_para(doc, "La solution se distingue des outils génériques de ticketing par son adaptation au flux interne de l'AGCE. Les règles ne sont pas uniquement configurées dans l'interface ; elles sont codées dans les services backend et dans les middlewares. Cette approche correspond mieux à un contexte institutionnel où la souveraineté des données, la confidentialité et la séparation des responsabilités sont essentielles.")
    add_para(doc, "Le choix d'une architecture modulaire facilite également l'évolution du système. Un module peut être amélioré sans réécrire l'ensemble de l'application, à condition de conserver les contrats d'API et les règles de sécurité. Cette organisation constitue une base favorable pour ajouter ultérieurement des SLA, des notifications avancées, des rapports exportables ou des tests automatisés.")

    add_title(doc, "Limites du travail", 2)
    add_para(doc, "Malgré les résultats obtenus, certaines limites doivent être reconnues. La validation reste principalement manuelle ; elle doit être complétée par des tests unitaires, des tests d'intégration et des tests end-to-end. La blacklist JWT est conservée en mémoire, ce qui reste acceptable pour un prototype mais doit être remplacé par une stratégie persistante ou par des tokens courts en production.")
    add_para(doc, "L'évaluation quantitative des gains de temps n'a pas encore été conduite sur un volume réel de tickets. De plus, l'industrialisation devrait intégrer des sauvegardes planifiées, du monitoring, une journalisation plus détaillée, une authentification multifactorielle et une politique complète d'audit des actions utilisateurs.")

    add_title(doc, "Synthèse du chapitre", 2)
    add_para(doc, "Ce chapitre a présenté la réalisation du système AGCE CRM, depuis l'environnement technique jusqu'aux tests de validation. L'application obtenue concrétise la conception proposée : elle fournit une interface par rôle, une API sécurisée, une base relationnelle, une messagerie temps réel, une gestion des tickets et des dashboards de supervision. Les tests effectués confirment l'adéquation générale de la solution avec les besoins identifiés, tout en mettant en évidence les améliorations nécessaires avant un déploiement industriel complet.")


def conclusion(doc):
    add_title(doc, "Conclusion générale", 1)
    add_para(doc, "Ce mémoire a porté sur la conception et la réalisation d'une application web sécurisée de gestion des tickets et de la relation client pour l'AGCE. Le problème initial résidait dans la dispersion des demandes, la difficulté de suivi, l'insuffisance de traçabilité et la visibilité limitée sur l'activité de support.")
    add_para(doc, "La solution AGCE CRM proposée centralise les organisations, contacts, tickets, échanges, réunions et indicateurs. Elle met en œuvre une séparation claire des responsabilités entre administrateur, Service Delivery, manager, IT et PKI. Elle améliore ainsi la structuration du cycle de vie d'une demande depuis sa création jusqu'à sa résolution.")
    add_para(doc, "Sur le plan technique, le projet a permis de mettre en pratique une architecture web moderne reposant sur React, Express.js, MySQL et Socket.IO, avec des mécanismes de sécurité fondés sur JWT, bcrypt et le contrôle d'accès par rôle. Sur le plan organisationnel, il apporte une meilleure visibilité, une traçabilité renforcée et un cadre de collaboration plus cohérent.")
    add_title(doc, "Perspectives futures", 1)
    add_para(doc, "Plusieurs évolutions peuvent renforcer la robustesse et la maturité de la solution.")
    rows = [
        ("Court terme", "Tests automatisés, amélioration des validations, consolidation des messages d'erreur, documentation API."),
        ("Sécurité", "Authentification multifactorielle, rotation des tokens, stockage persistant de révocation, audit des accès."),
        ("Exploitation", "Sauvegardes automatisées, monitoring, journalisation avancée, alertes d'incidents."),
        ("Fonctionnel", "Notifications avancées, statistiques plus détaillées, filtres analytiques, export de rapports."),
        ("Long terme", "Assistance IA pour la qualification des tickets, application mobile, déploiement pilote en production."),
    ]
    add_table(doc, ["Horizon", "Améliorations proposées"], rows, widths=[1800, 7560])
    add_caption(doc, "Tableau 8 : perspectives d'évolution prioritaires.")


def bibliography(doc):
    add_title(doc, "Bibliographie", 1)
    refs = [
        "Journal officiel de la République algérienne démocratique et populaire, Loi n° 15-04 du 1er février 2015 fixant les règles générales relatives à la signature et à la certification électroniques.",
        "Journal officiel de la République algérienne démocratique et populaire, Décret exécutif n° 16-135 du 25 avril 2016 fixant l'organisation et le fonctionnement de l'AGCE.",
        "Journal officiel de la République algérienne démocratique et populaire, Loi n° 18-07 du 10 juin 2018 relative à la protection des personnes physiques dans le traitement des données à caractère personnel.",
        "Autorité Gouvernementale de Certification Électronique, Présentation de l'AGCE. Disponible : https://www.agce.dz/",
        "R. S. Sandhu, E. J. Coyne, H. L. Feinstein et C. E. Youman, « Role-Based Access Control Models », Computer, vol. 29, n° 2, pp. 38-47, 1996.",
        "OWASP Foundation, Authorization Cheat Sheet, OWASP Cheat Sheet Series.",
        "OWASP Foundation, Authentication Cheat Sheet, OWASP Cheat Sheet Series.",
        "Object Management Group, OMG Unified Modeling Language Specification, version 2.5.1, 2017.",
        "K. Beck et al., Manifesto for Agile Software Development, Agile Alliance, 2001.",
        "Meta Open Source, React Documentation.",
        "Express.js, Express - Node.js Web Application Framework Documentation.",
        "Socket.IO, Socket.IO Documentation, version 4.",
        "Oracle, MySQL 8.4 Reference Manual.",
        "International Organization for Standardization, ISO/IEC 25010:2023 - SQuaRE product quality model.",
        "NIST, Security and Privacy Controls for Information Systems and Organizations, SP 800-53 Rev. 5, 2020.",
        "OWASP Foundation, JSON Web Token Cheat Sheet.",
        "N. Provos et D. Mazières, « A Future-Adaptable Password Scheme », USENIX Annual Technical Conference, 1999.",
        "Atlassian, Jira Service Management Cloud Documentation.",
        "Zendesk, Zendesk Documentation and Administrator Guide.",
        "GLPI Project, GLPI Documentation and IT Service Desk Features.",
        "osTicket, osTicket Documentation.",
        "Vite Team, Vite Documentation.",
        "M. Crispin, Internet Message Access Protocol - Version 4rev1, RFC 3501, IETF, 2003.",
        "Nodemailer Project, Mailparser Documentation.",
        "P. A. Grassi et al., Digital Identity Guidelines: Authentication and Lifecycle Management, NIST SP 800-63B, 2017.",
        "F. Buttle et S. Maklan, Customer Relationship Management: Concepts and Technologies, 4e éd., Routledge, 2019.",
        "W. Reinartz, M. Krafft et W. D. Hoyer, « The Customer Relationship Management Process: Its Measurement and Impact on Performance », Journal of Marketing Research, vol. 41, n° 3, pp. 293-305, 2004.",
        "AXELOS, ITIL Foundation: ITIL 4 Edition, The Stationery Office, 2019.",
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(ref)
        set_run_font(r, size=10.2)


def annexes(doc):
    add_title(doc, "Annexes", 1)
    add_title(doc, "Annexe A - Checklist de conformité scientifique", 2)
    rows = [
        ("Page de garde complète", "Oui"),
        ("Résumé en français, arabe et anglais", "Oui"),
        ("Problématique et question de recherche explicites", "Oui"),
        ("État de l'art avec comparaison", "Oui"),
        ("Méthodologie de recherche", "Oui"),
        ("Diagrammes UML et modèle logique", "Oui"),
        ("Captures et résultats de réalisation", "Oui"),
        ("Tests et limites", "Oui"),
        ("Conclusion et perspectives", "Oui"),
        ("Bibliographie enrichie", "Oui"),
    ]
    add_table(doc, ["Élément vérifié", "Statut"], rows, widths=[6500, 2860])
    add_caption(doc, "Tableau 9 : checklist finale de conformité au guide.")


def main():
    diagrams = make_diagrams()
    doc = Document()
    configure_document(doc)
    cover(doc)
    front_matter(doc)
    chapter_intro(doc)
    chapter_one(doc)
    chapter_two(doc, diagrams)
    chapter_three(doc)
    conclusion(doc)
    bibliography(doc)
    annexes(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
