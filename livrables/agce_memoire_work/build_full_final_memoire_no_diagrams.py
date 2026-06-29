from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/issamouladsmane/Desktop/full_stuck_ticket_gestion")
OUT_DIR = ROOT / "livrables"
WORK_DIR = OUT_DIR / "agce_memoire_work"
DOCX_OUT = OUT_DIR / "Memoire_AGCE_CRM_Complet_Final_Sans_Diagrammes.docx"
MD_OUT = OUT_DIR / "Memoire_AGCE_CRM_Complet_Final_Sans_Diagrammes.md"

BLUE = RGBColor(31, 78, 121)
DARK_BLUE = RGBColor(22, 47, 75)
MUTED = RGBColor(95, 105, 115)
GRAY_FILL = "F2F4F7"
LIGHT_BLUE_FILL = "E8EEF5"
PLACEHOLDER_FILL = "F7F7F7"


md_lines: list[str] = []


def set_run_font(run, size=None, bold=None, italic=None, color=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_bidi(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def style_table(table, widths=None, header=True):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    if widths:
        grid = table._tbl.tblGrid
        if grid is None:
            grid = OxmlElement("w:tblGrid")
            table._tbl.insert(0, grid)
        for child in list(grid):
            grid.remove(child)
        for w in widths:
            col = OxmlElement("w:gridCol")
            col.set(qn("w:w"), str(w))
            grid.append(col)
        for row in table.rows:
            for idx, w in enumerate(widths):
                if idx < len(row.cells):
                    set_cell_width(row.cells[idx], w)
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            for par in cell.paragraphs:
                par.paragraph_format.space_before = Pt(0)
                par.paragraph_format.space_after = Pt(2)
                for run in par.runs:
                    set_run_font(run, size=9.4)
        if header and r_idx == 0:
            for cell in row.cells:
                set_cell_shading(cell, GRAY_FILL)
                for par in cell.paragraphs:
                    for run in par.runs:
                        set_run_font(run, size=9.4, bold=True, color=DARK_BLUE)


def configure_document(doc: Document):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)


def add_header_footer(doc: Document):
    for section in doc.sections:
        section.different_first_page_header_footer = True
        section.header.is_linked_to_previous = False
        p = section.header.paragraphs[0]
        p.text = "AGCE CRM - Mémoire de fin d'études"
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in p.runs:
            set_run_font(run, size=8.5, color=MUTED)
        f = section.footer.paragraphs[0]
        f.alignment = WD_ALIGN_PARAGRAPH.CENTER
        f.text = "Licence Informatique - Option ISIL - 2025/2026"
        for run in f.runs:
            set_run_font(run, size=8.5, color=MUTED)


def page_break(doc):
    doc.add_page_break()
    md_lines.append("\n---\n")


def h(doc, text: str, level=1):
    doc.add_heading(text, level=level)
    md_lines.append(f"\n{'#' * level} {text}\n")


def p(doc, text: str, bold_prefix: str | None = None, rtl=False):
    par = doc.add_paragraph()
    if rtl:
        set_paragraph_bidi(par)
        par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if bold_prefix and text.startswith(bold_prefix):
        r1 = par.add_run(bold_prefix)
        set_run_font(r1, bold=True, size=11)
        r2 = par.add_run(text[len(bold_prefix):])
        set_run_font(r2, size=11)
    else:
        r = par.add_run(text)
        set_run_font(r, size=11, name="Arial" if rtl else "Calibri")
    md_lines.append(text + "\n")


def bullets(doc, items):
    for item in items:
        par = doc.add_paragraph(style="List Bullet")
        r = par.add_run(item)
        set_run_font(r, size=10.8)
        md_lines.append(f"- {item}\n")


def numbered(doc, items):
    for item in items:
        par = doc.add_paragraph(style="List Number")
        r = par.add_run(item)
        set_run_font(r, size=10.8)
        md_lines.append(f"1. {item}\n")


def caption(doc, text):
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = par.add_run(text)
    set_run_font(r, size=9.4, italic=True, color=MUTED)
    md_lines.append(f"_{text}_\n")


def table(doc, headers, rows, widths=None, title=None):
    if title:
        caption(doc, title)
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, head in enumerate(headers):
        t.rows[0].cells[i].text = str(head)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    style_table(t, widths=widths)
    md_lines.append("\n| " + " | ".join(headers) + " |\n")
    md_lines.append("|" + "|".join(["---"] * len(headers)) + "|\n")
    for row in rows:
        md_lines.append("| " + " | ".join(str(v).replace("\n", "<br>") for v in row) + " |\n")
    md_lines.append("\n")
    return t


def placeholder(doc, title: str, explanation: str):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.cell(0, 0)
    set_cell_shading(cell, PLACEHOLDER_FILL)
    set_cell_margins(cell, top=160, bottom=160, start=180, end=180)
    par = cell.paragraphs[0]
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = par.add_run(f"[{title}]")
    set_run_font(r, size=11.5, bold=True, color=DARK_BLUE)
    par2 = cell.add_paragraph()
    par2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = par2.add_run(explanation)
    set_run_font(r2, size=9.8, italic=True, color=MUTED)
    md_lines.append(f"\n[{title}]\n{explanation}\n\n")


def code_block(doc, text: str):
    par = doc.add_paragraph()
    for line in text.splitlines():
        r = par.add_run(line + "\n")
        set_run_font(r, size=9.2, name="Courier New")
    md_lines.append("```text\n" + text + "\n```\n")


def cover(doc):
    for line in [
        "République Algérienne Démocratique et Populaire",
        "Ministère de l'Enseignement Supérieur et de la Recherche Scientifique",
        "Université Saad Dahleb - Blida 1",
        "Faculté des Sciences",
        "Département d'Informatique",
    ]:
        par = doc.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = par.add_run(line)
        set_run_font(r, size=12, bold=True if "Université" in line else False)
        md_lines.append(line + "\n")

    doc.add_paragraph()
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = par.add_run("Mémoire de fin d'études")
    set_run_font(r, size=16, bold=True, color=BLUE)
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = par.add_run("Présenté en vue de l'obtention du diplôme de Licence en Informatique")
    set_run_font(r, size=12)
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = par.add_run("Option : Systèmes d'Information et Ingénierie Logicielle, ISIL")
    set_run_font(r, size=12)

    doc.add_paragraph()
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    c = t.cell(0, 0)
    set_cell_shading(c, LIGHT_BLUE_FILL)
    set_cell_margins(c, top=220, bottom=220, start=180, end=180)
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c.paragraphs[0].add_run("Conception et réalisation d'une application de gestion des incidents et des tickets des services de l'AGCE")
    set_run_font(r, size=16, bold=True, color=DARK_BLUE)
    doc.add_paragraph()

    rows = [
        ("Préparé par", "Mr. Laceb Karim\nMr. Ouladsmane Issam"),
        ("Encadrants", "Mrs. Arkam Meriem\nMr. Lekhchine Sami"),
        ("Organisme d'accueil", "Autorité Gouvernementale de Certification Électronique, AGCE"),
        ("Année universitaire", "2025/2026"),
    ]
    table(doc, ["Élément", "Information"], rows, widths=[2600, 6760])
    page_break(doc)


def preliminaries(doc):
    h(doc, "Dédicace de Mr. Laceb Karim", 1)
    p(doc, "À ses chers parents, pour leur soutien, leurs sacrifices et leur confiance constante. À sa famille, pour sa présence et ses encouragements durant le parcours universitaire. À ses enseignants, pour leur accompagnement et pour les connaissances transmises. Ce travail est dédié à toutes les personnes qui ont contribué, par leur aide ou leurs conseils, à l'aboutissement de ce projet.")
    page_break(doc)

    h(doc, "Dédicace de Mr. Ouladsmane Issam", 1)
    p(doc, "À ses chers parents, en témoignage de gratitude pour leur patience, leurs sacrifices et leur soutien moral. À sa famille et à ses proches, pour leurs encouragements dans les moments importants de ce parcours. Ce mémoire est dédié, avec respect et reconnaissance, à toutes les personnes qui ont cru en ce travail et qui ont contribué, directement ou indirectement, à sa réalisation.")
    page_break(doc)

    h(doc, "Remerciements", 1)
    p(doc, "Nous remercions d'abord Dieu, qui nous a accordé la force, la patience et la volonté nécessaires pour mener à bien ce travail.")
    p(doc, "Nous exprimons notre profonde reconnaissance à nos familles pour leur soutien moral, leur confiance et leurs encouragements constants tout au long de notre formation.")
    p(doc, "Nous adressons nos sincères remerciements à notre encadrante universitaire, Mrs. Arkam Meriem, pour son accompagnement, ses orientations méthodologiques, sa disponibilité et ses remarques constructives.")
    p(doc, "Nous remercions également notre encadrant professionnel, Mr. Lekhchine Sami, pour son accueil au sein de l'Autorité Gouvernementale de Certification Électronique, pour ses conseils techniques et pour les informations mises à notre disposition.")
    p(doc, "Nos remerciements s'adressent au personnel de l'AGCE pour son aide et sa collaboration, ainsi qu'à l'ensemble des enseignants du Département d'Informatique pour la formation reçue. Nous remercions enfin les membres du jury pour l'honneur qu'ils nous font en évaluant ce mémoire.")
    page_break(doc)

    h(doc, "Résumé", 1)
    p(doc, "Ce mémoire présente la conception et la réalisation d'une application web sécurisée de gestion des incidents et des tickets des services de l'Autorité Gouvernementale de Certification Électronique. Le projet répond à un besoin de centralisation des demandes clients, des organisations, des contacts, des tickets techniques, des messages internes, des réunions de suivi et des indicateurs de supervision.")
    p(doc, "La solution proposée, appelée AGCE CRM, repose sur une architecture client-serveur. Le frontend est développé avec React et Vite, tandis que le backend utilise Node.js, Express.js, MySQL, JSON Web Token, bcrypt et Socket.IO. L'application intègre également un service d'import des e-mails via Gmail IMAP afin de faciliter le suivi des demandes reçues par le Service Delivery.")
    p(doc, "Les résultats obtenus montrent que le système permet d'organiser le cycle de vie d'une demande depuis sa réception jusqu'à sa résolution, tout en appliquant une séparation stricte des rôles, une traçabilité des actions et un contrôle d'accès fiable. Les perspectives concernent notamment l'authentification multifactorielle, les tests automatisés, le monitoring, les sauvegardes et l'amélioration des tableaux de bord.")
    p(doc, "Mots-clés : AGCE, CRM, gestion de tickets, RBAC, React, Express.js, MySQL, Socket.IO, sécurité applicative.")
    page_break(doc)

    h(doc, "Abstract", 1)
    p(doc, "This thesis presents the design and implementation of a secure web application for incident and ticket management within the Governmental Electronic Certification Authority. The project addresses the need to centralize client requests, organizations, contacts, technical tickets, internal messages, follow-up meetings and supervision indicators.")
    p(doc, "The proposed solution, named AGCE CRM, is based on a client-server architecture. The frontend is developed with React and Vite, while the backend relies on Node.js, Express.js, MySQL, JSON Web Token, bcrypt and Socket.IO. The application also includes an email import service through Gmail IMAP to support the monitoring of requests received by the Service Delivery team.")
    p(doc, "The achieved results show that the system supports the full lifecycle of a request, from reception to resolution, while enforcing strict role separation, action traceability and reliable access control. Future improvements include multi-factor authentication, automated testing, monitoring, backups and enhanced dashboards.")
    p(doc, "Keywords: AGCE, CRM, ticket management, RBAC, React, Express.js, MySQL, Socket.IO, application security.")
    page_break(doc)

    h(doc, "ملخص", 1)
    p(doc, "يعرض هذا العمل تصميم وإنجاز تطبيق ويب آمن لإدارة الحوادث والتذاكر على مستوى سلطة التصديق الإلكتروني الحكومية. يستجيب المشروع لحاجة عملية تتمثل في مركزية طلبات الزبائن، والمنظمات، وجهات الاتصال، والتذاكر التقنية، والرسائل الداخلية، واجتماعات المتابعة، ومؤشرات الإشراف.", rtl=True)
    p(doc, "تعتمد الحلول المقترحة، المسماة AGCE CRM، على بنية عميل/خادم. تم تطوير الواجهة الأمامية باستعمال React و Vite، بينما يعتمد الخادم على Node.js و Express.js و MySQL و JSON Web Token و bcrypt و Socket.IO. كما يدمج النظام خدمة لاستيراد الرسائل الإلكترونية عبر Gmail IMAP من أجل دعم متابعة الطلبات الواردة إلى مصلحة Service Delivery.", rtl=True)
    p(doc, "تظهر النتائج أن النظام يسمح بتنظيم دورة حياة الطلب منذ استقباله إلى غاية حله، مع ضمان فصل واضح بين الأدوار، وتتبع العمليات، ومراقبة موثوقة للصلاحيات. وتشمل الآفاق المستقبلية المصادقة متعددة العوامل، والاختبارات الآلية، والمراقبة، والنسخ الاحتياطي، وتحسين لوحات القيادة.", rtl=True)
    p(doc, "الكلمات المفتاحية: AGCE، CRM، إدارة التذاكر، RBAC، React، Express.js، MySQL، Socket.IO، أمن التطبيقات.", rtl=True)
    page_break(doc)


def navigation_pages(doc):
    h(doc, "Table des matières", 1)
    toc = [
        "Introduction générale",
        "Chapitre 1 : Contexte général, organisme d'accueil et problématique",
        "1.1 Introduction du chapitre",
        "1.2 Présentation de l'AGCE",
        "1.3 Historique et cadre général",
        "1.4 Services concernés par le projet",
        "1.5 Contexte du projet",
        "1.6 Problématique",
        "1.7 Objectifs du projet",
        "1.8 Enjeux du projet",
        "1.9 Conclusion du chapitre",
        "Chapitre 2 : État de l'art et étude de l'existant",
        "Chapitre 3 : Analyse et conception",
        "Chapitre 4 : Réalisation, implémentation et tests",
        "Conclusion générale et perspectives",
        "Bibliographie",
        "Annexes",
    ]
    for item in toc:
        p(doc, item)
    page_break(doc)

    h(doc, "Liste des figures", 1)
    for item in [
        "Figure 1 : Organigramme du département concerné au sein de l'AGCE",
        "Figure 2 : Emplacement réservé au diagramme de cas d'utilisation",
        "Figure 3 : Emplacement réservé au diagramme de classes",
        "Figure 4 : Emplacement réservé au diagramme de séquence d'authentification",
        "Figure 5 : Emplacement réservé au diagramme de séquence de création d'un ticket",
        "Figure 6 : Aperçu de l'interface de l'application",
    ]:
        p(doc, item)
    page_break(doc)

    h(doc, "Liste des tableaux", 1)
    for item in [
        "Tableau 1 : Comparaison des solutions existantes",
        "Tableau 2 : Synthèse de la solution proposée",
        "Tableau 3 : Acteurs et permissions",
        "Tableau 4 : Besoins fonctionnels",
        "Tableau 5 : Besoins non fonctionnels",
        "Tableau 6 : Tables principales de la base de données",
        "Tableau 7 : Environnement logiciel",
        "Tableau 8 : Tests fonctionnels",
        "Tableau 9 : Tests de sécurité",
    ]:
        p(doc, item)
    page_break(doc)

    h(doc, "Liste des abréviations", 1)
    table(doc, ["Abréviation", "Signification"], [
        ("AGCE", "Autorité Gouvernementale de Certification Électronique"),
        ("CRM", "Customer Relationship Management"),
        ("CRL", "Certificate Revocation List"),
        ("ISIL", "Ingénierie des Systèmes d'Information et Logiciels"),
        ("JWT", "JSON Web Token"),
        ("LRAO", "Local Registration Authority Officer"),
        ("MFA", "Multi-Factor Authentication"),
        ("OCSP", "Online Certificate Status Protocol"),
        ("PFE", "Projet de Fin d'Études"),
        ("PKI", "Public Key Infrastructure"),
        ("RBAC", "Role-Based Access Control"),
        ("REST", "Representational State Transfer"),
        ("SD", "Service Delivery"),
        ("SSO", "Single Sign-On"),
        ("TLS/SSL", "Transport Layer Security / Secure Sockets Layer"),
        ("VPN", "Virtual Private Network"),
    ], widths=[2100, 7260])
    page_break(doc)


def introduction_generale(doc):
    h(doc, "Introduction générale", 1)
    p(doc, "La transformation numérique occupe aujourd'hui une place importante dans l'organisation des institutions publiques et privées. Elle ne se limite pas à la dématérialisation des documents ; elle implique aussi la mise en place de systèmes d'information capables de structurer les échanges, d'assurer le suivi des activités et de protéger les données traitées. Dans ce contexte, les applications de gestion interne deviennent des outils essentiels pour améliorer la qualité de service, la traçabilité et la prise de décision.")
    p(doc, "Les systèmes d'information jouent un rôle central dans cette évolution. Ils permettent de relier les acteurs, les données et les processus métier au sein d'une même plateforme. Lorsqu'ils sont correctement conçus, ils réduisent les pertes d'information, facilitent le contrôle des opérations et offrent une meilleure visibilité sur l'état réel de l'activité. Cette importance devient encore plus marquée dans les organismes qui manipulent des données sensibles ou qui interviennent dans la sécurité des échanges numériques.")
    p(doc, "La confiance numérique constitue l'un des fondements de l'administration électronique. Elle repose sur des mécanismes d'identification, d'authentification, de signature électronique, de certification et de protection de l'intégrité des documents numériques. En Algérie, ce domaine est notamment encadré par la loi n° 15-04 relative à la signature et à la certification électroniques. L'Autorité Gouvernementale de Certification Électronique, créée par le décret exécutif n° 16-135 du 25 avril 2016, intervient dans ce cadre en tant qu'organisme public chargé de la certification électronique au niveau gouvernemental.")
    p(doc, "L'AGCE occupe ainsi une position sensible dans l'écosystème numérique national. Elle participe à la sécurisation des échanges électroniques, à la gestion des certificats numériques, au fonctionnement de l'infrastructure à clés publiques et à l'accompagnement des institutions dans l'usage des services de confiance. Dans un tel environnement, la gestion des incidents, des demandes clients et des échanges internes doit être organisée de manière rigoureuse.")
    p(doc, "Le problème observé concerne la dispersion des informations entre plusieurs canaux : e-mails, discussions informelles, fichiers, appels, notes séparées et échanges internes non centralisés. Cette dispersion peut rendre difficile la reconstitution de l'historique d'un incident, retarder l'affectation d'une demande, limiter la visibilité du manager et créer une confusion dans les responsabilités. Le besoin dépasse donc la simple création d'un formulaire de ticket ; il s'agit de concevoir un système complet couvrant le cycle de vie d'une demande depuis sa réception jusqu'à sa résolution.")
    p(doc, "La problématique centrale de ce mémoire peut être formulée ainsi : comment concevoir et réaliser une application web sécurisée permettant à l'AGCE de centraliser la gestion des organisations, des contacts, des tickets, des messages, des réunions et des indicateurs de supervision, tout en garantissant une séparation stricte des rôles, une traçabilité des actions et un contrôle d'accès fiable ?")
    p(doc, "Pour répondre à cette problématique, le projet vise à concevoir et réaliser une application web sécurisée de type CRM/ticketing appelée AGCE CRM. Cette solution doit permettre de centraliser les organisations clientes, gérer les contacts, collecter les demandes entrantes, créer et suivre les tickets, affecter les incidents aux services IT ou PKI, conserver l'historique des affectations, organiser les échanges internes, gérer les réunions de suivi et fournir des tableaux de bord analytiques.")
    p(doc, "La méthodologie adoptée repose sur une démarche progressive : étude du contexte, analyse de l'existant, collecte des besoins, analyse fonctionnelle, conception, prototypage, développement frontend, développement backend, intégration, tests et validation. Cette démarche s'inspire d'une approche Agile, avec un découpage en tâches permettant d'améliorer progressivement la solution et de vérifier régulièrement la cohérence entre les besoins identifiés et les fonctionnalités réalisées.")
    p(doc, "Le mémoire est organisé en quatre chapitres. Le premier chapitre présente le contexte général, l'organisme d'accueil et la problématique. Le deuxième chapitre étudie l'état de l'art et compare les solutions existantes. Le troisième chapitre décrit l'analyse et la conception du système. Le quatrième chapitre présente la réalisation, l'implémentation et les tests. Enfin, la conclusion générale synthétise les apports du projet, ses limites et ses perspectives d'évolution.")
    page_break(doc)


def chapter_one(doc):
    h(doc, "Chapitre 1 : Contexte général, organisme d'accueil et problématique", 1)
    h(doc, "1.1 Introduction du chapitre", 2)
    p(doc, "Ce chapitre présente le cadre général du projet, l'organisme d'accueil et le problème auquel la solution proposée cherche à répondre. Il met en évidence le rôle de l'AGCE dans la confiance numérique, les services concernés par le projet et les enjeux liés à la gestion centralisée des incidents et des tickets.")
    h(doc, "1.2 Présentation de l'AGCE", 2)
    p(doc, "L'Autorité Gouvernementale de Certification Électronique est un organisme public algérien chargé de la certification électronique et de la signature électronique dans le secteur gouvernemental. Elle participe à la mise en place des services de confiance numérique nécessaires à l'authentification, à la sécurisation des échanges électroniques et à la protection de l'intégrité des documents numériques.")
    p(doc, "Dans le cadre de ses missions, l'AGCE intervient autour de l'infrastructure à clés publiques, appelée PKI. Cette infrastructure permet notamment de gérer les certificats électroniques, leur cycle de vie, leur publication, leur révocation et leur vérification. Elle constitue un élément technique essentiel pour établir la confiance dans les échanges numériques.")
    h(doc, "1.3 Historique et cadre général", 2)
    p(doc, "Le développement mondial des services numériques a renforcé le besoin de mécanismes fiables de confiance électronique. La signature électronique, la certification numérique, l'horodatage et la vérification des certificats permettent de sécuriser les transactions et de donner une valeur juridique et technique aux échanges dématérialisés.")
    p(doc, "En Algérie, la loi n° 15-04 fixe le cadre général relatif à la signature et à la certification électroniques. L'AGCE a été créée par le décret exécutif n° 16-135 du 25 avril 2016 afin d'assurer, dans le secteur gouvernemental, les missions liées à la certification électronique. Ce contexte montre que les données manipulées par l'organisme ont une importance particulière et doivent être traitées dans un environnement maîtrisé.")
    p(doc, "La mise en place d'une infrastructure nationale de confiance numérique s'inscrit dans une dynamique de modernisation des institutions publiques. Elle nécessite des outils internes capables d'organiser les demandes, de suivre les incidents et de conserver l'historique des opérations.")
    h(doc, "1.4 Services concernés par le projet", 2)
    p(doc, "Le projet concerne principalement la Direction des infrastructures de gestion de clés. Deux sous-directions sont particulièrement impliquées : la Sous-direction de l'enregistrement, notamment le Service Delivery, et la Sous-direction d'exploitation des infrastructures de gestion de clés, notamment le service PKI.")
    table(doc, ["Service / rôle", "Responsabilités principales"], [
        ("Service Delivery", "Réception, validation et enregistrement des demandes ; gestion des organisations et contacts ; création et affectation des tickets ; suivi des demandes ; support et assistance technique ; réception des demandes de révocation."),
        ("PKI", "Génération des certificats électroniques ; gestion du cycle de vie ; publication des certificats et listes de révocation ; vérification de validité ; horodatage ; traitement des demandes techniques liées à la PKI."),
        ("IT", "Traitement des incidents informatiques, problèmes applicatifs ou techniques et support interne."),
        ("Manager", "Supervision de l'activité, consultation des tableaux de bord, analyse des indicateurs et suivi global en lecture seule."),
        ("Administrateur", "Gestion des comptes employés, création, modification, désactivation et accès au tableau de bord administrateur."),
    ], widths=[2300, 7060], title="Tableau : services concernés par le projet.")
    h(doc, "1.5 Contexte du projet", 2)
    p(doc, "Dans un organisme comme l'AGCE, la gestion des tickets et de la relation client ne peut pas être limitée à un simple enregistrement de demandes. Les incidents peuvent concerner des certificats, des organisations, des contacts, des applications ou des opérations techniques liées à la PKI. Chaque demande doit être suivie, orientée vers le service compétent et documentée.")
    p(doc, "La dispersion des informations entre les e-mails, les appels, les discussions informelles et les fichiers séparés limite la capacité à suivre correctement les incidents. Elle peut également compliquer la mesure du temps de traitement et réduire la visibilité du manager sur l'activité réelle des services.")
    h(doc, "1.6 Problématique", 2)
    p(doc, "La problématique centrale est la suivante : comment concevoir et réaliser une application web sécurisée permettant à l'AGCE de centraliser la gestion des organisations, des contacts, des tickets, des messages, des réunions et des indicateurs de supervision, tout en garantissant une séparation stricte des rôles, une traçabilité des actions et un contrôle d'accès fiable ?")
    h(doc, "1.7 Objectifs du projet", 2)
    p(doc, "L'objectif principal est de concevoir et réaliser une application web sécurisée de type CRM/ticketing permettant de centraliser la gestion des incidents, des demandes clients et des échanges internes au sein de l'AGCE.")
    bullets(doc, [
        "centraliser les organisations clientes et leurs contacts ;",
        "collecter les demandes entrantes et créer des tickets structurés ;",
        "affecter les tickets aux services IT ou PKI selon leur nature ;",
        "assurer la traçabilité des affectations, commentaires, messages et réunions ;",
        "permettre au Manager une supervision en lecture seule ;",
        "sécuriser les accès selon les rôles et l'état actif des comptes ;",
        "protéger les données sensibles et respecter les exigences de souveraineté ;",
        "améliorer l'organisation du support et la visibilité analytique.",
    ])
    h(doc, "1.8 Enjeux du projet", 2)
    p(doc, "Les enjeux du projet sont fonctionnels, organisationnels et sécuritaires. La centralisation permet de regrouper les informations liées aux demandes. La traçabilité permet de reconstituer l'historique d'un incident. La séparation des rôles limite les actions aux responsabilités de chaque acteur. La visibilité analytique aide le manager à suivre l'activité. Enfin, la souveraineté des données justifie la mise en place d'une solution locale, sécurisée et adaptée au contexte de l'AGCE.")
    h(doc, "1.9 Conclusion du chapitre", 2)
    p(doc, "Ce chapitre a présenté le contexte de la confiance numérique, l'organisme d'accueil, les services concernés et la problématique du projet. Il a montré que la solution attendue doit dépasser le simple enregistrement de tickets pour proposer un système complet de suivi, de sécurité, de traçabilité et de supervision.")
    page_break(doc)


def chapter_two(doc):
    h(doc, "Chapitre 2 : État de l'art et étude de l'existant", 1)
    h(doc, "2.1 Introduction du chapitre", 2)
    p(doc, "Ce chapitre présente les notions liées aux CRM, à la gestion des tickets, à l'ITSM, au contrôle d'accès et à la souveraineté des données. Il compare ensuite plusieurs solutions existantes afin de justifier la réalisation d'une solution personnalisée pour l'AGCE.")
    h(doc, "2.2 Définition d'un CRM", 2)
    p(doc, "Un CRM, ou Customer Relationship Management, est un système destiné à organiser les informations relatives aux clients, aux contacts, aux interactions et aux demandes. Il permet de centraliser les données, de suivre l'historique des échanges et d'améliorer la qualité de la relation avec les utilisateurs ou les organismes partenaires.")
    p(doc, "Dans le contexte de l'AGCE, la notion de CRM est adaptée à un environnement institutionnel. Les clients peuvent être des organisations ou des interlocuteurs liés à des demandes de certification, de support ou d'assistance technique.")
    h(doc, "2.3 Gestion des tickets et ITSM", 2)
    p(doc, "La gestion des tickets consiste à transformer une demande ou un incident en objet suivi. Le cycle de vie d'un ticket comprend généralement la création, la qualification, la priorisation, l'affectation, le traitement, le suivi, la résolution et la clôture. Cette logique permet d'éviter que les demandes restent dispersées ou non suivies.")
    p(doc, "L'ITSM, ou gestion des services informatiques, propose une approche structurée de la fourniture de services. Dans le cadre du projet, cette logique se traduit par la définition de statuts, de rôles, d'affectations et d'indicateurs de supervision.")
    h(doc, "2.4 Sécurité, RBAC et souveraineté des données", 2)
    p(doc, "Le RBAC, ou contrôle d'accès basé sur les rôles, consiste à attribuer des permissions selon la fonction de l'utilisateur. Cette approche convient au projet, car les responsabilités de l'administrateur, du Service Delivery, du Manager, de PKI et de IT sont différentes.")
    p(doc, "La souveraineté des données constitue également un enjeu important. Une solution utilisée par une autorité de certification peut contenir des informations sur des organisations, des contacts, des incidents techniques, des échanges internes et des pièces jointes. Il est donc nécessaire de privilégier une solution dont les données restent sous contrôle de l'organisme.")
    h(doc, "2.5 Solutions existantes", 2)
    p(doc, "Plusieurs solutions de ticketing ou de support existent sur le marché. Jira Service Management et Zendesk proposent des plateformes puissantes, souvent orientées SaaS. GLPI et osTicket sont des solutions open source pouvant être déployées localement. Ces outils constituent des références utiles, mais ils ne répondent pas directement à toutes les contraintes métier et sécuritaires de l'AGCE.")
    h(doc, "2.6 Tableau comparatif des solutions existantes", 2)
    table(doc, ["Critère", "Solutions SaaS, Jira/Zendesk", "Solutions Open Source On-Premise, GLPI/osTicket", "Solution AGCE CRM proposée"], [
        ("Hébergement", "Cloud ou infrastructure du fournisseur.", "Déploiement local possible.", "Déploiement local adapté au contexte AGCE."),
        ("Souveraineté des données", "Dépend du fournisseur et de la localisation des données.", "Meilleur contrôle si l'installation est interne.", "Données conservées dans une base MySQL contrôlée par l'organisme."),
        ("Adaptation au contexte AGCE", "Nécessite une configuration avancée et parfois coûteuse.", "Adaptation possible, mais générique.", "Conçue autour des rôles SD, PKI, IT, Manager et ADMIN."),
        ("Personnalisation des rôles", "Possible, mais dépend des offres et modules.", "Possible avec configuration.", "RBAC intégré dans le frontend et le backend."),
        ("Intégration des tickets", "Gestion complète mais standardisée.", "Gestion standard des incidents.", "Tickets liés aux organisations, contacts, rooms, meetings et e-mails."),
        ("Messagerie temps réel", "Disponible selon modules ou intégrations.", "Souvent limitée ou externe.", "Rooms Socket.IO associées aux tickets."),
        ("Meetings", "Peut nécessiter une intégration externe.", "Fonction rarement centrale.", "Module de réunions intégré au suivi des tickets."),
        ("Tableaux de bord", "Riches mais génériques.", "Présents selon configuration.", "Dashboards personnalisés pour SD et Manager."),
        ("Contrôle RBAC", "Présent, mais dépend des plans et configurations.", "Présent, généralement générique.", "Contrôle d'accès spécifique aux responsabilités AGCE."),
        ("Coût et maîtrise technique", "Abonnement et dépendance fournisseur.", "Maîtrise plus forte mais adaptation nécessaire.", "Maîtrise du code, de la base et des règles métier."),
    ], widths=[1700, 2550, 2550, 2560], title="Tableau 1 : Comparaison des solutions existantes.")
    h(doc, "2.7 Limites des solutions existantes", 2)
    p(doc, "Les solutions SaaS comme Jira Service Management et Zendesk sont fonctionnellement riches, mais elles peuvent poser des questions de souveraineté, de confidentialité et de dépendance au fournisseur. Les solutions open source comme GLPI et osTicket offrent davantage de maîtrise locale, mais elles restent génériques et nécessitent une adaptation importante pour intégrer précisément le contexte PKI, les affectations IT/PKI, les meetings, les rooms de messagerie, l'inbox e-mail et les dashboards personnalisés.")
    h(doc, "2.8 Justification d'une solution personnalisée", 2)
    p(doc, "La solution AGCE CRM est justifiée par la nécessité de disposer d'un système local, sécurisé et aligné sur les processus internes. Elle permet de définir des règles métier spécifiques, de contrôler les accès selon les rôles, de conserver les données localement et d'intégrer les modules utiles dans une seule application cohérente.")
    table(doc, ["Dimension", "Apport de la solution proposée"], [
        ("Métier", "Adaptation aux rôles Service Delivery, PKI, IT, Manager et ADMIN."),
        ("Sécurité", "JWT, bcrypt, RBAC, vérification du compte actif et refus 403 côté backend."),
        ("Traçabilité", "Historique des affectations, commentaires, messages, meetings et activités."),
        ("Centralisation", "Organisations, contacts, tickets, e-mails et indicateurs regroupés."),
        ("Souveraineté", "Base MySQL locale et paramètres sensibles placés dans l'environnement."),
    ], widths=[2500, 6860], title="Tableau 2 : Synthèse de la solution proposée.")
    h(doc, "2.9 Conclusion du chapitre", 2)
    p(doc, "L'étude de l'existant montre que les solutions disponibles peuvent couvrir de nombreux besoins généraux, mais qu'elles ne répondent pas directement à toutes les contraintes de l'AGCE. Une solution personnalisée permet d'intégrer les règles métier, la souveraineté des données et le contrôle d'accès dans une architecture cohérente.")
    page_break(doc)


def chapter_three(doc):
    h(doc, "Chapitre 3 : Analyse et conception", 1)
    h(doc, "3.1 Introduction du chapitre", 2)
    p(doc, "Ce chapitre présente l'analyse des besoins, les acteurs, les permissions, les cas d'utilisation, les entités principales, l'architecture logique et les contrats d'API. Il décrit la conception sans générer de diagrammes visuels, conformément au choix de laisser des emplacements réservés pour les diagrammes définitifs.")
    h(doc, "3.2 Méthodologie utilisée", 2)
    p(doc, "La méthodologie adoptée repose sur une approche Agile. Le travail a été découpé en étapes : étude du contexte, analyse de l'existant, collecte des besoins, analyse fonctionnelle, conception, prototypage UI/UX, développement frontend, développement backend, intégration, tests et validation. Le prototypage a permis de clarifier les parcours utilisateur, tandis que GitHub a servi au suivi du code. Postman a été utilisé pour tester les routes API et MySQL Workbench pour inspecter la base de données.")
    h(doc, "3.3 Analyse des besoins fonctionnels", 2)
    table(doc, ["Besoin", "Description", "Acteur concerné"], [
        ("Authentification", "Connexion par identifiants, génération du token et redirection selon le rôle.", "Tous"),
        ("Gestion employés", "Création, modification et désactivation des comptes.", "ADMIN"),
        ("Gestion organisations", "Création, consultation et mise à jour des organisations clientes.", "SD, Manager en lecture"),
        ("Gestion contacts", "Enregistrement des contacts et rattachement à une organisation.", "SD, Manager en lecture"),
        ("Création tickets", "Création de tickets avec statut Pending.", "SD"),
        ("Consultation tickets", "Liste, recherche, filtre et détail des tickets autorisés.", "SD, Manager, PKI, IT"),
        ("Affectation tickets", "Affectation aux services IT ou PKI avec historique.", "SD"),
        ("Suivi statuts", "Passage entre Pending, In Progress, Warning, Critical et Resolved.", "Selon règles backend"),
        ("Commentaires", "Ajout et consultation de commentaires liés aux tickets.", "Acteurs autorisés"),
        ("Messagerie", "Échanges internes en temps réel dans des rooms associées aux tickets.", "Acteurs autorisés"),
        ("Meetings", "Création et suivi des réunions liées aux demandes.", "SD, Manager selon cas"),
        ("Inbox e-mail", "Consultation des e-mails clients importés par IMAP.", "SD"),
        ("Dashboard", "Indicateurs de suivi et supervision.", "ADMIN, SD, Manager, PKI, IT"),
        ("Accès selon rôle", "Contrôle frontend et backend des permissions.", "Tous"),
    ], widths=[2100, 5200, 2060], title="Tableau 4 : Besoins fonctionnels.")
    h(doc, "3.4 Analyse des besoins non fonctionnels", 2)
    table(doc, ["Besoin", "Description"], [
        ("Sécurité", "Protection des routes, authentification JWT, hachage bcrypt et contrôle RBAC."),
        ("Performance", "Chargement raisonnable des listes, requêtes SQL ciblées et tableaux de bord calculés."),
        ("Maintenabilité", "Organisation du backend en routes, controllers, services et repositories."),
        ("Évolutivité", "Possibilité d'ajouter de nouveaux statuts, indicateurs, notifications ou règles métier."),
        ("Traçabilité", "Conservation des actions, affectations, commentaires, messages et meetings."),
        ("Ergonomie", "Interface claire, sidebar par rôle, toasts et états loading/error."),
        ("Disponibilité", "Application web accessible depuis un navigateur moderne en environnement local."),
        ("Fiabilité", "Validation des champs, gestion des erreurs et cohérence des statuts."),
        ("Souveraineté des données", "Conservation locale des données dans MySQL sous contrôle de l'organisme."),
        ("Confidentialité", "Protection des informations clients, employés, incidents et pièces jointes."),
    ], widths=[2600, 6760], title="Tableau 5 : Besoins non fonctionnels.")
    h(doc, "3.5 Identification des acteurs", 2)
    table(doc, ["Acteur", "Rôle dans le système"], [
        ("ADMIN", "Gère les comptes employés et le tableau de bord administrateur ; ne traite pas les tickets métier."),
        ("SD - Service Delivery", "Organise les demandes entrantes, gère organisations, contacts, tickets, affectations, meetings, messages et inbox e-mail."),
        ("Manager", "Supervise l'activité, consulte les données en lecture seule et répond uniquement aux meetings qui lui sont affectés."),
        ("PKI", "Traite les tickets liés aux certificats, révocations, vérifications et opérations techniques PKI accessibles."),
        ("IT", "Traite les incidents techniques ou applicatifs affectés au service IT."),
    ], widths=[2100, 7260], title="Tableau 3 : Acteurs et permissions.")
    h(doc, "3.6 Matrice RBAC", 2)
    table(doc, ["Fonction", "ADMIN", "SD", "Manager", "PKI", "IT"], [
        ("Gérer employés", "Oui", "Selon route actuelle", "Non", "Non", "Non"),
        ("Gérer organisations", "Non", "Oui", "Lecture", "Non", "Non"),
        ("Gérer contacts", "Non", "Oui", "Lecture", "Non", "Non"),
        ("Créer tickets", "Non", "Oui", "Non", "Non", "Non"),
        ("Consulter tickets", "Non", "Oui", "Lecture globale", "Tickets PKI", "Tickets IT"),
        ("Affecter tickets", "Non", "Oui", "Lecture", "Non", "Non"),
        ("Modifier statut", "Non", "Oui selon implémentation", "Non", "Non", "Non"),
        ("Commenter", "Non", "Oui si autorisé", "Lecture", "Tickets PKI", "Tickets IT"),
        ("Accéder messages", "Non", "Oui", "Lecture", "Rooms PKI", "Rooms IT"),
        ("Accéder meetings", "Non", "Oui", "Lecture / réponse limitée", "Selon accès", "Selon accès"),
        ("Consulter dashboard", "Admin", "Oui", "Oui", "Oui", "Oui"),
        ("Consulter inbox e-mail", "Non", "Oui", "Non", "Non", "Non"),
    ], widths=[2100, 1250, 1550, 1650, 1405, 1405])
    h(doc, "3.7 Diagramme de cas d'utilisation", 2)
    placeholder(doc, "Emplacement réservé au diagramme de cas d'utilisation", "Le diagramme de cas d'utilisation sera inséré dans cette section afin de représenter les interactions entre les différents acteurs du système et les fonctionnalités principales de l'application. Il permettra de visualiser les droits associés à chaque rôle et la séparation des responsabilités entre l'administrateur, le Service Delivery, le Manager, le service PKI et le service IT.")
    h(doc, "3.8 Description textuelle des cas d'utilisation principaux", 2)
    cases = [
        ("S'authentifier", "Tous", "Accéder à l'application selon le rôle.", "L'utilisateur possède un compte actif.", "Saisie username/password, vérification backend, génération JWT, redirection.", "Accès à la section autorisée."),
        ("Gérer les employés", "ADMIN", "Créer, modifier ou désactiver un compte.", "ADMIN authentifié.", "Saisie des données, validation, hachage du mot de passe, enregistrement.", "Compte employé mis à jour."),
        ("Gérer les organisations", "SD", "Centraliser les organisations clientes.", "SD authentifié.", "Création ou modification d'une organisation.", "Organisation disponible pour contacts et tickets."),
        ("Gérer les contacts", "SD", "Rattacher des interlocuteurs aux organisations.", "Organisation existante.", "Saisie du contact et association.", "Contact enregistré."),
        ("Créer un ticket", "SD", "Enregistrer une demande ou un incident.", "Organisation ou contexte disponible.", "Saisie du ticket, statut Pending, création de room.", "Ticket créé et traçable."),
        ("Affecter un ticket", "SD", "Orienter le ticket vers IT ou PKI.", "Ticket non résolu.", "Sélection du service, contrôle backend, historique.", "Ticket affecté."),
        ("Consulter les tickets", "Acteur autorisé", "Suivre les tickets accessibles.", "Utilisateur authentifié.", "Chargement de la liste filtrée par rôle.", "Tickets affichés."),
        ("Modifier le statut", "SD", "Mettre à jour l'état du ticket.", "Ticket non résolu.", "Choix d'un statut valide.", "Statut mis à jour."),
        ("Ajouter un commentaire", "Acteur autorisé", "Documenter une action ou décision.", "Ticket accessible et non résolu.", "Saisie et enregistrement du commentaire.", "Historique enrichi."),
        ("Échanger dans une room", "Acteur autorisé", "Communiquer autour d'un ticket.", "Room accessible.", "Connexion Socket.IO, envoi et diffusion.", "Message reçu par les membres autorisés."),
        ("Planifier un meeting", "SD", "Organiser une réunion de suivi.", "Utilisateur et ticket disponibles.", "Création du meeting, choix invité/salle.", "Meeting enregistré."),
        ("Consulter le dashboard", "Manager / SD / autres rôles", "Visualiser les indicateurs.", "Utilisateur authentifié.", "Chargement des KPI selon rôle.", "Indicateurs affichés."),
        ("Consulter l'inbox e-mail", "SD", "Voir les demandes reçues par e-mail.", "Configuration IMAP disponible.", "Synchronisation et affichage des e-mails.", "E-mails consultables."),
    ]
    table(doc, ["Cas", "Acteur principal", "Objectif", "Précondition", "Scénario nominal", "Résultat attendu"], cases, widths=[1300, 1300, 1800, 1600, 2100, 1260])
    h(doc, "3.9 Diagramme de classes", 2)
    placeholder(doc, "Emplacement réservé au diagramme de classes", "Le diagramme de classes sera ajouté dans cette partie afin de représenter la structure statique du système, les principales entités manipulées, leurs attributs essentiels et leurs relations.")
    h(doc, "3.10 Description des entités principales", 2)
    table(doc, ["Entité", "Rôle"], [
        ("Employee", "Représente un utilisateur interne, son identité, son service, son statut et ses informations de connexion."),
        ("Organization", "Représente une organisation cliente ou partenaire."),
        ("Contact", "Représente un interlocuteur rattaché à une organisation."),
        ("Ticket", "Représente une demande ou un incident suivi par le système."),
        ("Room", "Représente l'espace de discussion lié à un ticket."),
        ("Message", "Représente un message envoyé dans une room."),
        ("Comment", "Représente une note structurée liée à un ticket."),
        ("Meeting", "Représente une réunion de suivi avec organisateur, invité et éventuellement ticket."),
        ("ActivityLog", "Représente une action significative conservée pour la traçabilité."),
        ("ClientEmail", "Représente un e-mail client importé dans l'inbox SD."),
        ("Attachment", "Représente une pièce jointe associée à un e-mail."),
        ("AssignmentHistory", "Représente l'historique d'affectation IT/PKI d'un ticket."),
    ], widths=[2300, 7060])
    h(doc, "3.11 Diagrammes de séquence", 2)
    placeholder(doc, "Emplacement réservé au diagramme de séquence 1 : authentification sécurisée", "Ce flux présentera la saisie des identifiants, la vérification du mot de passe, le contrôle du statut Active, la génération du JWT et la redirection selon le rôle.")
    placeholder(doc, "Emplacement réservé au diagramme de séquence 2 : création et affectation d'un ticket", "Ce flux présentera la création d'un ticket par le Service Delivery, son enregistrement en base, la création de la room associée et son affectation éventuelle à IT ou PKI.")
    placeholder(doc, "Emplacement réservé au diagramme de séquence 3 : échange de messages en temps réel", "Ce flux présentera la connexion Socket.IO, l'accès à la room, l'envoi du message, son stockage et sa diffusion aux acteurs autorisés.")
    h(doc, "3.12 Architecture générale du système", 2)
    p(doc, "L'architecture générale est de type client-serveur. L'utilisateur interagit avec un frontend React et Vite. Le frontend communique avec l'API Express.js au moyen de requêtes HTTP JSON contenant un token Bearer. Les échanges temps réel sont assurés par Socket.IO. Les services métier appliquent les règles fonctionnelles et les repositories exécutent les requêtes SQL vers MySQL. Un service Gmail IMAP permet d'importer les e-mails entrants du Service Delivery.")
    placeholder(doc, "Emplacement réservé au schéma d'architecture générale", "Le schéma d'architecture visuel sera ajouté ultérieurement. Il devra montrer l'utilisateur, le frontend React/Vite, l'API Express.js, Socket.IO, les services métier, les repositories SQL, MySQL et le service Gmail IMAP.")
    h(doc, "3.13 Modèle relationnel", 2)
    table(doc, ["Table", "Rôle dans le système"], [
        ("employees", "Stocke les employés, identifiants, services et statuts."),
        ("organizations", "Stocke les organisations clientes."),
        ("contacts", "Stocke les contacts rattachés aux organisations."),
        ("tickets", "Stocke les demandes, statuts, niveaux et informations de suivi."),
        ("rooms", "Stocke les conversations liées aux tickets."),
        ("messages", "Stocke les messages envoyés dans les rooms."),
        ("room_message_reads", "Stocke le suivi lu/non lu des messages par employé."),
        ("comments", "Stocke les commentaires historisés des tickets."),
        ("ticket_assignment_history", "Stocke l'historique des affectations IT/PKI."),
        ("meetings", "Stocke les réunions de suivi."),
        ("activity_logs", "Stocke les actions significatives du système."),
        ("client_emails", "Stocke les e-mails clients importés."),
        ("client_email_attachments", "Stocke les pièces jointes des e-mails."),
        ("client_email_reads", "Stocke l'état de lecture des e-mails par employé."),
    ], widths=[2800, 6560], title="Tableau 6 : Tables principales de la base de données.")
    h(doc, "3.14 Contrats d'API et intégration système", 2)
    p(doc, "L'intégration frontend/backend repose sur des routes REST échangeant des données JSON. Les requêtes sécurisées utilisent l'en-tête Authorization avec un Bearer token. Les réponses suivent une structure normalisée indiquant le succès, le message et les données éventuelles. Les erreurs utilisent des codes HTTP adaptés : 400 pour les données invalides, 401 pour l'authentification, 403 pour l'autorisation, 404 pour une ressource inexistante et 500 pour une erreur serveur.")
    table(doc, ["Endpoint représentatif", "Fonction", "Sécurité principale"], [
        ("POST /api/auth/login", "Connexion", "Identifiants valides et compte actif"),
        ("GET /api/employees/me", "Profil courant", "JWT valide"),
        ("POST /api/tickets", "Création ticket", "Service Delivery"),
        ("PUT /api/tickets/:id/assign", "Affectation IT/PKI", "SD, ticket non résolu"),
        ("POST /api/tickets/:id/comments", "Ajout commentaire", "Accès ticket et non résolu"),
        ("GET /api/rooms", "Rooms accessibles", "Utilisateur authentifié et autorisé"),
        ("GET /api/dashboard/manager", "Dashboard manager", "Manager"),
        ("GET /api/client-emails", "Inbox e-mail", "Service Delivery"),
    ], widths=[2600, 3400, 3360])
    h(doc, "3.15 Conclusion du chapitre", 2)
    p(doc, "Ce chapitre a présenté l'analyse fonctionnelle et non fonctionnelle, les acteurs, les règles RBAC, les cas d'utilisation, les entités et l'architecture du système. La conception proposée prépare une implémentation modulaire, sécurisée et adaptée aux responsabilités internes de l'AGCE.")
    page_break(doc)


def chapter_four(doc):
    h(doc, "Chapitre 4 : Réalisation, implémentation et tests", 1)
    h(doc, "4.1 Introduction du chapitre", 2)
    p(doc, "Ce chapitre présente la mise en œuvre technique de l'application AGCE CRM. Il décrit l'environnement de développement, l'installation, l'implémentation du frontend, l'implémentation du backend, la sécurité, la messagerie temps réel, l'import des e-mails, la base de données, les tests et les résultats obtenus.")
    h(doc, "4.2 Environnement matériel", 2)
    table(doc, ["Élément", "Description prudente"], [
        ("Ordinateur portable", "Machine de développement utilisée pour exécuter le frontend, le backend et la base de données."),
        ("Processeur", "Processeur compatible avec l'exécution de Node.js, MySQL et navigateur moderne."),
        ("Mémoire", "Mémoire suffisante pour le développement web local et les tests manuels."),
        ("Système d'exploitation", "Environnement macOS ou Windows selon la machine de développement."),
    ], widths=[2600, 6760])
    h(doc, "4.3 Environnement logiciel", 2)
    table(doc, ["Outil", "Rôle"], [
        ("Visual Studio Code", "Édition du code frontend et backend."),
        ("Node.js", "Exécution de l'environnement JavaScript."),
        ("npm", "Installation et gestion des dépendances."),
        ("MySQL Server", "Stockage relationnel des données."),
        ("MySQL Workbench", "Consultation et administration de la base."),
        ("Navigateur web", "Exécution et test de l'interface utilisateur."),
        ("Postman", "Test des endpoints API."),
        ("Git/GitHub", "Versionnement et suivi du code."),
    ], widths=[2600, 6760], title="Tableau 7 : Environnement logiciel.")
    h(doc, "4.4 Installation et configuration", 2)
    p(doc, "L'installation commence par Node.js et MySQL Server. Une base de données nommée agce_crm est créée dans MySQL. Le frontend et le backend sont installés séparément avec npm install. Le frontend est lancé avec npm run dev dans le dossier frontend_ticket_gestion, tandis que le backend est lancé avec npm run dev dans le dossier backend_ticket_gestion.")
    code_block(doc, "PORT=2300\nJWT_SECRET=une_cle_secrete_longue_et_privee\nDB_HOST=localhost\nDB_NAME=agce_crm\nDB_USER=utilisateur_mysql\nDB_PWD=mot_de_passe_mysql\nGMAIL_IMAP_USER=adresse_de_reception_sd@gmail.com\nGMAIL_APP_PASSWORD=mot_de_passe_application_gmail\nGMAIL_SYNC_INTERVAL_MS=60000")
    h(doc, "4.5 Implémentation du frontend", 2)
    p(doc, "Le frontend est développé avec React et Vite. Il est organisé autour de composants réutilisables, de routes protégées et d'une sidebar adaptée au rôle de l'utilisateur. Le composant RoleBasedRoute vérifie le token, le rôle et l'état actif du compte avant d'autoriser l'accès à une page.")
    p(doc, "L'interface comprend des dashboards, des formulaires, des listes, des modals, des toasts de notification, des pages d'erreur et une page Access Denied. Le thème sombre/clair est géré par des variables CSS et une préférence locale associée à l'utilisateur.")
    h(doc, "4.6 Concepts React utilisés", 2)
    table(doc, ["Concept", "Utilisation dans le projet"], [
        ("Composants et props", "Découpage de l'interface en blocs réutilisables et transmission des données parent-enfant."),
        ("useState", "Gestion des formulaires, chargements, erreurs, onglet actif, ticket sélectionné, modal et thème."),
        ("useEffect", "Chargement des données, vérification du compte actif, récupération des tickets, room Socket.IO, redirection et thème."),
        ("useRef", "Scroll automatique des messages, clic extérieur, connexion socket et room active."),
        ("useMemo", "Mémorisation des filtrages, KPI dashboard, tickets filtrés, room sélectionnée et route de retour."),
        ("useCallback", "Mémorisation de joinRoom, sendMessage, loadRooms et chargements partagés."),
        ("Événements", "Gestion des clics, soumissions de formulaires et changements de champs."),
        ("Affichage conditionnel", "Boutons selon rôle, loading, error, état vide, statut ticket, room non lue et onglet actif."),
    ], widths=[2500, 6860])
    h(doc, "4.7 Implémentation du backend", 2)
    p(doc, "Le backend repose sur Express.js. Le fichier index.js assure le démarrage du serveur HTTP, l'initialisation de Socket.IO, les schémas complémentaires et le service Gmail. Le fichier app.js monte les routes Express. La connexion MySQL est centralisée dans config/db.js.")
    table(doc, ["Couche", "Responsabilité"], [
        ("Routes", "Déclaration des endpoints, méthodes HTTP et middlewares."),
        ("Controllers", "Réception des requêtes et formatage des réponses."),
        ("Services", "Règles métier, validations fonctionnelles et orchestration."),
        ("Repositories", "Requêtes SQL vers MySQL."),
        ("Middlewares", "Vérification JWT, rôles, accès aux tickets et blocage des actions interdites."),
    ], widths=[2500, 6860])
    h(doc, "4.8 Authentification et sécurité", 2)
    p(doc, "La sécurité est un élément central du projet. L'authentification repose sur un username et un mot de passe. Les mots de passe sont hachés avec bcrypt. Après vérification, le backend génère un JWT contenant les informations nécessaires à la session. Le token est stocké côté client et envoyé dans l'en-tête Authorization des requêtes protégées.")
    p(doc, "Le middleware auth.js vérifie le token, son expiration, sa présence dans la blacklist et l'état actif du compte. Les middlewares requireServiceDelivery, requireManager, requireServiceDeliveryOrManager, requireTicketAccess et blockResolvedTicket appliquent les règles d'autorisation. Une requête non autorisée est refusée par le backend, généralement avec un code 403.")
    p(doc, "La sécurité ne repose donc pas uniquement sur le masquage des boutons dans l'interface. Le frontend adapte les routes, le menu et les actions visibles, mais le backend contrôle réellement le rôle, l'accès à la ressource et l'état du ticket.")
    h(doc, "4.9 Messagerie temps réel", 2)
    p(doc, "La messagerie temps réel utilise Socket.IO. Chaque ticket dispose d'une room de conversation accessible uniquement aux acteurs autorisés. Lorsqu'un utilisateur ouvre une conversation, le frontend rejoint la room correspondante. Les messages sont enregistrés en base puis diffusés aux membres connectés. Le système conserve également l'état lu/non lu par employé.")
    h(doc, "4.10 Import des e-mails clients", 2)
    p(doc, "L'inbox e-mail du Service Delivery est intégrée au système au moyen d'une connexion Gmail IMAP. ImapFlow assure la communication avec la boîte e-mail, tandis que mailparser extrait le sujet, l'expéditeur, le contenu et les pièces jointes. Les e-mails reconnus peuvent être stockés dans MySQL et consultés par l'équipe SD avec un statut lu/non lu individuel.")
    h(doc, "4.11 Base de données", 2)
    table(doc, ["Table", "Rôle"], [
        ("employees", "Comptes internes, services, identifiants et statut actif/inactif."),
        ("organizations", "Organisations clientes."),
        ("contacts", "Contacts liés aux organisations."),
        ("tickets", "Demandes et incidents avec statut, niveau et contexte."),
        ("rooms", "Conversations associées aux tickets."),
        ("messages", "Messages temps réel."),
        ("room_message_reads", "État de lecture des messages."),
        ("comments", "Commentaires liés aux tickets."),
        ("ticket_assignment_history", "Historique des affectations."),
        ("meetings", "Réunions de suivi."),
        ("activity_logs", "Traçabilité des actions importantes."),
        ("client_emails", "E-mails clients importés."),
        ("client_email_attachments", "Pièces jointes."),
        ("client_email_reads", "Lecture individuelle des e-mails."),
    ], widths=[2800, 6560])
    h(doc, "4.12 Interfaces principales de l'application", 2)
    h(doc, "Aperçu de l'interface de l'application", 3)
    t = doc.add_table(rows=4, cols=2)
    t.style = "Table Grid"
    entries = [
        ("Capture 1\nInterface d'accueil / Tableau de bord\n[Emplacement réservé à l'image]", "Capture 2\nGestion des tickets\n[Emplacement réservé à l'image]"),
        ("", ""),
        ("Capture 3\nGestion des messages / e-mails\n[Emplacement réservé à l'image]", "Capture 4\nDétails d'un ticket / meetings\n[Emplacement réservé à l'image]"),
        ("", ""),
    ]
    for r, row in enumerate(entries):
        for c, text in enumerate(row):
            cell = t.cell(r, c)
            cell.text = text
            set_cell_margins(cell, top=180, bottom=180, start=160, end=160)
            set_cell_shading(cell, PLACEHOLDER_FILL if text else "FFFFFF")
            for par in cell.paragraphs:
                par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in par.runs:
                    set_run_font(run, size=10.2, bold="[Emplacement" not in run.text, color=DARK_BLUE if "[Emplacement" not in run.text else MUTED)
    style_table(t, widths=[4680, 4680], header=False)
    md_lines.append("\n| Capture 1 | Capture 2 |\n|---|---|\n| Interface d'accueil / Tableau de bord<br>[Emplacement réservé à l'image] | Gestion des tickets<br>[Emplacement réservé à l'image] |\n\n")
    md_lines.append("| Capture 3 | Capture 4 |\n|---|---|\n| Gestion des messages / e-mails<br>[Emplacement réservé à l'image] | Détails d'un ticket / meetings<br>[Emplacement réservé à l'image] |\n\n")
    p(doc, "Ces captures d'écran permettront d'illustrer les principales interfaces développées dans l'application. Elles mettront en évidence l'ergonomie du système, la clarté de la navigation, la cohérence visuelle de la solution et l'organisation des fonctionnalités selon les rôles.")
    h(doc, "4.13 Tests fonctionnels", 2)
    table(doc, ["Test", "Scénario", "Résultat attendu", "Résultat obtenu", "Statut"], [
        ("TF01", "Connexion réussie", "Accès au dashboard du rôle", "Accès obtenu", "Validé"),
        ("TF02", "Connexion avec mauvais mot de passe", "Refus de connexion", "Erreur affichée", "Validé"),
        ("TF03", "Accès à une route interdite", "Access Denied ou redirection", "Accès bloqué", "Validé"),
        ("TF04", "Désactivation d'un employé", "Compte refusé", "Session invalidée", "Validé"),
        ("TF05", "Création d'organisation", "Organisation enregistrée", "Liste actualisée", "Validé"),
        ("TF06", "Création de contact", "Contact associé", "Contact visible", "Validé"),
        ("TF07", "Création de ticket", "Ticket Pending créé", "Ticket et room créés", "Validé"),
        ("TF08", "Affectation IT/PKI", "Historique enregistré", "Service affecté", "Validé"),
        ("TF09", "Modification statut", "Statut mis à jour", "Statut actualisé", "Validé"),
        ("TF10", "Ajout commentaire", "Commentaire conservé", "Historique enrichi", "Validé"),
        ("TF11", "Envoi message", "Diffusion temps réel", "Message reçu", "Validé"),
        ("TF12", "Création meeting", "Meeting Pending", "Meeting enregistré", "Validé"),
        ("TF13", "Consultation dashboard", "Indicateurs affichés", "KPI visibles", "Validé"),
    ], widths=[900, 2600, 2500, 2300, 1060], title="Tableau 8 : Tests fonctionnels.")
    h(doc, "4.14 Tests de sécurité", 2)
    table(doc, ["Test", "Scénario", "Résultat attendu", "Résultat obtenu"], [
        ("TS01", "Accès sans token", "Refus 401", "Requête refusée"),
        ("TS02", "Token invalide", "Refus 401", "Requête refusée"),
        ("TS03", "Token expiré", "Retour login ou refus", "Accès refusé"),
        ("TS04", "Compte désactivé", "Refus 403", "Accès bloqué"),
        ("TS05", "Rôle non autorisé", "Refus 403", "Action interdite"),
        ("TS06", "Modification d'un ticket résolu", "Blocage", "Action refusée"),
        ("TS07", "Accès PKI à un ticket IT", "Refus", "Accès interdit"),
        ("TS08", "Accès IT à un ticket PKI", "Refus", "Accès interdit"),
        ("TS09", "Vérification 403", "Message clair", "Erreur retournée"),
    ], widths=[900, 3300, 2700, 2460], title="Tableau 9 : Tests de sécurité.")
    h(doc, "4.15 Résultats obtenus", 2)
    p(doc, "La solution réalisée permet de centraliser les organisations, les contacts, les tickets, les commentaires, les messages, les meetings et les e-mails clients. Elle améliore l'organisation du support en structurant le cycle de vie des demandes. Elle renforce la traçabilité grâce aux historiques d'affectation et aux échanges conservés. Elle applique une séparation claire des rôles et offre au manager une visibilité globale sur l'activité.")
    p(doc, "Le système contribue également à la sécurité applicative par l'utilisation de JWT, bcrypt, middlewares de rôle, vérification du compte actif et blocage des tickets résolus. Les données étant conservées localement dans MySQL, la solution répond aussi à l'enjeu de souveraineté des données.")
    h(doc, "4.16 Difficultés rencontrées", 2)
    bullets(doc, [
        "structuration du code entre frontend, backend et base de données ;",
        "synchronisation des contrats API avec les composants React ;",
        "définition précise des permissions par rôle ;",
        "gestion de Socket.IO et des rooms autorisées ;",
        "sécurisation des routes sensibles ;",
        "intégration Gmail IMAP et traitement des pièces jointes ;",
        "cohérence des relations MySQL ;",
        "tests des accès selon les rôles et les états des tickets.",
    ])
    h(doc, "4.17 Limites du projet", 2)
    p(doc, "Le projet présente certaines limites. Les tests ont été principalement réalisés de manière manuelle et doivent être complétés par des tests automatisés. L'application n'a pas encore été évaluée sur un grand volume de données en production. La blacklist JWT est actuellement en mémoire et devrait être remplacée par un mécanisme persistant. Le monitoring, les sauvegardes automatiques, l'authentification multifactorielle, les statistiques prédictives et le déploiement industriel restent à renforcer.")
    h(doc, "4.18 Conclusion du chapitre", 2)
    p(doc, "Ce chapitre a présenté la réalisation technique de l'application AGCE CRM, son environnement, son implémentation, ses mécanismes de sécurité et ses tests. Les résultats confirment que la solution répond aux besoins principaux, tout en laissant des perspectives d'amélioration pour une version plus industrialisée.")
    page_break(doc)


def conclusion_and_biblio(doc):
    h(doc, "Conclusion générale et perspectives", 1)
    p(doc, "Ce mémoire a présenté la conception et la réalisation d'une application web sécurisée de gestion des incidents et des tickets des services de l'AGCE. Le projet s'inscrit dans un contexte où la confiance numérique, la traçabilité, la sécurité et la souveraineté des données constituent des exigences importantes.")
    p(doc, "La problématique initiale portait sur la centralisation des organisations, des contacts, des tickets, des messages, des réunions et des indicateurs de supervision, tout en garantissant une séparation stricte des rôles et un contrôle d'accès fiable. Pour y répondre, nous avons conçu une solution de type CRM/ticketing adaptée au contexte de l'AGCE.")
    p(doc, "Le travail réalisé couvre l'analyse du besoin, l'étude de l'existant, la conception, l'implémentation frontend et backend, la base de données, la sécurité, la messagerie temps réel, l'inbox e-mail, les dashboards et les tests. La solution obtenue améliore la centralisation des informations, la traçabilité des opérations, l'organisation du support et la visibilité managériale.")
    p(doc, "Les limites du projet concernent principalement l'absence de tests à grande échelle, l'industrialisation encore incomplète, le besoin d'un monitoring avancé, la persistance de la blacklist JWT et le renforcement des mécanismes de sécurité.")
    h(doc, "Perspectives", 2)
    bullets(doc, [
        "intégrer une authentification multifactorielle ;",
        "préparer un déploiement en production ;",
        "mettre en place des sauvegardes automatiques ;",
        "ajouter du monitoring et de la journalisation avancée ;",
        "générer des rapports PDF ou Excel ;",
        "améliorer la recherche avancée ;",
        "ajouter des notifications en temps réel ;",
        "renforcer l'intégration des e-mails ;",
        "améliorer les dashboards et les indicateurs prédictifs ;",
        "mettre en place une gestion avancée des SLA ;",
        "rendre persistante ou distribuée la blacklist JWT.",
    ])
    page_break(doc)

    h(doc, "Bibliographie", 1)
    refs = [
        "République Algérienne Démocratique et Populaire. Loi n° 15-04 relative à la signature et à la certification électroniques.",
        "République Algérienne Démocratique et Populaire. Décret exécutif n° 16-135 du 25 avril 2016 fixant l'organisation et le fonctionnement de l'AGCE.",
        "React. Documentation officielle de React. https://react.dev/",
        "Vite. Documentation officielle de Vite. https://vite.dev/",
        "Express.js. Documentation officielle d'Express. https://expressjs.com/",
        "MySQL. Documentation officielle de MySQL. https://dev.mysql.com/doc/",
        "Socket.IO. Documentation officielle de Socket.IO. https://socket.io/docs/",
        "JWT.io. Introduction aux JSON Web Tokens. https://jwt.io/",
        "bcrypt. Documentation du module bcrypt pour Node.js. https://www.npmjs.com/package/bcrypt",
        "OWASP Foundation. Authentication Cheat Sheet. https://cheatsheetseries.owasp.org/",
        "OWASP Foundation. Authorization Cheat Sheet. https://cheatsheetseries.owasp.org/",
        "Sandhu, R. S., Coyne, E. J., Feinstein, H. L., & Youman, C. E. (1996). Role-Based Access Control Models. Computer, 29(2), 38-47.",
        "Buttle, F., & Maklan, S. (2019). Customer Relationship Management: Concepts and Technologies. Routledge.",
        "AXELOS. ITIL Foundation: ITIL 4 Edition. The Stationery Office.",
        "Jira Service Management. Documentation officielle. https://support.atlassian.com/jira-service-management-cloud/",
        "Zendesk. Documentation officielle. https://support.zendesk.com/",
        "GLPI Project. Documentation officielle. https://glpi-project.org/documentation/",
        "osTicket. Documentation officielle. https://docs.osticket.com/",
    ]
    for ref in refs:
        par = doc.add_paragraph(style="List Number")
        r = par.add_run(ref)
        set_run_font(r, size=10.2)
        md_lines.append(f"1. {ref}\n")
    page_break(doc)


def annexes(doc):
    h(doc, "Annexes", 1)
    h(doc, "Annexe A : Hiérarchie et structure interne de l'AGCE", 2)
    p(doc, "Cette annexe est réservée à la présentation de l'organisation interne de l'AGCE et du département concerné par le projet.")
    placeholder(doc, "Emplacement réservé à l'organigramme de l'AGCE", "L'organigramme officiel ou simplifié sera inséré ici.")
    h(doc, "Annexe B : Produits, solutions et services de l'AGCE", 2)
    p(doc, "L'AGCE intervient dans les services de certification électronique, de signature numérique, d'authentification, de gestion des certificats, de publication des certificats, de listes de révocation, de vérification et d'horodatage. Ces services s'inscrivent dans la construction de la confiance numérique au niveau gouvernemental.")
    h(doc, "Annexe C : Cadre réglementaire et souveraineté des données", 2)
    p(doc, "La certification électronique et la signature électronique nécessitent un cadre réglementaire, des mécanismes de sécurité et une protection rigoureuse des données. Dans ce contexte, la souveraineté des données justifie le choix d'une solution locale, contrôlée et adaptée aux responsabilités de l'organisme.")
    h(doc, "Annexe D : Diagrammes UML", 2)
    placeholder(doc, "Emplacement réservé au diagramme de cas d'utilisation", "Diagramme à ajouter par l'étudiant.")
    placeholder(doc, "Emplacement réservé au diagramme de classes", "Diagramme à ajouter par l'étudiant.")
    placeholder(doc, "Emplacement réservé au diagramme de séquence d'authentification", "Diagramme à ajouter par l'étudiant.")
    placeholder(doc, "Emplacement réservé au diagramme de séquence de création d'un ticket", "Diagramme à ajouter par l'étudiant.")
    placeholder(doc, "Emplacement réservé au diagramme de séquence de messagerie temps réel", "Diagramme à ajouter par l'étudiant.")
    h(doc, "Annexe E : Captures d'écran supplémentaires", 2)
    placeholder(doc, "Emplacement réservé aux captures d'écran complémentaires", "Ajouter ici les interfaces de connexion, dashboard, tickets, messages, meetings et inbox e-mail.")
    h(doc, "Annexe F : Extraits de code importants", 2)
    for title in [
        "route d'authentification",
        "middleware JWT",
        "contrôle RBAC",
        "requête SQL",
        "composant RoleBasedRoute",
        "hook useChatRoom",
    ]:
        placeholder(doc, f"Emplacement réservé : {title}", "Insérer ici un extrait court et commenté du code correspondant.")


def main():
    WORK_DIR.mkdir(parents=True, exist_ok=True)
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    cover(doc)
    preliminaries(doc)
    navigation_pages(doc)
    introduction_generale(doc)
    chapter_one(doc)
    chapter_two(doc)
    chapter_three(doc)
    chapter_four(doc)
    conclusion_and_biblio(doc)
    annexes(doc)
    add_header_footer(doc)
    doc.save(DOCX_OUT)
    MD_OUT.write_text("\n".join(md_lines), encoding="utf-8")
    print(DOCX_OUT)
    print(MD_OUT)


if __name__ == "__main__":
    main()
